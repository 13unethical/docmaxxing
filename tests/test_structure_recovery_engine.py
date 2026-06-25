"""Tests for the AI Structure Recovery Engine."""

from __future__ import annotations

import io

from docx import Document

from services.structure_recovery_engine import (
    headings_exist,
    paragraphs_from_text,
    recover_structure,
)


HUMANIZED_RESEARCH_PAPER = """
Machine Learning in Healthcare Diagnostics

This paper examines how machine learning models support clinical decision-making in diagnostic imaging. The aim of this study is to explore recent advances and evaluate their reliability in hospital settings. The following sections outline the background, methods, and implications of automated diagnosis tools.

A substantial body of literature has investigated artificial intelligence in radiology (Smith, 2021). Previous studies demonstrate that convolutional neural networks can match specialist accuracy on selected tasks (Jones & Lee, 2020). According to Patel (2019), deployment challenges remain significant in resource-limited environments. Numerous studies highlight data quality as a primary constraint on model generalisation.

Participants were recruited from two metropolitan hospitals over six months. Data collection involved retrospective review of annotated scans and structured interviews with radiologists. Ethical approval was obtained before analysis began. A qualitative approach was combined with quantitative performance metrics to compare model outputs against expert consensus.

The results indicate that the proposed ensemble classifier achieved 91% sensitivity on the held-out test set. Table 1 summarises performance across disease categories. Findings reveal improved consistency compared with single-model baselines. Statistically significant gains were observed for early-stage lesion detection.

These findings suggest that ensemble methods may reduce false negatives in screening workflows. However, limitations include sample size and single-site data collection. Future research should examine multi-centre validation and clinician trust. In contrast to earlier pilots, this study incorporated structured user feedback.

In conclusion, machine learning can augment diagnostic workflows when paired with governance and clinician oversight. This paper has demonstrated both technical promise and organisational prerequisites for safe adoption.

Smith, J. (2021). AI in radiology. Medical Informatics, 12(3), 45–60.
Jones, A., & Lee, B. (2020). Deep learning for diagnosis. Health AI Journal, 8(1), 10–22.
Patel, R. (2019). Deployment barriers in hospitals. Digital Health Review, 4(2), 88–95.
""".strip()


def test_paragraphs_preserved_not_merged():
    result = recover_structure(
        text=HUMANIZED_RESEARCH_PAPER,
        document_type="research_paper",
        prefer_ai=False,
    )
    assert "error" not in result
    paragraphs = result["paragraphs"]
    assert len(paragraphs) >= 8
    assert result["paragraph_count"] == len(paragraphs)
    assert result["recovery_mode"] == "reconstructed"
    assert result["headings_present"] is False


def test_reconstructs_major_sections():
    result = recover_structure(
        text=HUMANIZED_RESEARCH_PAPER,
        document_type="research_paper",
        prefer_ai=False,
    )
    titles = [node["title"].lower() for node in result["structure_tree"]]
    joined = " ".join(titles)
    assert "introduction" in joined or "title" in joined
    assert "methodology" in joined or "results" in joined
    assert "conclusion" in joined
    assert "references" in joined


def test_section_confidence_scores():
    result = recover_structure(
        text=HUMANIZED_RESEARCH_PAPER,
        document_type="research_paper",
        prefer_ai=False,
    )
    for node in result["structure_tree"]:
        assert 0 <= node["confidence"] <= 1
        assert node.get("paragraph_indices")
    assert result["overall_confidence"] > 0


def test_preserves_existing_headings():
    text = (
        "Research Title\n\n"
        "Introduction\n\n"
        "This essay introduces the topic.\n\n"
        "Methodology\n\n"
        "We surveyed fifty students.\n\n"
        "Conclusion\n\n"
        "In conclusion, the study was informative.\n\n"
        "References\n\n"
        "Author, A. (2024). Sample. Journal, 1(1), 1–10."
    )
    result = recover_structure(text=text, document_type="essay")
    assert result["headings_present"] is True
    assert result["recovery_mode"] == "preserved"
    titles = [n["title"] for n in result["structure_tree"]]
    assert "Introduction" in titles
    assert "Methodology" in titles


def test_word_heading_styles_detected():
    doc = Document()
    doc.add_heading("Chapter Title", 0)
    doc.add_paragraph("This chapter introduces the dissertation topic and scope.")
    doc.add_heading("Introduction", 1)
    doc.add_paragraph("The aim of this chapter is to frame the research problem.")
    doc.add_heading("Conclusion", 1)
    doc.add_paragraph("In conclusion, the chapter summarises key arguments.")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    doc = Document(buf)
    result = recover_structure(doc=doc, document_type="dissertation_chapter")
    assert result["headings_present"] is True
    assert result["recovery_mode"] == "preserved"
    assert len(result["structure_tree"]) >= 3


def test_infer_reflection_document_type():
    text = (
        "Weekly Learning Journal\n\n"
        "This week I reflected on my placement experience in the ward.\n\n"
        "I learned how communication affects patient trust. My experience "
        "showed me the importance of reflective practice in nursing.\n\n"
        "In conclusion, I will apply these insights in future shifts."
    )
    result = recover_structure(text=text, prefer_ai=False)
    assert result["inferred_document_type"] in ("reflection", "learning_journal")


def test_enriched_api_fields_on_heuristic_recovery():
    result = recover_structure(
        text=HUMANIZED_RESEARCH_PAPER,
        document_type="research_paper",
        prefer_ai=False,
    )
    assert "document_type" in result
    assert "sections" in result
    assert "headings" in result
    assert "confidence_scores" in result
    assert len(result["sections"]) == len(result["structure_tree"])
    paras = paragraphs_from_text(
        "Climate policy has evolved rapidly in the last decade.\n\n"
        "Governments now coordinate targets across multiple sectors and institutions.\n\n"
        "This paragraph discusses implementation challenges in developing economies.\n\n"
        "Another paragraph explains monitoring frameworks and reporting cycles."
    )
    assert headings_exist(paras) is False
