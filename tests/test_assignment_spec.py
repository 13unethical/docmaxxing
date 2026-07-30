"""Tests for AssignmentSpec — requirement-driven generation contract."""

from __future__ import annotations

from docx import Document

from services.assignment_formatting import _docx_from_markdown, _job_from_requirement
from services.assignment_spec import (
    build_assignment_spec,
    needs_expansion,
    validate_draft_against_spec,
)
from services.revision_engine.section_parser import render_sections


def _lj_requirement() -> dict:
    return {
        "id": "req-1",
        "project_id": "proj-1",
        "title": "Learning Journal",
        "assignment_type": "Learning Journal",
        "word_count": 1200,
        "required_sections": [
            "Cover page",
            "Introduction",
            "Journal Entry 1",
            "Journal Entry 2",
            "Journal Entry 3",
            "Journal Entry 4",
            "Reflection",
            "References",
        ],
        "section_word_budgets": {
            "Introduction": 100,
            "Journal Entry 1": 200,
            "Journal Entry 2": 200,
            "Journal Entry 3": 200,
            "Journal Entry 4": 200,
            "Reflection": 300,
        },
        "formatting": {
            "font_family": "Times New Roman",
            "font_size": "12",
            "line_spacing": "double-spaced",
            "margins": "1 inch from each sides",
            "alignment": "left",
        },
        "learning_outcomes": ["LO1", "LO5"],
    }


def test_build_assignment_spec_from_learning_journal_brief():
    spec = build_assignment_spec(_lj_requirement(), project_id="proj-1")
    assert spec.total_word_target == 1200
    assert spec.section_word_targets["Introduction"] == 100
    assert spec.section_word_targets["Reflection"] == 300
    assert spec.formatting.font_family == "Times New Roman"
    assert spec.formatting.font_size_pt == 12
    assert spec.formatting.line_spacing == 2.0
    assert spec.formatting.alignment == "left"
    assert spec.formatting.margins_inches == 1.0
    assert spec.formatting.cover_page_required is True
    assert spec.formatting.references_on_new_page is True
    cover = next(s for s in spec.sections if s.title == "Cover page")
    assert cover.writable is False and cover.target_words == 0


def test_validate_rejects_short_total_and_missing_sections():
    spec = build_assignment_spec(_lj_requirement())
    content = "## Introduction\n\nToo short.\n\n## Reflection\n\nAlso short."
    result = validate_draft_against_spec(content=content, spec=spec)
    assert result.passed is False
    assert result.total_passed is False
    assert any("Journal Entry 1" in m for m in result.missing_sections)
    assert any("too short" in b.lower() or "outside" in b.lower() for b in result.blocking_issues)


def test_validate_excludes_references_from_word_budget():
    """Brief '2000 words' means essay body — bibliography must not inflate the gate."""
    from services.assignment_spec.validate import count_body_words, count_words

    spec = build_assignment_spec(
        {
            "title": "Essay",
            "assignment_type": "Essay",
            "word_count": 2000,
            "required_sections": ["Introduction", "Body", "Conclusion", "References"],
            "section_word_budgets": {
                "Introduction": 300,
                "Body": 1400,
                "Conclusion": 300,
            },
        }
    )
    # ~2000 body words + a fat references block that would push a naive total to ~2500.
    intro = " ".join(["intro"] * 300)
    body = " ".join(["body"] * 1400)
    conclusion = " ".join(["end"] * 300)
    refs = " ".join(["Smith", "2020", "Journal", "Article", "Title", "Extra", "Words"] * 80)
    content = (
        f"## Introduction\n\n{intro}\n\n"
        f"## Body\n\n{body}\n\n"
        f"## Conclusion\n\n{conclusion}\n\n"
        f"## References\n\n{refs}\n"
    )
    assert count_words(content) > 2200
    assert 1800 <= count_body_words(content) <= 2200
    result = validate_draft_against_spec(content=content, spec=spec)
    assert result.total_passed is True
    assert result.reference_words > 0
    assert result.total_words == count_body_words(content)


def test_docx_from_markdown_peels_heading_from_body():
    """## Title\\nBody must not become one Heading 2 paragraph (bold leakage)."""
    doc = _docx_from_markdown(
        "Learning Journal",
        "## Introduction\nThis is body prose that must be Normal style.",
    )
    styles = [p.style.name for p in doc.paragraphs if p.text.strip()]
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    assert "Heading 1" in styles
    assert "Heading 2" in styles
    assert "Normal" in styles
    intro = next(p for p in doc.paragraphs if p.style and p.style.name == "Heading 2")
    assert intro.text == "Introduction"
    body = next(p for p in doc.paragraphs if p.style and p.style.name == "Normal")
    assert "body prose" in body.text
    assert not any(t.startswith("Introduction This") for t in texts)


def test_format_job_consumes_analyzer_formatting_fields():
    job = _job_from_requirement(_lj_requirement())
    assert job.font_family == "Times New Roman"
    assert job.font_size_pt == 12
    assert job.line_spacing == 2.0
    assert job.alignment == "left"
    assert job.margin_preset == "normal"


def test_render_sections_uses_blank_line_after_heading():
    rendered = render_sections(
        [{"title": "Introduction", "body": "Purpose of the journal."}]
    )
    assert rendered.startswith("## Introduction\n\nPurpose")


def test_needs_expansion_uses_ten_percent_tolerance():
    assert needs_expansion(89, 100) is True   # below 90 (±10%)
    assert needs_expansion(90, 100) is False
    assert needs_expansion(200, 200) is False
