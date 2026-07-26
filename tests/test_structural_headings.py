"""Structural heading detection — never promote in-sentence keywords."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from formatter import FormatJob
from formatter.document_reconstruction import reconstruct_blocks, reconstruct_document_before_format
from tests.conftest import run_format_pipeline

SHEIN = Path(
    "/Users/nazirov/Desktop/Reading-Summary-with-Opinion-on-Shein-accused-of-shaming-"
    "customers-into-buying-more-than-they-can-afford.docx"
)

TITLE = (
    'Reading Summary with Opinion on "Shein accused of \'shaming\' customers into '
    'buying more than they can afford"'
)


@pytest.mark.skipif(not SHEIN.exists(), reason="Shein reading-summary fixture missing")
def test_shein_title_preserved_and_no_in_sentence_headings():
    doc = Document(str(SHEIN))
    recon = reconstruct_document_before_format(doc, document_type="essay", prefer_ai=False)
    job = FormatJob(
        font_family="Times New Roman",
        font_size_pt=12,
        line_spacing=2.0,
        alignment="justify",
        first_line_indent=True,
        space_before_pt=0,
        space_after_pt=0,
        margin_preset="normal",
        page_number_position="bottom_center",
        auto_headings=True,
        heading_all_caps=False,
        auto_justify_refs=True,
        requirement_headings=True,
    )
    from formatter import format_document_full

    format_document_full(doc, job, recon.assignments)
    texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    styles = {
        p.text.strip(): (p.style.name if p.style else "")
        for p in doc.paragraphs
        if p.text.strip()
    }

    assert texts[0] == TITLE
    assert styles[TITLE].startswith("Heading")
    assert "Document" not in texts

    heading_texts = [t for t, style in styles.items() if style.startswith("Heading") and t != TITLE]
    for bad in ("introduction", "discussion", "references", "conclusion"):
        assert bad not in {h.lower() for h in heading_texts}

    # In-sentence "references" must remain prose, not a section heading.
    joined = "\n".join(texts)
    assert "The references used to prepare this analysis" in joined.replace("\n", " ")


def test_in_sentence_references_never_becomes_heading():
    text = (
        "Finally, fast fashion companies should act responsibly. "
        "The references used to prepare this analysis are described below."
    )
    blocks = reconstruct_blocks([text], document_type="essay")
    assert not any(b.kind == "heading" and "reference" in b.text.lower() for b in blocks)
    assert any("The references used" in b.text for b in blocks)


def test_standalone_introduction_is_heading():
    blocks = reconstruct_blocks(
        ["Introduction", "This essay examines digital marketing ethics."],
        document_type="essay",
    )
    assert blocks[0].kind in {"heading", "title"}
    assert blocks[0].text == "Introduction"
    assert blocks[1].text.startswith("This essay")


def test_merged_introduction_at_paragraph_start_still_splits():
    blocks = reconstruct_blocks(
        ["Introduction Artificial intelligence has transformed education."],
        document_type="essay",
    )
    assert blocks[0].text == "Introduction"
    assert blocks[1].text.startswith("Artificial intelligence")
