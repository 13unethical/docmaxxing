from __future__ import annotations

from pathlib import Path

from docx import Document

from formatter_v2.render.model import DocumentModel
from formatter_v2.spec import ParagraphRole
from formatter_v2.structure.from_heuristics import HeuristicsExtractor
from formatter_v2.structure.from_word_styles import WordStylesExtractor


FIXTURES_DIR = Path(__file__).resolve().parent / "fixtures"
TEXT_FIXTURE = FIXTURES_DIR / "test_essay.txt"
DOCX_FIXTURE = FIXTURES_DIR / "test_essay_styled.docx"


def _ensure_fixtures_exist() -> None:
    """Create tiny fixtures locally for extractor comparison."""
    FIXTURES_DIR.mkdir(parents=True, exist_ok=True)
    if TEXT_FIXTURE.is_file() and DOCX_FIXTURE.is_file():
        return

    lines = [
        "Introduction",
        "• Item one",
        "• Item two",
        "Table 1. Summary of results",
        "Conclusion",
        "Some body text.",
    ]

    TEXT_FIXTURE.write_text("\n".join(lines), encoding="utf-8")

    doc = Document()
    p = doc.add_paragraph(lines[0])
    p.style = "Heading 1"
    doc.add_paragraph(lines[1], style="Normal")
    doc.add_paragraph(lines[2], style="Normal")

    doc.add_paragraph(lines[3], style="Normal")
    p = doc.add_paragraph(lines[4])
    p.style = "Heading 1"
    doc.add_paragraph(lines[5], style="Normal")
    doc.save(str(DOCX_FIXTURE))


def _roles(model: DocumentModel) -> list[str]:
    return [b.role for b in model.body]


def test_word_styles_document_detects_bullet_list_in_normal_paragraphs() -> None:
    doc = Document()
    doc.add_paragraph("• Item one", style="Normal")
    model = WordStylesExtractor().extract(doc)
    assert model.body[0].role == ParagraphRole.LIST_BULLET


def test_word_styles_document_detects_table_caption_in_normal_paragraph() -> None:
    doc = Document()
    doc.add_paragraph("Table 1. Results are shown", style="Normal")
    model = WordStylesExtractor().extract(doc)
    assert model.body[0].role == ParagraphRole.TABLE_CAPTION


def test_word_styles_explicit_heading_is_not_overridden_by_heuristics() -> None:
    doc = Document()
    p = doc.add_paragraph("• Not a list, still a heading", style="Heading 1")
    assert p.style.name == "Heading 1"
    model = WordStylesExtractor().extract(doc)
    assert model.body[0].role == ParagraphRole.HEADING_1


def test_both_extractors_produce_same_roles_for_equivalent_input() -> None:
    _ensure_fixtures_exist()
    lines = TEXT_FIXTURE.read_text(encoding="utf-8").splitlines()
    doc = Document(str(DOCX_FIXTURE))

    heur = HeuristicsExtractor().extract(lines)
    word = WordStylesExtractor().extract(doc)

    assert _roles(heur) == _roles(word)

