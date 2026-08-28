"""Tests for validation-based check pipeline and weighted scoring."""

from __future__ import annotations

from pathlib import Path

from services.check_explanation import (
    _collect_allowed_numbers,
    _filter_compliance_claims_without_brief,
    _filter_sentences_with_unknown_numbers,
    explain_check_results,
)
from services.check_pipeline import run_check_pipeline
from services.check_requirements import normalize_requirements, parse_word_count_spec
from services.check_scoring import compute_readiness_score, validations_to_categories
from services.check_text import document_word_count, split_document_paragraphs
from services.check_validator import validate_all_requirements
from services.document_checker import check_document

FIXTURES = Path(__file__).resolve().parent / "fixtures"


def test_parse_word_count_range():
    wmin, wmax, conf = parse_word_count_spec("1800-2200 words")
    assert wmin == 1800
    assert wmax == 2200
    assert conf >= 0.9


def test_short_document_gets_low_readiness_score():
    requirements = (
        "Write between 1800 and 2200 words. Use APA 7. "
        "Include 10 peer-reviewed articles. "
        "Required sections: Introduction, Body 1, Body 2, Body 3, Counterargument, Conclusion."
    )
    text = "Hello.\n\nAI is good.\n\nThanks."
    parsed = {
        "word_count": "1800-2200 words",
        "citation_style": "APA",
        "references_required": True,
        "required_sections": [
            "Introduction",
            "Body 1",
            "Body 2",
            "Body 3",
            "Counterargument",
            "Conclusion",
        ],
        "confidence_score": 0.9,
    }
    result = check_document(
        text=text,
        requirements=requirements,
        document_type="essay",
        parsed_requirements=parsed,
    )
    assert result["score"] < 45
    validations = {v["id"]: v for v in result["validations"]}
    assert validations["word_count"]["completion_pct"] < 15
    assert validations["word_count"]["status"] == "FAIL"
    assert validations["references"]["detected"] == "0"
    assert validations["references"]["status"] in ("FAIL", "PARTIAL")


def test_word_count_completion_math():
    req = normalize_requirements(
        "1800-2200 words",
        parsed_payload={"word_count": "1800-2200 words"},
    )
    metrics = {"word_count": 134}
    validations = validate_all_requirements(req, metrics)
    wc = next(v for v in validations if v["id"] == "word_count")
    assert wc["completion_pct"] == 7
    assert round(wc["points_earned"], 1) == 1.9


def test_weighted_score_sums_completion():
    validations = [
        {"weight": 25, "completion": 0.07, "status": "FAIL"},
        {"weight": 20, "completion": 0.0, "status": "FAIL"},
        {"weight": 15, "completion": 0.0, "status": "FAIL"},
        {"weight": 10, "completion": 0.9, "status": "PASS"},
    ]
    meta = compute_readiness_score(validations)
    total_weight = 70.0
    earned = 25 * 0.07 + 10 * 0.9
    assert meta["score"] == int(round(earned / total_weight * 100))
    assert meta["applicable_weight"] == total_weight
    assert meta["checks_applied"] == 4


def test_score_normalised_by_applicable_weight():
    validations = [
        {"weight": 25, "completion": 1.0, "status": "PASS"},
        {"weight": 10, "completion": 1.0, "status": "PASS"},
    ]
    meta = compute_readiness_score(validations)
    assert meta["score"] == 100
    assert meta["applicable_weight"] == 35.0
    assert meta["checks_applied"] == 2


def test_perfect_document_without_brief_scores_high():
    para = (
        "This paragraph develops one main idea with enough words to count as a substantive "
        "body paragraph for the academic structure heuristic used by the checker (Karimov, 2021)."
    )
    text = (
        "Introduction\n\n"
        + para
        + "\n\nLiterature Review\n\n"
        + para
        + "\n\nDiscussion\n\n"
        + para
        + "\n\nConclusion\n\n"
        + para
        + "\n\nReferences\n\n"
        "Karimov, A. (2021). Urban water management in Uzbekistan. Central Asian Studies, 14(2), 88-104.\n"
    )
    result = check_document(text=text, requirements="", document_type="essay")
    assert result["score"] >= 70
    assert result["applicable_weight"] >= 30
    assert result["score"] < 100 or result["checks_applied"] >= 4


def test_unchecked_category_is_not_reported_as_complete():
    result = check_document(
        text=(FIXTURES / "test_essay.txt").read_text(),
        requirements="",
        document_type="essay",
    )
    cats = result["categories"]
    assert cats["requirements_match"]["status"] == "NOT_CHECKED"
    assert cats["requirements_match"]["score"] is None
    assert cats["formatting"]["status"] == "NOT_CHECKED"
    checked = [k for k, v in cats.items() if v["status"] == "CHECKED"]
    assert "structure" in checked
    assert "references" in checked or "clarity_organization" in checked
    assert any(item["id"] == "sections" for item in result["not_checked"])
    assert result["checks_applied"] >= 4


def test_llm_summary_cannot_introduce_numbers_absent_from_requirements():
    allowed = _collect_allowed_numbers(
        structured={"peer_reviewed_refs": None},
        metrics={"word_count": 134, "reference_entries": 0},
        validations=[
            {
                "label": "References",
                "required": "1",
                "detected": "0",
                "completion_pct": 0,
                "weight": 15,
                "points_earned": 0,
                "points_possible": 15,
            }
        ],
        readiness_score=22,
    )
    polluted = (
        "Readiness score is 22/100. Add at least 10 scholarly references before submission. "
        "Word count is 134 words."
    )
    cleaned = _filter_sentences_with_unknown_numbers(polluted, allowed)
    assert "10 scholarly" not in cleaned
    assert "134" in cleaned or "word count" in cleaned.lower()


def test_text_extraction_reads_full_document():
    text = (FIXTURES / "test_essay.txt").read_text()
    words = document_word_count(text)
    paragraphs = split_document_paragraphs(text)
    assert words >= 14
    assert len(paragraphs) == 6
    result = check_document(text=text, requirements="", document_type="essay")
    assert result["meta"]["word_count"] == words
    assert result["meta"]["paragraph_count"] == len(paragraphs)
    assert "issues" not in result
    assert "priorities" not in result
    assert "summary" not in result
    assert result["meta"]["compliance_analysis"]["summary"]
    assert "health_score" not in result.get("structure_analysis", {})
    assert "outline_coverage_score" in result["structure_analysis"]


def test_sections_checklist():
    req = normalize_requirements(
        "Include Introduction and Conclusion",
        parsed_payload={"required_sections": ["Introduction", "Conclusion", "Body 1"]},
    )
    metrics = {
        "detected_sections": [
            {"title": "Introduction", "canonical": "introduction", "body_word_count": 80},
            {"title": "Conclusion", "canonical": "conclusion", "body_word_count": 40},
        ],
        "word_count": 500,
    }
    validations = validate_all_requirements(req, metrics)
    sections = next(v for v in validations if v["id"] == "sections")
    assert sections["detected"] == "2/3"
    assert sections["details"]["missing"] == ["Body 1"]


def test_empty_section_heading_does_not_count_as_present():
    empty = (
        "Introduction\n\n"
        "References\n\n"
        "Smith, J. (2020). Coastal governance in practice.\n"
    )
    empty_result = check_document(
        text=empty,
        requirements="Required sections: Introduction, Conclusion, References.",
        document_type="essay",
        parsed_requirements={"required_sections": ["Introduction", "Conclusion", "References"]},
    )
    empty_sections = next(v for v in empty_result["validations"] if v["id"] == "sections")
    empty_map = {c["section"]: c["present"] for c in empty_sections["details"]["checklist"]}
    assert empty_map["Introduction"] is False
    assert empty_map["Conclusion"] is False
    assert empty_map["References"] is True

    body = " ".join(["Climate"] * 40)
    filled = (
        f"1. Introduction\n\n{body}\n\n"
        "7 References\n\n"
        "Karimov, A. (2021). Urban water management in Uzbekistan. Central Asian Studies, 14(2), 88-104.\n"
    )
    filled_result = check_document(
        text=filled,
        requirements="Required sections: Introduction, References.",
        document_type="essay",
        parsed_requirements={"required_sections": ["Introduction", "References"]},
    )
    filled_sections = next(v for v in filled_result["validations"] if v["id"] == "sections")
    filled_map = {c["section"]: c["present"] for c in filled_sections["details"]["checklist"]}
    assert filled_map["Introduction"] is True
    assert filled_map["References"] is True


def test_pipeline_action_plan_orders_by_gain():
    text = "Intro\n\nBody text here with some words.\n\nConclusion"
    requirements = "1800-2200 words. 10 peer reviewed articles. APA."
    pipeline = run_check_pipeline(
        text=text,
        requirements=requirements,
        paragraphs=[p for p in text.split("\n\n") if p.strip()],
        doc=None,
        document_type="essay",
        parsed_requirements={"word_count": "1800-2200 words", "citation_style": "APA"},
    )
    plan = pipeline["action_plan"]
    assert plan
    assert plan[0]["estimated_improvement"] >= (plan[-1]["estimated_improvement"] if len(plan) > 1 else 0)


def _check_client():
    from app import app

    app.config["TESTING"] = True
    return app.test_client()


def test_uploaded_file_takes_priority_over_pasted_text():
    docx_path = FIXTURES / "test_essay_styled.docx"
    pasted_noise = " ".join(["noise"] * 500)
    client = _check_client()
    with docx_path.open("rb") as fh:
        res = client.post(
            "/api/check-document",
            data={
                "requirements": "",
                "pasted_text": pasted_noise,
                "document_type": "essay",
                "file": (fh, docx_path.name),
            },
            content_type="multipart/form-data",
        )
    assert res.status_code == 200
    data = res.get_json()
    from services.check_text import document_word_count
    from formatter.document_io import extract_text_from_document_bytes

    expected_words = document_word_count(
        extract_text_from_document_bytes(docx_path.read_bytes(), docx_path.name)
    )
    assert data["meta"]["word_count"] == expected_words
    assert data["meta"]["word_count"] >= 100
    assert data["meta"]["document_source"]["type"] == "file"
    assert data["meta"]["document_source"]["filename"] == docx_path.name
    assert "Checked:" in data["meta"]["document_source"]["label"]


def test_ai_review_makes_no_compliance_claims_without_brief():
    polluted = (
        "The essay meets the required section count but lacks references. "
        "We found two headings in the text."
    )
    cleaned = _filter_compliance_claims_without_brief(polluted)
    assert "required section count" not in cleaned.lower()
    assert "headings" in cleaned.lower()

    from unittest.mock import patch

    with patch("services.check_explanation.gemini_enabled", return_value=False):
        explanation = explain_check_results(
            requirements="",
            validations=[
                {
                    "id": "sections",
                    "label": "Required sections",
                    "required": "2 sections",
                    "detected": "2/2",
                    "completion_pct": 100,
                    "completion": 1.0,
                    "status": "PASS",
                    "weight": 20,
                }
            ],
            readiness_score=86,
            metrics={"word_count": 23, "heading_count": 2, "body_paragraph_count": 0},
            document_type="essay",
            has_assignment_brief=False,
        )
    summary = explanation["compliance_analysis"]["summary"].lower()
    assert "required section" not in summary
    assert "meets the required" not in summary
    assert "assignment brief" in summary
    assert explanation["source"] == "local"


def _citation_body(*sentences: str) -> str:
    pad = " ".join(["Climate"] * 40)
    extra = " ".join(sentences)
    return (
        f"1. Introduction\n\n{pad} {extra}\n\n"
        "6. Conclusion\n\n"
        f"{pad}\n\n"
    )


def test_uncited_reference_is_reported():
    text = (
        _citation_body("(Karimov, 2021).")
        + "7 References\n\n"
        "Karimov, A. (2021). Urban water management in Uzbekistan. Central Asian Studies, 14(2), 88-104.\n\n"
        "Nazarova, D. (2019). Water governance in transition. Tashkent: Fan Press.\n"
    )
    result = check_document(
        text=text,
        requirements="Use APA. Include a reference list.",
        document_type="essay",
        parsed_requirements={"citation_style": "APA", "references_required": True},
    )
    cite = next(v for v in result["validations"] if v["id"] == "in_text_citations")
    labels = [str(item.get("label") or "").lower() for item in cite["details"]["uncited"]]
    assert any("nazarova" in label for label in labels)
    assert "listed but not cited" in cite["detected"].lower()
    assert cite["status"] != "PASS"


def test_intext_citation_without_reference_is_reported():
    text = (
        _citation_body("(Smith, 2020).")
        + "7 References\n\n"
        "Karimov, A. (2021). Urban water management in Uzbekistan. Central Asian Studies, 14(2), 88-104.\n"
    )
    result = check_document(
        text=text,
        requirements="Use APA. Include a reference list.",
        document_type="essay",
        parsed_requirements={"citation_style": "APA", "references_required": True},
    )
    cite = next(v for v in result["validations"] if v["id"] == "in_text_citations")
    labels = [str(item.get("label") or "").lower() for item in cite["details"]["missing"]]
    assert any("smith" in label for label in labels)
    assert any("Cited but not in the list" in line for line in cite["details"]["mismatches"])
    assert "cited in text but missing from your reference list" in cite["detected"].lower()
    assert cite["status"] != "PASS"


def test_fully_matched_citations_pass():
    text = (
        _citation_body("Karimov (2021) and (Nazarova & Ibrahim, 2019).")
        + "7 References\n\n"
        "Karimov, A. (2021). Urban water management in Uzbekistan. Central Asian Studies, 14(2), 88-104.\n\n"
        "Nazarova, D. and Ibrahim, S. (2019). Water governance in transition. Tashkent: Fan Press.\n"
    )
    result = check_document(
        text=text,
        requirements="Use APA. Include a reference list.",
        document_type="essay",
        parsed_requirements={"citation_style": "APA", "references_required": True},
    )
    cite = next(v for v in result["validations"] if v["id"] == "in_text_citations")
    assert cite["details"]["uncited"] == []
    assert cite["details"]["missing"] == []
    assert cite["detected"] == "2 sources listed · 2 cited in text"
    assert cite["status"] == "PASS"
    assert cite["details"]["mode"] == "author_year"


def test_check_uses_v2_structure_extractor():
    from docx import Document
    from formatter_v2.structure.from_word_styles import document_has_structural_styles

    text = "Introduction\n\n" + ("Climate policy " * 20)
    result = check_document(text=text, requirements="", document_type="essay")
    assert result["meta"]["metrics"]["structure_extractor"] == "heuristics"

    doc = Document()
    doc.add_paragraph("Climate Adaptation", style="Title")
    doc.add_paragraph("Introduction", style="Heading 1")
    doc.add_paragraph("Sea-level rise reshapes municipal budgets and planning cycles.")
    assert document_has_structural_styles(doc)
    styled_text = "\n\n".join(p.text for p in doc.paragraphs if (p.text or "").strip())
    styled = check_document(text=styled_text, requirements="", document_type="essay", doc=doc)
    assert styled["meta"]["metrics"]["structure_extractor"] == "word_styles"


def test_reference_entries_are_not_counted_as_sections():
    pad = " ".join(["Climate"] * 40)
    text = (
        f"Introduction\n\n{pad}\n\n"
        "References\n\n"
        "1. Peppas NA, Bures P, Leobandung W, Ichikawa H. Hydrogels in pharmaceutical formulations. 2000.\n\n"
        "2. Hoffman AS. Hydrogels for biomedical applications. 2012.\n\n"
        "3. Caló E, Khutoryanskiy VV. Biomedical applications of hydrogels. 2015.\n"
    )
    result = check_document(text=text, requirements="", document_type="essay")
    titles = [str(s.get("title") or "") for s in result["meta"]["metrics"]["detected_sections"]]
    assert not any("Peppas" in title for title in titles)
    assert not any("Hoffman" in title for title in titles)
    assert any("reference" in title.lower() for title in titles)


def test_list_items_are_not_counted_as_sections():
    pad = " ".join(["Climate"] * 40)
    text = (
        f"Introduction\n\n{pad}\n\n"
        "The following operational checklist summarises fieldwork priorities:\n\n"
        "1. Time limitations: In the given 80 days, there was insufficient access.\n\n"
        "2. Cultural Boundaries: Data collection and interaction were constrained.\n\n"
        "3. Political and socio-economic issues: Most of the sites were closed.\n\n"
        "4. Budget Restrictions: For effective large scope sampling funds were limited.\n\n"
        f"Conclusion\n\n{pad}\n"
    )
    result = check_document(text=text, requirements="", document_type="essay")
    titles = [str(s.get("title") or "") for s in result["meta"]["metrics"]["detected_sections"]]
    assert not any("Time limitations" in title for title in titles)
    assert not any("Budget Restrictions" in title for title in titles)
    assert any("introduction" in title.lower() for title in titles)


def test_introduction_body_words_counted_correctly():
    intro_body = " ".join(["Adaptation"] * 40)
    text = (
        f"Introduction\n\n{intro_body}\n\n"
        "Background and context\n\n"
        + " ".join(["History"] * 20)
        + "\n\nConclusion\n\n"
        + " ".join(["Closing"] * 20)
        + "\n"
    )
    result = check_document(
        text=text,
        requirements="Required sections: Introduction.",
        document_type="essay",
        parsed_requirements={"required_sections": ["Introduction"]},
    )
    intro = next(
        s
        for s in result["meta"]["metrics"]["detected_sections"]
        if "introduction" in str(s.get("title") or "").lower()
    )
    assert int(intro["body_word_count"]) >= 30
    sections = next(v for v in result["validations"] if v["id"] == "sections")
    present = {c["section"]: c["present"] for c in sections["details"]["checklist"]}
    assert present["Introduction"] is True


def test_numeric_citation_style_matched_by_number():
    text = (
        _citation_body("Prior work used hydrogels [1] and later reviews (2).")
        + "References\n\n"
        "[1] Peppas NA, Bures P, Leobandung W, Ichikawa H. Hydrogels in pharmaceutical formulations. Eur J Pharm Biopharm. 2000.\n\n"
        "[2] Hoffman AS. Hydrogels for biomedical applications. Adv Drug Deliv Rev. 2012.\n"
    )
    result = check_document(
        text=text,
        requirements="Use IEEE. Include a reference list.",
        document_type="essay",
        parsed_requirements={"citation_style": "IEEE", "references_required": True},
    )
    cite = next(v for v in result["validations"] if v["id"] == "in_text_citations")
    assert cite["details"]["mode"] == "numeric"
    assert cite["details"]["listed"] == 2
    assert cite["details"]["cited"] == 2
    assert cite["details"]["uncited"] == []
    assert cite["details"]["missing"] == []
    assert cite["status"] == "PASS"


def test_author_year_style_matched_by_surname():
    text = (
        _citation_body("Karimov (2021) challenges the earlier account.")
        + "7 References\n\n"
        "Karimov, A. (2021). Urban water management in Uzbekistan. Central Asian Studies, 14(2), 88-104.\n"
    )
    result = check_document(
        text=text,
        requirements="Use APA. Include a reference list.",
        document_type="essay",
        parsed_requirements={"citation_style": "APA", "references_required": True},
    )
    cite = next(v for v in result["validations"] if v["id"] == "in_text_citations")
    assert cite["details"]["mode"] == "author_year"
    assert cite["details"]["uncited"] == []
    assert cite["details"]["missing"] == []
    assert cite["status"] == "PASS"


def test_bare_year_in_parentheses_is_not_a_citation():
    text = (
        _citation_body("The drought peaked (2020) across the basin, especially during 2004–2023.")
        + "7 References\n\n"
        "Karimov, A. (2021). Urban water management in Uzbekistan. Central Asian Studies, 14(2), 88-104.\n"
    )
    result = check_document(
        text=text,
        requirements="Use APA. Include a reference list.",
        document_type="essay",
        parsed_requirements={"citation_style": "APA", "references_required": True},
    )
    cite = next(v for v in result["validations"] if v["id"] == "in_text_citations")
    assert cite["details"]["mode"] == "author_year"
    assert cite["details"]["cited"] == 0
    labels = [str(item.get("label") or "").lower() for item in cite["details"]["uncited"]]
    assert any("karimov" in label for label in labels)
    assert cite["status"] != "PASS"


def test_unknown_citation_mode_reports_cannot_verify():
    pad = " ".join(["Climate"] * 40)
    text = (
        f"Introduction\n\n{pad}\n\n"
        "References\n\n"
        "Course packet, week 3 seminar notes.\n\n"
        "Library guide, further reading folder.\n\n"
        "Moodle collection, unmarked teaching slides.\n"
    )
    result = check_document(
        text=text,
        requirements="Use APA. Include a reference list.",
        document_type="essay",
        parsed_requirements={"citation_style": "APA", "references_required": True},
    )
    cite = next(v for v in result["validations"] if v["id"] == "in_text_citations")
    assert cite["status"] == "CANNOT_VERIFY"
    assert cite["detected"].lower() == "couldn't verify"
    assert cite["details"].get("verifiable") is False
    from services.check_scoring import compute_readiness_score

    meta = compute_readiness_score(result["validations"])
    applied = [
        v["id"]
        for v in result["validations"]
        if v.get("status") not in {"SKIP", "NOT_APPLICABLE", "NOT_CHECKED", "CANNOT_VERIFY"}
        and float(v.get("weight") or 0) > 0
    ]
    assert "in_text_citations" not in applied
    assert meta["applicable_weight"] == sum(
        float(v["weight"])
        for v in result["validations"]
        if v["id"] in applied
    )

