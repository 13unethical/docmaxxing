"""Tests for AI structure recovery normalization and document rebuild."""

from __future__ import annotations

import io

from docx import Document

from formatter.pipeline import FormatJob, format_document_full
from formatter.structure_rebuild import rebuild_document_from_recovery
from services.ai_structure_recovery import ai_result_to_recovery_payload


SAMPLE_AI_RESPONSE = {
    "document_type": "research_paper",
    "document_type_confidence": 0.88,
    "sections": [
        {
            "title": "Title",
            "heading_text": "Machine Learning in Healthcare",
            "confidence": 0.9,
            "paragraph_indices": [1],
            "insert_heading": False,
        },
        {
            "title": "Introduction",
            "heading_text": "Introduction",
            "confidence": 0.85,
            "paragraph_indices": [2],
            "insert_heading": True,
        },
        {
            "title": "Literature Review",
            "heading_text": "Literature Review",
            "confidence": 0.8,
            "paragraph_indices": [3],
            "insert_heading": True,
        },
        {
            "title": "Conclusion",
            "heading_text": "Conclusion",
            "confidence": 0.82,
            "paragraph_indices": [4],
            "insert_heading": True,
        },
        {
            "title": "References",
            "heading_text": "References",
            "confidence": 0.9,
            "paragraph_indices": [5],
            "insert_heading": True,
        },
    ],
}

ORIGINAL_PARAGRAPHS = [
    "Machine Learning in Healthcare",
    "This paper examines clinical decision-making in diagnostic imaging.",
    "Smith (2021) found that neural networks can match specialist accuracy.",
    "In conclusion, machine learning can augment diagnostic workflows.",
    "Smith, J. (2021). AI in radiology. Medical Informatics, 12(3), 45–60.",
]


def test_ai_result_to_recovery_payload_shape():
    result = ai_result_to_recovery_payload(
        SAMPLE_AI_RESPONSE,
        original_paragraphs=ORIGINAL_PARAGRAPHS,
    )
    assert result["recovery_mode"] == "ai_reconstructed"
    assert result["document_type"] == "research_paper"
    assert len(result["sections"]) == 5
    assert result["headings"] == [s["heading_text"] for s in result["sections"]]
    assert result["confidence_scores"] == [s["confidence"] for s in result["sections"]]
    assert result["ai_powered"] is True


def test_ai_paragraph_splits_expand_indices():
    data = dict(SAMPLE_AI_RESPONSE)
    data["paragraph_splits"] = [
        {
            "index": 3,
            "segments": [
                "Smith (2021) found that neural networks can match specialist accuracy.",
                "Jones (2020) reported similar findings in oncology.",
            ],
        }
    ]
    paragraphs = list(ORIGINAL_PARAGRAPHS)
    result = ai_result_to_recovery_payload(data, original_paragraphs=paragraphs)
    assert result["paragraph_count"] == 6
    lit = next(s for s in result["sections"] if s["title"] == "Literature Review")
    assert len(lit["paragraph_indices"]) == 2


def test_rebuild_document_inserts_headings():
    recovery = ai_result_to_recovery_payload(
        SAMPLE_AI_RESPONSE,
        original_paragraphs=ORIGINAL_PARAGRAPHS,
    )
    doc = Document()
    for para in ORIGINAL_PARAGRAPHS:
        doc.add_paragraph(para)

    rebuild_document_from_recovery(doc, recovery)
    texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    assert texts[0] == "Machine Learning in Healthcare"
    assert "Introduction" in texts
    assert "Literature Review" in texts
    assert "Conclusion" in texts
    assert "References" in texts
    assert texts.count("Introduction") == 1


def test_rebuild_preserves_all_body_paragraphs():
    recovery = ai_result_to_recovery_payload(
        SAMPLE_AI_RESPONSE,
        original_paragraphs=ORIGINAL_PARAGRAPHS,
    )
    doc = Document()
    for para in ORIGINAL_PARAGRAPHS:
        doc.add_paragraph(para)
    rebuild_document_from_recovery(doc, recovery)

    joined = "\n".join(p.text for p in doc.paragraphs)
    assert "clinical decision-making" in joined
    assert "Smith (2021)" in joined
    assert "Smith, J. (2021)" in joined


def test_ai_heading_level_survives_rebuild_into_word_style():
    recovery = {
        "recovery_mode": "ai_reconstructed",
        "ai_powered": True,
        "paragraphs": [
            "This heading has many words and would be body by heuristic rules",
            "Body paragraph with normal content.",
        ],
        "structure_tree": [
            {
                "title": "Methods",
                "heading_text": "This heading has many words and would be body by heuristic rules",
                "level": 2,
                "confidence": 0.88,
                "source": "ai_inferred",
                "paragraph_indices": [1],
                "insert_heading": False,
            },
            {
                "title": "Body",
                "level": 2,
                "confidence": 0.7,
                "source": "ai_inferred",
                "paragraph_indices": [2],
                "insert_heading": False,
            },
        ],
    }
    doc = Document()
    doc.add_paragraph("placeholder")
    apply_result = rebuild_document_from_recovery(doc, recovery)
    assert apply_result is not None

    job = FormatJob(
        font_family="Times New Roman",
        font_size_pt=12,
        line_spacing=1.5,
        alignment="left",
        first_line_indent=False,
        space_before_pt=0,
        space_after_pt=0,
        margin_preset="normal",
        page_number_position="none",
        auto_headings=True,
        heading_all_caps=False,
        auto_justify_refs=False,
    )
    format_document_full(
        doc,
        job,
        apply_result.assignments,
        recovery_mode="ai_reconstructed",
        ai_powered=True,
    )

    assert doc.paragraphs[0].style.name == "Heading 2"
