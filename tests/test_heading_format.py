"""Consistent academic heading detection and DOCX styling."""

from __future__ import annotations

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from formatter import FormatJob, format_document_full
from services.document_structure_engine import detect_heading_level


HEADINGS = [
    "1. Introduction",
    "2. Background of the organisation",
    "3. Management process and manager functions at different organisational levels",
    "4.Main problem",
    "5. Theories that help to explain the problem",
    "6. Recommendations",
    "7. Conclusion",
]


def test_numbered_section_headings_detected():
    for text in HEADINGS:
        assert detect_heading_level(text, True) == 2, text


def test_learning_journal_heading_lines_detected():
    lines = [
        "A digital innovation journal entry.Journal Entry – Digital Innovations.",
        "Journal Entry 2 Industrial Revolution",
        "The students make journal entries on Internationalisation and Globalisation of Business.",
        "This journal entry focuses on the history of branding and the actions of consumers.",
    ]
    for text in lines:
        assert detect_heading_level(text, True) == 2, text


def test_learning_journal_headings_get_16pt_not_source_bold():
    doc = Document()
    for line in [
        "A digital innovation journal entry.Journal Entry – Digital Innovations.",
        "This journal entry focuses on the history of branding and the actions of consumers.",
    ]:
        p = doc.add_paragraph(line)
        for run in p.runs:
            run.font.bold = True
            run.font.size = Pt(12)

    job = FormatJob(
        font_family="Times New Roman",
        font_size_pt=12,
        line_spacing=1.5,
        alignment="left",
        first_line_indent=False,
        space_before_pt=0,
        space_after_pt=0,
        margin_preset="normal",
        page_number_position="none",
        auto_headings=True,
        heading_all_caps=False,
        auto_justify_refs=False,
    )
    format_document_full(doc, job, None)
    for p in doc.paragraphs:
        assert p.runs[0].font.bold is True
        assert p.runs[0].font.size.pt == 16


def test_all_headings_left_bold_16pt():
    doc = Document()
    for heading in HEADINGS:
        doc.add_paragraph(heading)
        doc.add_paragraph("Body text for " + heading)

    job = FormatJob(
        font_family="Times New Roman",
        font_size_pt=12,
        line_spacing=1.5,
        alignment="justify",
        first_line_indent=False,
        space_before_pt=0,
        space_after_pt=0,
        margin_preset="normal",
        page_number_position="none",
        auto_headings=True,
        heading_all_caps=False,
        auto_justify_refs=False,
    )
    format_document_full(doc, job, None)

    for idx in range(0, len(HEADINGS) * 2, 2):
        heading_para = doc.paragraphs[idx]
        body_para = doc.paragraphs[idx + 1]
        assert heading_para.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.LEFT
        assert heading_para.runs[0].font.bold is True
        assert heading_para.runs[0].font.size.pt == 16
        assert heading_para.runs[0].font.name == "Times New Roman"
        assert body_para.runs[0].font.size.pt == 12
        assert body_para.runs[0].font.bold is not True
