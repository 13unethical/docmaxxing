"""Tests for the Delivery Engine."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from docx.shared import Pt

from services.ai_detection_engine import AIDetectionEngineService
from services.assignment_pipeline.service import AssignmentPipelineService
from services.assignment_project.models import ProjectStatus
from services.assignment_project.service import ProjectService
from services.blueprint_engine import BlueprintEngineService
from services.delivery_engine import DeliveryEngineService, DeliveryStatus
from services.delivery_engine.packager import resolve_client_format
from services.humanizer_engine import HumanizerEngineService
from services.research_engine import ResearchEngineService
from services.reviewer_engine import ReviewerEngineService
from services.revision_engine import RevisionEngineService
from services.writer_engine import WriterEngineService
from services.writer_engine.models import WriterSectionStatus
from tests.assignment_helpers import prepare_project_for_research


def _draft() -> dict:
    return {
        "id": "draft-1",
        "project_id": "proj-1",
        "title": "Climate Policy Essay",
        "content": "## Introduction\nThis essay discusses climate policy.\n\n## Conclusion\nPolicy remains contested.",
        "total_words": 1200,
    }


def _requirement() -> dict:
    return {
        "assignment_type": "Essay",
        "citation_style": "APA 7",
        "difficulty": "★★★☆☆",
        "word_count": 1200,
    }


def _research_plan() -> dict:
    return {"estimated_completion_time": "6–8 hours", "estimated_difficulty": "Moderate"}


def _blueprint() -> dict:
    return {"estimated_completion_time": "6–8 hours", "sections": []}


def _review_report() -> dict:
    return {"overall_score": 88, "passed": True}


def _detection_report() -> dict:
    return {"overall_ai_score": 9.5, "final_status": "passed"}


def test_resolve_client_format_defaults_to_docx():
    assert resolve_client_format({}) == "docx"
    assert resolve_client_format({"submission_format": "Word-processed"}) == "docx"
    assert resolve_client_format({"submission_format": "PDF"}) == "pdf"


def test_delivery_packages_client_file_only():
    engine = DeliveryEngineService()
    package = engine.prepare_package(
        final_draft=_draft(),
        requirement_json=_requirement(),
        research_plan=_research_plan(),
        blueprint=_blueprint(),
        review_report=_review_report(),
        detection_report=_detection_report(),
        project_id="proj-1",
        revision_attempts=1,
        humanization_attempts=6,
        completion_time="6–8 hours",
    )

    assert package.status == DeliveryStatus.READY
    assert len(package.files) == 1
    assert package.files[0].filename.endswith(".docx")
    assert package.client_format == "docx"
    assert package.project_summary.project_name == "Climate-Policy-Essay"
    assert package.project_summary.word_count == 1200
    assert package.project_summary.total_revisions == 1
    assert package.project_summary.total_humanization_attempts == 6
    assert package.project_summary.overall_review_score == 88
    assert package.project_summary.final_ai_score == 9.5
    assert package.package_download_url
    assert all(item.ready for item in package.files)

    root = Path(package.files[0].storage_path).parent
    assert (root / "debug" / "requirement.json").is_file()
    # Client-facing list must not include JSON artifacts.
    assert all(not f.filename.endswith(".json") for f in package.files)


def test_delivery_uses_formatted_docx(tmp_path):
    """Client download must ship Format Engine DOCX, not a markdown rebuild."""
    formatted = tmp_path / "formatted.docx"
    doc = Document()
    doc.add_heading("Formatted Title", level=1)
    p = doc.add_paragraph("Body with academic formatting.")
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
    doc.save(formatted)

    engine = DeliveryEngineService()
    package = engine.prepare_package(
        final_draft={
            **_draft(),
            "content": "## Document\nUnformatted fallback text that must not win.",
        },
        requirement_json=_requirement(),
        research_plan=_research_plan(),
        blueprint=_blueprint(),
        review_report=_review_report(),
        detection_report=_detection_report(),
        project_id="proj-formatted-zip",
        formatted_document_path=str(formatted),
    )

    docx_entry = next(f for f in package.files if f.file_type == "final_assignment_docx")
    assert docx_entry.label == "Formatted Assignment"
    on_disk = Document(docx_entry.storage_path)
    assert on_disk.paragraphs[0].text == "Formatted Title"
    assert "Unformatted fallback" not in "\n".join(p.text for p in on_disk.paragraphs)


def test_delivery_pdf_when_brief_asks(tmp_path):
    engine = DeliveryEngineService()
    package = engine.prepare_package(
        final_draft=_draft(),
        requirement_json={**_requirement(), "submission_format": "PDF"},
        research_plan=_research_plan(),
        blueprint=_blueprint(),
        review_report=_review_report(),
        detection_report=_detection_report(),
        project_id="proj-pdf",
    )
    assert package.client_format == "pdf"
    assert len(package.files) == 1
    assert package.files[0].filename.endswith(".pdf")


def test_delivery_requires_all_inputs():
    engine = DeliveryEngineService()
    with pytest.raises(ValueError):
        engine.prepare_package(
            final_draft={},
            requirement_json=_requirement(),
            research_plan=_research_plan(),
            blueprint=_blueprint(),
            review_report=_review_report(),
            detection_report=_detection_report(),
        )


def test_delivery_file_lookup():
    engine = DeliveryEngineService()
    package = engine.prepare_package(
        final_draft=_draft(),
        requirement_json=_requirement(),
        research_plan=_research_plan(),
        blueprint=_blueprint(),
        review_report=_review_report(),
        detection_report=_detection_report(),
        project_id="proj-lookup",
    )
    first_file = package.files[0]
    found = engine.get_file(first_file.id)
    assert found.filename == first_file.filename
    with pytest.raises(KeyError):
        engine.get_file("missing-file")


def test_project_delivery_pipeline():
    from services.writer_engine import MockSectionWriter
    from services.writer_engine.mock_reviewer import MockSectionReviewer

    pipeline = AssignmentPipelineService()
    writer = WriterEngineService(writer=MockSectionWriter(), reviewer=MockSectionReviewer())
    humanizer = HumanizerEngineService()
    projects = ProjectService(
        pipeline=pipeline,
        research=ResearchEngineService(),
        blueprint=BlueprintEngineService(),
        writer=writer,
        reviewer=ReviewerEngineService(),
        revision=RevisionEngineService(draft_store=writer.drafts),
        humanizer=humanizer,
        ai_detection=AIDetectionEngineService(),
        delivery=DeliveryEngineService(),
    )
    bundle = projects.create_project(files=[{"file_type": "assignment_brief", "original_filename": "brief.pdf"}])
    prepare_project_for_research(projects, bundle.project.id)
    projects.run_research(bundle.project.id)
    projects.run_blueprint(bundle.project.id)
    session = projects.start_writer(bundle.project.id)
    while session.status.value == "active":
        if session.current_section_id and session.section_by_id(session.current_section_id).status == WriterSectionStatus.REVISION:
            session = projects.revise_writer_section(bundle.project.id, session.current_section_id)
        else:
            session = projects.advance_writer(bundle.project.id)
    projects.merge_writer_draft(bundle.project.id)
    projects.run_academic_review(bundle.project.id)

    hz = projects.start_humanizer(bundle.project.id)
    while hz.status.value == "active":
        hz = projects.advance_humanizer(bundle.project.id)
    projects.merge_humanized_draft(bundle.project.id)

    detection = projects.start_ai_detection(bundle.project.id)
    while detection.status.value == "active":
        detection = projects.advance_ai_detection(bundle.project.id)
    projects.finalize_ai_detection(bundle.project.id)

    package = projects.run_delivery(bundle.project.id)
    project = projects.get_project(bundle.project.id).project

    assert package.status == DeliveryStatus.READY
    assert len(package.files) == 1
    assert package.files[0].file_type in {"final_assignment_docx", "final_assignment_pdf"}
    assert project.status == ProjectStatus.COMPLETED
    assert projects.get_delivery_package(bundle.project.id).id == package.id
