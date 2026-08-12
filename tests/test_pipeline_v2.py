"""Formatter V2 pipeline + feature-flagged Flask route."""

from __future__ import annotations

import io
from pathlib import Path
from unittest.mock import patch

from docx import Document

from formatter_v2.pipeline import format_document_v2, select_extractor
from formatter_v2.spec import (
    AppendixConfig,
    Margins,
    ParagraphRole,
    StyleName,
    UserOverrides,
)
from formatter_v2.structure.from_heuristics import HeuristicsExtractor
from formatter_v2.structure.from_word_styles import (
    WordStylesExtractor,
    document_has_structural_styles,
)


def _styled_docx_bytes() -> bytes:
    doc = Document()
    doc.add_paragraph("Climate Adaptation in Coastal Cities", style="Title")
    doc.add_paragraph("Introduction", style="Heading 1")
    doc.add_paragraph("Sea-level rise reshapes municipal budgets.")
    doc.add_paragraph("References")
    doc.add_paragraph("Smith, J. (2020). Coastal governance.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_docx_with_word_styles_uses_style_extractor_not_heuristics() -> None:
    raw = _styled_docx_bytes()
    doc = Document(io.BytesIO(raw))
    assert document_has_structural_styles(doc)
    extractor, name, _ = select_extractor(doc)
    assert name == "word_styles"
    assert isinstance(extractor, WordStylesExtractor)
    result = format_document_v2(doc, UserOverrides(), StyleName.APA7)
    assert result.extractor_name == "word_styles"


def test_plain_text_falls_back_to_heuristics() -> None:
    lines = [
        "Climate Adaptation in Coastal Cities Are Changing Fast",
        "Municipal budgets are under pressure from compound flood risk.",
        "References",
        "Smith, J. (2020). Coastal governance.",
    ]
    extractor, name, document = select_extractor(lines)
    assert name == "heuristics"
    assert document is None
    assert isinstance(extractor, HeuristicsExtractor)
    result = format_document_v2(lines, UserOverrides(), StyleName.HARVARD)
    assert result.extractor_name == "heuristics"


def test_references_latch_from_word_styles() -> None:
    doc = Document()
    doc.add_paragraph("Introduction", style="Heading 1")
    doc.add_paragraph("Body text here.")
    doc.add_paragraph("Works Cited", style="Heading 1")
    doc.add_paragraph("Doe, J. Example Book.")
    doc.add_paragraph("Roe, A. Another Book.")
    model = WordStylesExtractor().extract(doc)
    assert model.references
    assert model.references[0].role == ParagraphRole.REFERENCES_HEADING
    assert model.references[0].text == "Works Cited"
    assert all(b.role == ParagraphRole.REFERENCES_ENTRY for b in model.references[1:])
    assert not any(b.role == ParagraphRole.REFERENCES_ENTRY for b in model.body)


def test_references_latch_from_heuristics() -> None:
    lines = [
        "Introduction",
        "Body paragraph about methods.",
        "Bibliography",
        "Alpha, A. First source.",
        "Beta, B. Second source.",
    ]
    model = HeuristicsExtractor().extract(lines)
    assert model.references[0].role == ParagraphRole.REFERENCES_HEADING
    assert model.references[0].text == "Bibliography"
    assert [b.text for b in model.references[1:]] == [
        "Alpha, A. First source.",
        "Beta, B. Second source.",
    ]


def test_unrecognised_paragraph_defaults_to_body() -> None:
    lines = [
        "This ordinary paragraph has no special markers and should stay body text.",
    ]
    model = HeuristicsExtractor().extract(lines)
    assert len(model.body) == 1
    assert model.body[0].role == ParagraphRole.BODY


def test_pipeline_returns_notices_from_resolver() -> None:
    overrides = UserOverrides(
        margins=Margins(top_in=1.5, bottom_in=1.5, left_in=1.5, right_in=1.5)
    )
    result = format_document_v2(
        ["Introduction", "Body text about sensors."],
        overrides,
        StyleName.IEEE,
    )
    assert result.notices
    assert any(n.field == "margins" for n in result.notices)


def test_route_returns_404_when_flag_disabled() -> None:
    with patch.dict("os.environ", {"FORMATTER_V2_ENABLED": "0"}, clear=False):
        from app import app

        client = app.test_client()
        response = client.post(
            "/api/format-v2",
            data={"pasted_text": "Hello world", "format_style": "harvard"},
        )
        assert response.status_code == 404


def test_route_returns_json_and_download_when_flag_enabled() -> None:
    import tempfile

    from formatter_v2.document_store import reset_document_store

    with patch.dict("os.environ", {"FORMATTER_V2_ENABLED": "1"}, clear=False):
        from app import app

        app.config["TESTING"] = True
        reset_document_store(root=Path(tempfile.mkdtemp()))
        client = app.test_client()
        response = client.post(
            "/api/format-v2",
            data={
                "pasted_text": "Introduction\n\nBody paragraph about climate risk.\n",
                "format_style": "apa7",
            },
        )
        assert response.status_code == 200
        assert response.is_json
        data = response.get_json()
        assert data["document_id"]
        doc = client.get(f"/api/format-v2/download/{data['document_id']}")
        assert doc.status_code == 200
        assert (
            doc.mimetype
            == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        assert doc.data[:2] == b"PK"


def test_pipeline_output_opens_as_valid_docx() -> None:
    result = format_document_v2(
        _styled_docx_bytes(),
        UserOverrides(),
        StyleName.HARVARD,
    )
    reloaded = Document(io.BytesIO(result.docx_bytes))
    texts = [p.text for p in reloaded.paragraphs if p.text.strip()]
    assert texts
    assert any("Climate Adaptation" in t or "Introduction" in t for t in texts)


def test_pipeline_output_has_appendix_after_references() -> None:
    fixture = Path(__file__).resolve().parent / "fixtures" / "test_essay_styled.docx"
    fixture.parent.mkdir(parents=True, exist_ok=True)

    # Keep this fixture representative of the real pipeline path:
    # Word-styled headings + references + trailing appendix block.
    doc = Document()
    doc.add_paragraph("Introduction", style="Heading 1")
    doc.add_paragraph("Body text before references.")
    doc.add_paragraph("References", style="Heading 1")
    doc.add_paragraph("Smith, J. (2020). Coastal governance.")
    doc.add_paragraph("Doe, A. (2019). Flood maps.")
    doc.add_paragraph("Appendix A")
    doc.add_paragraph("Survey instrument and coding guide.")
    doc.save(str(fixture))

    result = format_document_v2(
        str(fixture),
        UserOverrides(appendices=AppendixConfig(enabled=True)),
        StyleName.HARVARD,
    )
    out = Document(io.BytesIO(result.docx_bytes))

    refs_idx = next(
        i
        for i, p in enumerate(out.paragraphs)
        if (p.style.name or "").strip() == "DM References Heading"
    )
    appendix_idx = next(
        i
        for i, p in enumerate(out.paragraphs)
        if (p.style.name or "").strip() == "DM Appendix Heading"
    )
    assert appendix_idx > refs_idx
