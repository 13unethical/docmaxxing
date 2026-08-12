"""XML-level tests for Formatter V2 full-document builder."""

from __future__ import annotations

import io
from datetime import date

import pytest
from docx import Document
from docx.oxml.ns import qn

from formatter_v2.fixtures.sample_full_document import (
    captioned_figure_blocks,
    sample_full_document,
)
from formatter_v2.profiles import load_profile
from formatter_v2.render.builder import build_document
from formatter_v2.render.document import Block
from formatter_v2.render.model import DocumentModel
from formatter_v2.render.styles import style_name_for_role
from formatter_v2.resolve import resolve_format_spec
from formatter_v2.spec import (
    AbbreviationList,
    AppendixConfig,
    CoverPage,
    PageNumbering,
    PageNumberPosition,
    ParagraphRole,
    ReferencesConfig,
    StyleName,
    TableOfContents,
    UserOverrides,
)


def _save_reload(doc) -> Document:
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return Document(buf)


def _spec(
    style: StyleName,
    *,
    overrides: UserOverrides | None = None,
) -> FormatSpec:
    profile = load_profile(style)
    return resolve_format_spec(profile, overrides or UserOverrides()).spec


def _full_overrides(**kwargs) -> UserOverrides:
    base = dict(
        cover_page=CoverPage(enabled=True, title="Climate Adaptation in Coastal Cities"),
        table_of_contents=TableOfContents(
            enabled=True, max_depth=3, heading_text="Table of Contents"
        ),
        abbreviations=AbbreviationList(
            enabled=True,
            heading_text="List of Abbreviations",
            entries={
                "IPCC": "Intergovernmental Panel on Climate Change",
                "SLR": "Sea-Level Rise",
            },
        ),
        appendices=AppendixConfig(
            enabled=True, lettered=True, page_break_before_each=True
        ),
    )
    base.update(kwargs)
    return UserOverrides(**base)


def _para_styles(doc: Document) -> list[str]:
    return [p.style.name if p.style else "" for p in doc.paragraphs]


def _para_texts(doc: Document) -> list[str]:
    return [p.text for p in doc.paragraphs]


def test_document_order_is_cover_front_toc_body_references_appendices() -> None:
    spec = _spec(StyleName.HARVARD, overrides=_full_overrides())
    model = sample_full_document(spec)
    doc = _save_reload(build_document(model, spec))
    styles = _para_styles(doc)
    texts = _para_texts(doc)

    def first(name: str) -> int:
        for i, s in enumerate(styles):
            if s == name:
                return i
        raise AssertionError(f"missing style {name}")

    cover_i = first(style_name_for_role(ParagraphRole.COVER_TITLE))
    abstract_i = first(style_name_for_role(ParagraphRole.ABSTRACT_HEADING))
    toc_i = first(style_name_for_role(ParagraphRole.TOC_HEADING))
    abbrev_heading_i = next(i for i, t in enumerate(texts) if t == "List of Abbreviations")
    heading1 = style_name_for_role(ParagraphRole.HEADING_1)
    body_i = next(
        i
        for i, (t, s) in enumerate(zip(texts, styles))
        if t == "Introduction" and s == heading1
    )
    refs_i = first(style_name_for_role(ParagraphRole.REFERENCES_HEADING))
    app_i = next(i for i, t in enumerate(texts) if t.startswith("Appendix A"))

    assert cover_i < abstract_i < toc_i < abbrev_heading_i < body_i < refs_i < app_i


def test_disabled_sections_are_omitted_entirely() -> None:
    profile = load_profile(StyleName.HARVARD)
    overrides = UserOverrides(
        cover_page=CoverPage(enabled=False, title="Ignored"),
        table_of_contents=TableOfContents(enabled=False),
        abbreviations=AbbreviationList(enabled=False, entries={"X": "Y"}),
        appendices=AppendixConfig(enabled=False),
        references=ReferencesConfig(enabled=False, heading_text="References"),
    )
    spec = resolve_format_spec(profile, overrides).spec
    model = sample_full_document(spec)
    doc = _save_reload(build_document(model, spec))
    texts = "\n".join(_para_texts(doc))
    styles = _para_styles(doc)
    assert style_name_for_role(ParagraphRole.COVER_TITLE) not in styles
    assert style_name_for_role(ParagraphRole.TOC_HEADING) not in styles
    assert "List of Abbreviations" not in texts
    assert style_name_for_role(ParagraphRole.REFERENCES_HEADING) not in styles
    assert "\nAppendix\n" in f"\n{texts}\n"
    assert "Introduction" in texts


def test_references_heading_starts_on_new_page() -> None:
    spec = _spec(StyleName.APA7, overrides=_full_overrides())
    assert spec.roles[ParagraphRole.REFERENCES_HEADING].page_break_before is True
    model = sample_full_document(spec)
    doc = _save_reload(build_document(model, spec))
    style = doc.styles[style_name_for_role(ParagraphRole.REFERENCES_HEADING)]
    assert style.paragraph_format.page_break_before is True


def test_each_appendix_starts_on_new_page() -> None:
    spec = _spec(StyleName.HARVARD, overrides=_full_overrides())
    model = sample_full_document(spec)
    doc = _save_reload(build_document(model, spec))
    xml = doc.element.body.xml
    assert "w:type=\"page\"" in xml or "w:type='page'" in xml
    texts = _para_texts(doc)
    assert texts.index("Appendix A") < texts.index("Appendix B")


def test_appendices_are_lettered_a_b_c() -> None:
    spec = _spec(
        StyleName.HARVARD,
        overrides=_full_overrides(
            appendices=AppendixConfig(
                enabled=True, lettered=True, page_break_before_each=True
            )
        ),
    )
    model = sample_full_document(spec)
    model.appendices = [
        Block(ParagraphRole.APPENDIX_HEADING, "x"),
        Block(ParagraphRole.BODY, "one"),
        Block(ParagraphRole.APPENDIX_HEADING, "y"),
        Block(ParagraphRole.BODY, "two"),
        Block(ParagraphRole.APPENDIX_HEADING, "z"),
        Block(ParagraphRole.BODY, "three"),
    ]
    doc = _save_reload(build_document(model, spec))
    texts = _para_texts(doc)
    assert "Appendix A" in texts
    assert "Appendix B" in texts
    assert "Appendix C" in texts


def test_appendix_renders_after_references_with_config_disabled() -> None:
    spec = _spec(
        StyleName.HARVARD,
        overrides=UserOverrides(
            appendices=AppendixConfig(enabled=False),
            references=ReferencesConfig(enabled=True, heading_text="References"),
        ),
    )
    model = DocumentModel(
        body=[Block(ParagraphRole.HEADING_1, "Introduction"), Block(ParagraphRole.BODY, "Body text.")],
        references=[
            Block(ParagraphRole.REFERENCES_HEADING, "References"),
            Block(ParagraphRole.REFERENCES_ENTRY, "Smith, J. (2020). Coastal governance."),
        ],
        appendices=[
            Block(ParagraphRole.APPENDIX_HEADING, "Appendix A"),
            Block(ParagraphRole.BODY, "Supplementary material."),
        ],
    )
    doc = _save_reload(build_document(model, spec))
    texts = _para_texts(doc)
    refs_i = texts.index("References")
    app_i = texts.index("Appendix A")
    assert app_i > refs_i


def test_appendix_lettering_applies_only_when_config_enabled() -> None:
    model = DocumentModel(
        appendices=[
            Block(ParagraphRole.APPENDIX_HEADING, "Appendix C"),
            Block(ParagraphRole.BODY, "Appendix C body."),
        ]
    )
    enabled_spec = _spec(
        StyleName.HARVARD,
        overrides=UserOverrides(
            appendices=AppendixConfig(enabled=True, lettered=True, page_break_before_each=False)
        ),
    )
    disabled_spec = _spec(
        StyleName.HARVARD,
        overrides=UserOverrides(
            appendices=AppendixConfig(enabled=False, lettered=True, page_break_before_each=True)
        ),
    )
    enabled_doc = _save_reload(build_document(model, enabled_spec))
    disabled_doc = _save_reload(build_document(model, disabled_spec))
    enabled_texts = _para_texts(enabled_doc)
    disabled_texts = _para_texts(disabled_doc)
    assert "Appendix A" in enabled_texts
    assert "Appendix C" in disabled_texts


def test_update_fields_flag_present_when_toc_enabled() -> None:
    spec = _spec(
        StyleName.APA7,
        overrides=_full_overrides(
            table_of_contents=TableOfContents(enabled=True, max_depth=2)
        ),
    )
    doc = _save_reload(build_document(sample_full_document(spec), spec))
    settings_xml = doc.settings.element.xml
    assert "updateFields" in settings_xml
    assert 'w:val="true"' in settings_xml or "w:val='true'" in settings_xml


def test_toc_field_present_with_correct_depth() -> None:
    spec = _spec(
        StyleName.APA7,
        overrides=_full_overrides(
            table_of_contents=TableOfContents(
                enabled=True, max_depth=2, field_based=True
            )
        ),
    )
    doc = _save_reload(build_document(sample_full_document(spec), spec))
    body_xml = doc.element.body.xml
    assert "TOC" in body_xml
    assert "1-2" in body_xml


def test_front_matter_uses_roman_numerals_when_restart_enabled() -> None:
    numbering = PageNumbering(
        position=PageNumberPosition.TOP_RIGHT,
        restart_after_front_matter=True,
        skip_first_page=False,
    )
    spec = _spec(StyleName.HARVARD, overrides=_full_overrides(page_numbering=numbering))
    assert spec.page_numbering.restart_after_front_matter is True
    doc = _save_reload(build_document(sample_full_document(spec), spec))
    assert len(doc.sections) >= 2
    pg = doc.sections[0]._sectPr.find(qn("w:pgNumType"))
    assert pg is not None
    assert pg.get(qn("w:fmt")) == "lowerRoman"


def test_body_restarts_at_arabic_one_when_restart_enabled() -> None:
    numbering = PageNumbering(
        position=PageNumberPosition.TOP_RIGHT,
        restart_after_front_matter=True,
        skip_first_page=False,
    )
    spec = _spec(StyleName.HARVARD, overrides=_full_overrides(page_numbering=numbering))
    doc = _save_reload(build_document(sample_full_document(spec), spec))
    assert len(doc.sections) >= 2
    pg = doc.sections[1]._sectPr.find(qn("w:pgNumType"))
    assert pg is not None
    assert pg.get(qn("w:fmt")) == "decimal"
    assert pg.get(qn("w:start")) == "1"


def test_single_section_when_restart_disabled() -> None:
    numbering = PageNumbering(
        position=PageNumberPosition.TOP_RIGHT,
        restart_after_front_matter=False,
    )
    spec = _spec(StyleName.HARVARD, overrides=_full_overrides(page_numbering=numbering))
    doc = _save_reload(build_document(sample_full_document(spec), spec))
    assert len(doc.sections) == 1


def test_table_and_figure_counters_are_independent() -> None:
    spec = _spec(StyleName.HARVARD, overrides=_full_overrides())
    model = DocumentModel(
        body=[
            Block(ParagraphRole.TABLE_CAPTION, "First table"),
            Block(ParagraphRole.TABLE_CAPTION, "Second table"),
            Block(ParagraphRole.FIGURE_CAPTION, "First figure"),
            Block(ParagraphRole.FIGURE_CAPTION, "Second figure"),
        ]
    )
    doc = _save_reload(build_document(model, spec))
    texts = _para_texts(doc)
    assert any(t.startswith("Table 1") for t in texts)
    assert any(t.startswith("Table 2") for t in texts)
    assert any(t.startswith("Figure 1") or t.startswith("Fig. 1") for t in texts)
    assert any(t.startswith("Figure 2") or t.startswith("Fig. 2") for t in texts)


def test_apa_figure_caption_paragraph_precedes_figure() -> None:
    spec = _spec(StyleName.APA7, overrides=_full_overrides())
    assert spec.captions.figure_position == "above"
    blocks = captioned_figure_blocks(
        "Flood exposure under RCP4.5",
        "[Figure: map placeholder]",
        spec,
    )
    model = DocumentModel(body=blocks)
    doc = _save_reload(build_document(model, spec))
    styles = _para_styles(doc)
    texts = _para_texts(doc)
    cap_style = style_name_for_role(ParagraphRole.FIGURE_CAPTION)
    cap_i = styles.index(cap_style)
    fig_i = next(i for i, t in enumerate(texts) if "map placeholder" in t)
    assert cap_i < fig_i


def test_mla_first_page_has_four_line_header_block() -> None:
    spec = _spec(StyleName.MLA9, overrides=_full_overrides())
    model = sample_full_document(spec)
    assert model.cover is not None
    model.cover.student_name = "Alex Morgan"
    model.cover.lecturer = "Dr. Sam Rivera"
    model.cover.course = "Environmental Policy 301"
    model.cover.submission_date = date(2026, 5, 15)
    doc = _save_reload(build_document(model, spec))
    texts = _para_texts(doc)
    assert texts[0] == "Alex Morgan"
    assert texts[1] == "Dr. Sam Rivera"
    assert texts[2] == "Environmental Policy 301"
    assert "2026" in texts[3]


def test_mla_has_no_cover_page_even_if_requested() -> None:
    overrides = _full_overrides(
        cover_page=CoverPage(enabled=True, title="Should Appear As Doc Title"),
    )
    spec = _spec(StyleName.MLA9, overrides=overrides)
    spec.cover_page.enabled = True
    model = sample_full_document(spec)
    model.cover = CoverPage(
        enabled=True,
        title="Should Appear As Doc Title",
        student_name="Alex Morgan",
        lecturer="Dr. Sam Rivera",
        course="ENV 301",
        submission_date=date(2026, 5, 15),
    )
    doc = _save_reload(build_document(model, spec))
    styles = _para_styles(doc)
    assert style_name_for_role(ParagraphRole.COVER_TITLE) not in styles
    assert style_name_for_role(ParagraphRole.DOC_TITLE) in styles
    assert any("Should Appear" in t and "Doc Title" in t for t in _para_texts(doc))


def test_apa_cover_page_is_numbered_page_one() -> None:
    spec = _spec(StyleName.APA7, overrides=_full_overrides())
    assert spec.page_numbering.skip_first_page is False
    model = sample_full_document(spec)
    doc = _save_reload(build_document(model, spec))
    section = doc.sections[0]
    assert section.different_first_page_header_footer is False
    header_xml = section.header._element.xml
    assert "PAGE" in header_xml or "w:fldChar" in header_xml
    assert style_name_for_role(ParagraphRole.COVER_TITLE) in _para_styles(doc)


def test_builder_and_renderer_use_same_paragraph_function() -> None:
    from pathlib import Path

    import formatter_v2.render.builder as builder_mod
    import formatter_v2.render.document as document_mod

    assert hasattr(document_mod, "add_paragraph")
    assert builder_mod.add_paragraph is document_mod.add_paragraph
    source = Path(builder_mod.__file__).read_text(encoding="utf-8")
    assert "document.add_paragraph" not in source


@pytest.mark.parametrize(
    ("style", "expected"),
    [
        (StyleName.APA7, "May 15, 2026"),
        (StyleName.IEEE, "May 15, 2026"),
        (StyleName.MLA9, "15 May 2026"),
        (StyleName.HARVARD, "15 May 2026"),
        (StyleName.CHICAGO17, "15 May 2026"),
    ],
)
def test_cover_date_format_per_style(style: StyleName, expected: str) -> None:
    from formatter_v2.render.dates import format_cover_date

    profile = load_profile(style)
    assert format_cover_date(date(2026, 5, 15), profile.date_format) == expected
    spec = _spec(style)
    assert spec.date_format == profile.date_format


def test_mla_title_appears_after_header_block_before_body() -> None:
    spec = _spec(StyleName.MLA9, overrides=_full_overrides())
    model = sample_full_document(spec)
    assert model.cover is not None
    model.cover.student_name = "Alex Morgan"
    model.cover.lecturer = "Dr. Sam Rivera"
    model.cover.course = "Environmental Policy 301"
    model.cover.submission_date = date(2026, 5, 15)
    model.cover.title = "Climate Adaptation in Coastal Cities"
    model.front_matter = []
    doc = _save_reload(build_document(model, spec))
    texts = _para_texts(doc)
    styles = _para_styles(doc)
    assert texts[0] == "Alex Morgan"
    assert texts[1] == "Dr. Sam Rivera"
    assert texts[2] == "Environmental Policy 301"
    assert texts[3] == "15 May 2026"
    title_style = style_name_for_role(ParagraphRole.DOC_TITLE)
    assert styles[4] == title_style
    assert "Climate Adaptation" in texts[4]
    body_i = next(i for i, t in enumerate(texts) if t == "Introduction")
    assert 4 < body_i


def test_mla_header_block_has_no_first_line_indent() -> None:
    spec = _spec(StyleName.MLA9, overrides=_full_overrides())
    model = sample_full_document(spec)
    assert model.cover is not None
    model.cover.student_name = "Alex Morgan"
    model.cover.lecturer = "Dr. Sam Rivera"
    model.cover.course = "ENV 301"
    model.cover.submission_date = date(2026, 5, 15)
    doc = _save_reload(build_document(model, spec))
    for para in doc.paragraphs[:4]:
        indent = para.paragraph_format.first_line_indent
        assert indent is None or abs(indent.inches) < 0.001


def test_toc_field_placeholder_is_instructional_not_numeric() -> None:
    spec = _spec(
        StyleName.APA7,
        overrides=_full_overrides(
            table_of_contents=TableOfContents(
                enabled=True, max_depth=2, field_based=True
            )
        ),
    )
    doc = _save_reload(build_document(sample_full_document(spec), spec))
    body_xml = doc.element.body.xml
    assert "Right-click here and choose Update Field" in body_xml
    # Numeric "1" alone as TOC placeholder would be wrong; instructional text must win.
    assert "table of contents" in body_xml.lower()


def test_static_toc_lists_body_headings_in_order() -> None:
    spec = _spec(
        StyleName.HARVARD,
        overrides=_full_overrides(
            table_of_contents=TableOfContents(
                enabled=True, max_depth=3, field_based=False, page_break_after=False
            ),
            abbreviations=AbbreviationList(enabled=False),
        ),
    )
    model = DocumentModel(
        body=[
            Block(ParagraphRole.HEADING_1, "Introduction"),
            Block(ParagraphRole.BODY, "Body A."),
            Block(ParagraphRole.HEADING_2, "Methods"),
            Block(ParagraphRole.BODY, "Body B."),
            Block(ParagraphRole.HEADING_1, "Conclusion"),
        ]
    )
    doc = _save_reload(build_document(model, spec))
    toc_style = style_name_for_role(ParagraphRole.TOC_ENTRY)
    entries = [p.text for p in doc.paragraphs if p.style and p.style.name == toc_style]
    assert entries == ["Introduction", "Methods", "Conclusion"]


def test_static_toc_respects_max_depth() -> None:
    spec = _spec(
        StyleName.HARVARD,
        overrides=_full_overrides(
            table_of_contents=TableOfContents(
                enabled=True, max_depth=1, field_based=False, page_break_after=False
            ),
            abbreviations=AbbreviationList(enabled=False),
        ),
    )
    model = DocumentModel(
        body=[
            Block(ParagraphRole.HEADING_1, "Top"),
            Block(ParagraphRole.HEADING_2, "Nested"),
            Block(ParagraphRole.HEADING_3, "Deeper"),
        ]
    )
    doc = _save_reload(build_document(model, spec))
    toc_style = style_name_for_role(ParagraphRole.TOC_ENTRY)
    entries = [p.text for p in doc.paragraphs if p.style and p.style.name == toc_style]
    assert entries == ["Top"]


def test_static_toc_has_no_page_numbers() -> None:
    spec = _spec(
        StyleName.HARVARD,
        overrides=_full_overrides(
            table_of_contents=TableOfContents(
                enabled=True, max_depth=3, field_based=False, page_break_after=False
            ),
            abbreviations=AbbreviationList(enabled=False),
        ),
    )
    model = DocumentModel(
        body=[
            Block(ParagraphRole.HEADING_1, "Introduction"),
            Block(ParagraphRole.BODY, "Body."),
        ]
    )
    doc = _save_reload(build_document(model, spec))
    toc_style = style_name_for_role(ParagraphRole.TOC_ENTRY)
    for para in doc.paragraphs:
        if para.style and para.style.name == toc_style:
            assert para.text == "Introduction"
            assert not any(ch.isdigit() for ch in para.text if ch != "")
    body_xml = doc.element.body.xml
    assert "TOC \\o" not in body_xml and 'TOC \\o' not in body_xml


def test_field_toc_still_available_when_explicitly_enabled() -> None:
    spec = _spec(
        StyleName.APA7,
        overrides=_full_overrides(
            table_of_contents=TableOfContents(
                enabled=True, max_depth=3, field_based=True
            )
        ),
    )
    assert spec.table_of_contents.field_based is True
    doc = _save_reload(build_document(sample_full_document(spec), spec))
    assert "TOC" in doc.element.body.xml
    assert "Right-click here and choose Update Field" in doc.element.body.xml
