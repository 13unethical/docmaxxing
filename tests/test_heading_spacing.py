"""Academic heading spacing rules."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from docx.enum.text import WD_LINE_SPACING

from tests.conftest import run_format_pipeline
from formatter import FormatJob, format_document_full
from formatter.heading_spacing import resolve_paragraph_spacing
from styles.profile import (
    ACADEMIC_HEADING2_SPACE_AFTER_PT,
    ACADEMIC_HEADING2_SPACE_BEFORE_PT,
    ACADEMIC_TITLE_SPACE_AFTER_PT,
)

IDEAL_ESSAY = Path("/Users/nazirov/Desktop/formatted_document_20260702-192427.docx")
BROKEN_ESSAY = Path("/Users/nazirov/Desktop/formatted_document_20260702-202350.docx")


def _pt(value) -> float:
    return round(value.pt, 1) if value is not None else 0.0


@pytest.mark.skipif(not IDEAL_ESSAY.is_file() or not BROKEN_ESSAY.is_file(), reason="essay fixtures on Desktop")
def test_essay_spacing_matches_ideal_docx():
    """Re-format broken essay (correct structure) — spacing must match ideal."""
    from formatter.markdown_cleanup import strip_markdown_text

    def _norm_text(text: str) -> str:
        cleaned = strip_markdown_text(text or "")
        return (cleaned or "").strip()

    ideal = Document(str(IDEAL_ESSAY))
    source = Document(str(BROKEN_ESSAY))
    job = FormatJob(
        font_family="Times New Roman",
        font_size_pt=12,
        line_spacing=1.5,
        alignment="justify",
        first_line_indent=False,
        space_before_pt=0,
        space_after_pt=12,
        margin_preset="normal",
        page_number_position="none",
        auto_headings=True,
        heading_all_caps=False,
        auto_justify_refs=True,
    )
    run_format_pipeline(source, job, document_type="essay")

    assert len(source.paragraphs) == len(ideal.paragraphs)
    assert sum(1 for p in source.paragraphs if p.style.name == "Heading 1") == 1
    assert sum(1 for p in source.paragraphs if p.style.name == "Heading 2") == 5

    for idx, (gen_p, ideal_p) in enumerate(zip(source.paragraphs, ideal.paragraphs)):
        assert _norm_text(gen_p.text) == _norm_text(ideal_p.text)
        assert gen_p.style.name == ideal_p.style.name
        gpf = gen_p.paragraph_format
        if gen_p.style.name == "Heading 1":
            assert _pt(gpf.space_before) == 0
            assert _pt(gpf.space_after) == ACADEMIC_TITLE_SPACE_AFTER_PT
        elif gen_p.style.name == "Heading 2":
            assert _pt(gpf.space_before) == ACADEMIC_HEADING2_SPACE_BEFORE_PT
            assert _pt(gpf.space_after) == ACADEMIC_HEADING2_SPACE_AFTER_PT
        elif gen_p.style.name == "Normal":
            next_is_heading = idx + 1 < len(source.paragraphs) and (
                source.paragraphs[idx + 1].style.name or ""
            ).startswith("Heading")
            expected_after = 0 if next_is_heading else 12
            assert _pt(gpf.space_after) == expected_after
        if (gen_p.style.name or "").startswith("Heading"):
            assert gpf.line_spacing_rule in (WD_LINE_SPACING.SINGLE, None)
            assert gpf.keep_with_next is True


def test_heading_single_line_and_keep_with_next():
    doc = Document()
    doc.add_paragraph("Body before heading.")
    p = doc.add_paragraph("Introduction")
    p.style = "Heading 2"

    job = FormatJob(
        font_family="Times New Roman",
        font_size_pt=12,
        line_spacing=1.5,
        alignment="left",
        first_line_indent=False,
        space_before_pt=0,
        space_after_pt=12,
        margin_preset="normal",
        page_number_position="none",
        auto_headings=True,
        heading_all_caps=False,
        auto_justify_refs=False,
    )
    format_document_full(doc, job, None)
    heading = doc.paragraphs[1]
    pf = heading.paragraph_format
    assert pf.line_spacing_rule in (WD_LINE_SPACING.SINGLE, None)
    assert pf.keep_with_next is True


def test_resolve_paragraph_spacing_body_gets_space_after():
    sb, sa = resolve_paragraph_spacing(
        level=0,
        prev_level=0,
        next_level=0,
        prev_has_text=True,
        font_size_pt=12,
        line_spacing=1.5,
        body_space_before_pt=0,
        body_space_after_pt=0,
        format_style="harvard",
    )
    assert sb == 0
    assert sa == 12


def test_resolve_paragraph_spacing_body_before_heading_has_no_trailing_gap():
    sb, sa = resolve_paragraph_spacing(
        level=0,
        prev_level=0,
        next_level=2,
        prev_has_text=True,
        font_size_pt=12,
        line_spacing=1.5,
        body_space_before_pt=0,
        body_space_after_pt=0,
        format_style="harvard",
    )
    assert sa == 0


def test_resolve_paragraph_spacing_heading_after_body_gets_space_before():
    sb, sa = resolve_paragraph_spacing(
        level=2,
        prev_level=0,
        next_level=0,
        prev_has_text=True,
        font_size_pt=12,
        line_spacing=1.5,
        body_space_before_pt=0,
        body_space_after_pt=0,
        format_style="harvard",
    )
    assert sb == ACADEMIC_HEADING2_SPACE_BEFORE_PT
    assert sa == ACADEMIC_HEADING2_SPACE_AFTER_PT


def test_harvard_formatting_restores_body_and_heading_spacing():
    doc = Document()
    doc.add_paragraph("Some body text before the section with enough words.")
    doc.add_paragraph("Introduction")
    doc.add_paragraph("First paragraph of introduction section here.")
    doc.add_paragraph("Second body paragraph continues the discussion.")

    job = FormatJob(
        font_family="Times New Roman",
        font_size_pt=12,
        line_spacing=1.5,
        alignment="justify",
        first_line_indent=False,
        space_before_pt=0,
        space_after_pt=12,
        margin_preset="normal",
        page_number_position="none",
        auto_headings=True,
        heading_all_caps=False,
        auto_justify_refs=True,
    )
    format_document_full(doc, job, None)

    body = doc.paragraphs[0]
    heading = doc.paragraphs[1]
    body_after_heading = doc.paragraphs[2]
    body2 = doc.paragraphs[3]

    assert body.paragraph_format.space_after.pt == 0
    assert heading.paragraph_format.space_before.pt == ACADEMIC_HEADING2_SPACE_BEFORE_PT
    assert heading.paragraph_format.space_after.pt == ACADEMIC_HEADING2_SPACE_AFTER_PT
    assert body_after_heading.paragraph_format.space_after.pt == 12
    assert body2.paragraph_format.space_after.pt == 12
