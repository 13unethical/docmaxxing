"""Numbered heading vs list discrimination + Word-styles content latch."""

from __future__ import annotations

from docx import Document

from formatter_v2.pipeline import format_document_v2
from formatter_v2.spec import ParagraphRole, StyleName, UserOverrides
from formatter_v2.structure.from_heuristics import HeuristicsExtractor
from formatter_v2.structure.from_word_styles import WordStylesExtractor
from formatter_v2.structure.numbered import numbered_section_notices

# Real-case shape from 2000 w.docx: eight city sections separated by body,
# plus four consecutive labelled list items.
_CITY_SECTIONS_FIXTURE = [
    "Introduction",
    "This essay compares adaptation strategies across eight coastal cities.",
    "1. Rotterdam, Netherlands — Floodgates and adaptive infrastructure",
    "Rotterdam invested early in barriers and flexible public space design.",
    "2. New York, United States — Zoning reform after major storms",
    "Post-storm rebuilding revised coastal zoning and insurance rules.",
    "3. Lagos, Nigeria — Informal settlements and drainage upgrades",
    "Municipal works focused on channels serving low-lying neighbourhoods.",
    "4. Jakarta, Indonesia — Land subsidence and coastal defence",
    "Subsidence compounds sea-level rise and forces large relocation plans.",
    "5. Miami, United States — Nuisance flooding and property markets",
    "Repeated sunny-day floods reshaped investor expectations downtown.",
    "6. Shanghai, China — Delta megacity protection systems",
    "Large engineered defences sit alongside neighbourhood-scale measures.",
    "7. London, United Kingdom — Thames Barrier and estuary planning",
    "Estuary strategies couple a barrier with upstream storage options.",
    "8. Tokyo, Japan — Underground cisterns and dense urban runoff",
    "Underground storage reduces peak flows during extreme rainfall events.",
    "The following operational checklist summarises fieldwork priorities:",
    "1. Scope: confirm municipal contacts and data access windows",
    "2. Visit: schedule site walks with local engineers and planners",
    "3. Record: photograph assets and note maintenance regimes on site",
    "4. Report: draft comparative findings for the methods appendix",
]


def test_isolated_numbered_short_line_is_heading() -> None:
    lines = [
        "Overview of the case.",
        "1. Rotterdam, Netherlands — Floodgates",
        "Supporting paragraph after the section title continues here.",
    ]
    model = HeuristicsExtractor().extract(lines)
    assert model.body[1].role == ParagraphRole.HEADING_1
    assert model.body[1].text.startswith("1. Rotterdam")


def test_consecutive_numbered_items_are_list() -> None:
    lines = [
        "Checklist:",
        "1. Scope: confirm municipal contacts and data access",
        "2. Visit: schedule site walks with local engineers",
        "3. Record: photograph assets and note maintenance",
        "4. Report: draft comparative findings for appendix",
    ]
    model = HeuristicsExtractor().extract(lines)
    listed = [b for b in model.body if b.role == ParagraphRole.LIST_NUMBER]
    assert len(listed) == 4


def test_numbered_sections_separated_by_body_are_headings() -> None:
    model = HeuristicsExtractor().extract(_CITY_SECTIONS_FIXTURE)
    headings = [
        b
        for b in model.body
        if b.role == ParagraphRole.HEADING_1 and b.text.startswith(tuple(f"{i}." for i in range(1, 9)))
    ]
    # Eight city titles should be headings, not list items.
    city_heads = [
        b
        for b in model.body
        if isinstance(b.text, str)
        and b.text[0].isdigit()
        and "—" in b.text
        and b.role == ParagraphRole.HEADING_1
    ]
    assert len(city_heads) == 8
    assert not any(
        isinstance(b.text, str) and "—" in b.text and b.role == ParagraphRole.LIST_NUMBER
        for b in model.body
    )


def test_numbered_item_with_colon_label_is_list() -> None:
    model = HeuristicsExtractor().extract(_CITY_SECTIONS_FIXTURE)
    checklist = [
        b
        for b in model.body
        if b.role == ParagraphRole.LIST_NUMBER and isinstance(b.text, str) and ": " in b.text
    ]
    assert len(checklist) == 4


def test_dotted_multilevel_numbering_maps_to_heading_depth() -> None:
    lines = [
        "1. Background",
        "Paragraph under chapter one.",
        "1.1 Coastal risk",
        "Details under section 1.1 continue in this paragraph.",
        "1.1.1 Sensor networks",
        "Details under subsection 1.1.1 continue here.",
        "2. Methods",
        "Paragraph under chapter two.",
    ]
    model = HeuristicsExtractor().extract(lines)
    by_text = {b.text: b.role for b in model.body if isinstance(b.text, str)}
    assert by_text["1. Background"] == ParagraphRole.HEADING_1
    assert by_text["1.1 Coastal risk"] == ParagraphRole.HEADING_2
    assert by_text["1.1.1 Sensor networks"] == ParagraphRole.HEADING_3
    assert by_text["2. Methods"] == ParagraphRole.HEADING_1


def test_long_numbered_paragraph_is_body_not_heading() -> None:
    long = (
        "1. This unusually long numbered paragraph keeps going well past one "
        "hundred characters so it must not be promoted to a heading role even "
        "though it starts with a number and has no list neighbours nearby."
    )
    assert len(long) >= 100
    model = HeuristicsExtractor().extract(
        [
            "Introduction",
            long,
            "A following ordinary paragraph.",
        ]
    )
    assert model.body[1].role == ParagraphRole.BODY
    assert model.body[1].text == long


def test_word_styles_document_without_refs_heading_still_latches() -> None:
    doc = Document()
    doc.add_paragraph("Introduction", style="Heading 1")
    doc.add_paragraph("Body text about coastal sensors and budgets.")
    # No References heading — only a trailing bibliography under Normal / List Number.
    for line in (
        "[1] Smith, J. (2020). Coastal governance.",
        "[2] Doe, A. (2019). Flood risk mapping.",
        "[3] Roe, B. (2021). Adaptation finance.",
    ):
        doc.add_paragraph(line, style="List Number")
    model = WordStylesExtractor().extract(doc)
    assert len(model.references) == 3
    assert all(b.role == ParagraphRole.REFERENCES_ENTRY for b in model.references)
    assert not any(b.role == ParagraphRole.LIST_NUMBER for b in model.body)


def test_numbered_heading_right_after_numbered_list_is_heading() -> None:
    lines = [
        "Methods",
        "The study proceeded in three steps:",
        "1. Collect municipal budget data for coastal cities.",
        "2. Estimate flood exposure under RCP scenarios carefully.",
        "3. Compare adaptation instruments across jurisdictions.",
        "4. Findings",
        "Exposure rose in every modelled city under the high-emissions pathway.",
    ]
    model = HeuristicsExtractor().extract(lines)
    by_text = {b.text: b.role for b in model.body if isinstance(b.text, str)}
    assert by_text["1. Collect municipal budget data for coastal cities."] == (
        ParagraphRole.LIST_NUMBER
    )
    assert by_text["3. Compare adaptation instruments across jurisdictions."] == (
        ParagraphRole.LIST_NUMBER
    )
    assert by_text["4. Findings"] == ParagraphRole.HEADING_1


def test_list_item_followed_by_list_item_stays_list() -> None:
    lines = [
        "Checklist:",
        "1. Scope: confirm municipal contacts and data access",
        "2. Visit: schedule site walks with local engineers",
        "3. Record: photograph assets and note maintenance",
    ]
    model = HeuristicsExtractor().extract(lines)
    listed = [b for b in model.body if b.role == ParagraphRole.LIST_NUMBER]
    assert len(listed) == 3
    assert all(b.role == ParagraphRole.LIST_NUMBER for b in listed)


def test_last_list_item_followed_by_body_stays_list() -> None:
    last = (
        "3. Record: photograph assets and note maintenance regimes on site "
        "including serial numbers, access constraints, and follow-up owners"
    )
    assert ":" in last and len(last) >= 60
    lines = [
        "Fieldwork proceeded as follows:",
        "1. Scope: confirm municipal contacts and data access windows",
        "2. Visit: schedule site walks with local engineers and planners",
        last,
        "The following chapter discusses comparative findings in detail.",
    ]
    model = HeuristicsExtractor().extract(lines)
    listed = [b for b in model.body if b.role == ParagraphRole.LIST_NUMBER]
    assert len(listed) == 3
    assert listed[-1].text == last
    assert model.body[-1].role == ParagraphRole.BODY


def test_numbered_section_run_emits_info_notice() -> None:
    model = HeuristicsExtractor().extract(_CITY_SECTIONS_FIXTURE)
    notices = numbered_section_notices(model)
    assert len(notices) == 1
    assert notices[0].severity == "info"
    assert notices[0].field == "structure.numbered_sections"

    result = format_document_v2(
        _CITY_SECTIONS_FIXTURE,
        UserOverrides(),
        StyleName.HARVARD,
    )
    assert any(n.field == "structure.numbered_sections" for n in result.notices)


def test_numbered_heading_right_after_numbered_list_is_heading() -> None:
    lines = [
        "Methods",
        "The study proceeded in three steps:",
        "1. Collect municipal budget data for coastal cities.",
        "2. Estimate flood exposure under RCP scenarios carefully.",
        "3. Compare adaptation instruments across jurisdictions.",
        "4. Findings",
        "Exposure rose in every modelled city along the delta.",
    ]
    model = HeuristicsExtractor().extract(lines)
    by_text = {b.text: b.role for b in model.body if isinstance(b.text, str)}
    assert by_text["1. Collect municipal budget data for coastal cities."] == (
        ParagraphRole.LIST_NUMBER
    )
    assert by_text["3. Compare adaptation instruments across jurisdictions."] == (
        ParagraphRole.LIST_NUMBER
    )
    assert by_text["4. Findings"] == ParagraphRole.HEADING_1
    assert by_text["Exposure rose in every modelled city along the delta."] == (
        ParagraphRole.BODY
    )


def test_list_item_followed_by_list_item_stays_list() -> None:
    lines = [
        "Checklist:",
        "1. Scope: confirm municipal contacts and data access",
        "2. Visit: schedule site walks with local engineers",
        "3. Record: photograph assets and note maintenance",
    ]
    model = HeuristicsExtractor().extract(lines)
    listed = [b for b in model.body if b.role == ParagraphRole.LIST_NUMBER]
    assert len(listed) == 3
    assert not any(b.role == ParagraphRole.HEADING_1 and str(b.text)[0].isdigit() for b in model.body)


def test_last_list_item_followed_by_body_stays_list() -> None:
    last = (
        "3. Record: photograph assets and note maintenance regimes on site, "
        "including serial numbers and access constraints for follow-up visits"
    )
    assert len(last) >= 60 and ":" in last
    lines = [
        "Fieldwork proceeded as follows:",
        "1. Scope: confirm municipal contacts and data access windows",
        "2. Visit: schedule site walks with local engineers and planners",
        last,
        "The comparative draft was then written up for the methods appendix.",
    ]
    model = HeuristicsExtractor().extract(lines)
    listed = [b for b in model.body if b.role == ParagraphRole.LIST_NUMBER]
    assert len(listed) == 3
    assert listed[-1].text == last
    assert model.body[-1].role == ParagraphRole.BODY
