"""Regression cases: heading/body reconstruction must split merged paragraphs."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from formatter import FormatJob, format_document_full
from formatter.document_reconstruction import reconstruct_blocks
from tests.conftest import run_format_pipeline

IDEAL_ESSAY = Path("/Users/nazirov/Desktop/formatted_document_20260702-192427.docx")
BROKEN_ESSAY = Path("/Users/nazirov/Desktop/formatted_document_20260702-202350.docx")

CASES = {
    1: (
        "Introduction This is the introduction with enough body text here.",
        ["Introduction", "This is the introduction with enough body text here."],
    ),
    2: (
        "Body Paragraph 1: Benefits of AI Artificial intelligence improves learning outcomes.",
        ["Body Paragraph 1: Benefits of AI", "Artificial intelligence improves learning outcomes."],
    ),
    3: (
        "Journal Entry 1: Reflection on Week 1 Today I learned about entrepreneurship.",
        ["Journal Entry 1: Reflection on Week 1", "Today I learned about entrepreneurship."],
    ),
    4: (
        "References Smith, J. (2022). Example article. Journal, 1(1), 1-10.",
        ["References", "Smith, J. (2022). Example article. Journal, 1(1), 1-10."],
    ),
}


def _job() -> FormatJob:
    return FormatJob(
        font_family="Times New Roman",
        font_size_pt=12,
        line_spacing=1.5,
        alignment="justify",
        first_line_indent=False,
        space_before_pt=0,
        space_after_pt=12,
        margin_preset="normal",
        page_number_position="none",
        auto_headings=True,
        heading_all_caps=False,
        auto_justify_refs=True,
        requirement_headings=False,
    )


@pytest.mark.parametrize("case_id", [1, 2, 3, 4])
def test_split_embedded_heading_paragraph_cases(case_id: int):
    merged, expected_parts = CASES[case_id]
    blocks = reconstruct_blocks([merged], document_type="essay")
    texts = [b.text for b in blocks]
    assert texts == expected_parts


@pytest.mark.parametrize("case_id", [1, 2, 3, 4])
def test_format_document_splits_merged_cases(case_id: int):
    merged, expected_parts = CASES[case_id]
    doc = Document()
    doc.add_paragraph(merged)
    run_format_pipeline(doc, _job(), document_type="essay")
    texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    assert texts == expected_parts
    assert doc.paragraphs[0].style.name.startswith("Heading")


def test_ai_assignments_do_not_block_reconstruction():
    merged = CASES[1][0]
    doc = Document()
    doc.add_paragraph(merged)
    run_format_pipeline(doc, _job(), document_type="essay")
    texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    assert texts == CASES[1][1]


def test_case_5_humanizer_merged_paragraphs_reconstructed():
    """Each section merged into its own paragraph (typical humanizer output)."""
    doc = Document()
    for case_id in (1, 2, 3, 4):
        doc.add_paragraph(CASES[case_id][0])
    run_format_pipeline(doc, _job(), document_type="essay")
    expected: list[str] = []
    for case_id in (1, 2, 3, 4):
        expected.extend(CASES[case_id][1])
    texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    assert texts == expected
    heading_count = sum(1 for p in doc.paragraphs if (p.style.name or "").startswith("Heading"))
    assert heading_count == 4


@pytest.mark.skipif(not IDEAL_ESSAY.is_file() or not BROKEN_ESSAY.is_file(), reason="essay fixtures")
def test_humanized_essay_reconstruction_matches_ideal_structure():
    from formatter.document_reconstruction import reconstruct_document_before_format
    from formatter.markdown_cleanup import strip_markdown_text

    def _norm_text(text: str) -> str:
        cleaned = strip_markdown_text(text or "")
        return (cleaned or "").strip()

    ideal = Document(str(IDEAL_ESSAY))
    merged = Document(str(BROKEN_ESSAY))
    job = _job()
    recon = reconstruct_document_before_format(merged, document_type="essay", prefer_ai=False)
    format_document_full(merged, job, recon.assignments)

    assert len(merged.paragraphs) == len(ideal.paragraphs)
    for gen_p, ideal_p in zip(merged.paragraphs, ideal.paragraphs):
        assert _norm_text(gen_p.text) == _norm_text(ideal_p.text)
        assert gen_p.style.name == ideal_p.style.name
