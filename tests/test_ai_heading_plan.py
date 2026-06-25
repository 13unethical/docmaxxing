"""Tests for AI heading source-of-truth through format pipeline."""

from __future__ import annotations

import io

from docx import Document

from formatter import FormatJob, format_document_full
from formatter.heading_plan import ParagraphHeadingAssignment, resolve_paragraph_heading_level
from formatter.structure_rebuild import rebuild_document_from_recovery
from services.ai_structure_recovery import ai_result_to_recovery_payload


SAMPLE_AI_RESPONSE = {
    "document_type": "learning_journal",
    "document_type_confidence": 0.88,
    "sections": [
        {
            "title": "Title",
            "heading_text": "Learning Journal",
            "confidence": 0.9,
            "paragraph_indices": [1],
            "insert_heading": False,
            "level": 1,
        },
        {
            "title": "Introduction",
            "heading_text": "Introduction",
            "confidence": 0.85,
            "paragraph_indices": [2],
            "insert_heading": True,
            "level": 2,
        },
        {
            "title": "Journal Entry 3",
            "heading_text": "Journal Entry 3: Entrepreneurial Response to Change",
            "confidence": 0.91,
            "paragraph_indices": [3],
            "insert_heading": True,
            "level": 2,
        },
        {
            "title": "References",
            "heading_text": "References",
            "confidence": 0.9,
            "paragraph_indices": [4],
            "insert_heading": True,
            "level": 2,
        },
    ],
}

ORIGINAL_PARAGRAPHS = [
    "Learning Journal",
    "This journal introduces the module themes.",
    "Journal Entry 3: Entrepreneurial Response to Change\nBody about innovation.",
    "Smith, J. (2021). AI. Journal, 1(1), 1–10.",
]


def test_resolve_priority_ai_wins_over_heuristic():
    assignment = ParagraphHeadingAssignment(text="Intro", level=2, source="ai", confidence=0.9)
    level, source, recovered = resolve_paragraph_heading_level(
        assignment=assignment,
        word_style_level=None,
        heuristic_level=3,
        auto_headings=True,
    )
    assert level == 2
    assert source == "ai"
    assert recovered == 2


def test_resolve_priority_word_style_over_heuristic():
    level, source, recovered = resolve_paragraph_heading_level(
        assignment=None,
        word_style_level=2,
        heuristic_level=3,
        auto_headings=True,
    )
    assert level == 2
    assert source == "word_style"
    assert recovered is None


def test_ai_level_2_survives_to_docx_heading_2():
    recovery = ai_result_to_recovery_payload(
        SAMPLE_AI_RESPONSE,
        original_paragraphs=ORIGINAL_PARAGRAPHS,
    )
    assert recovery["recovery_mode"] == "ai_reconstructed"
    assert recovery["ai_powered"] is True

    doc = Document()
    for para in ORIGINAL_PARAGRAPHS:
        doc.add_paragraph(para)

    apply_result = rebuild_document_from_recovery(doc, recovery)
    assert apply_result is not None
    assert apply_result.ai_powered is True

    je3_paras = [a for a in apply_result.assignments if "Journal Entry 3" in a.text]
    assert len(je3_paras) == 1
    assert je3_paras[0].level == 2
    body_paras = [a for a in apply_result.assignments if a.text.startswith("Body about")]
    assert len(body_paras) == 1
    assert body_paras[0].level is None

    ai_locked = [a for a in apply_result.assignments if a.is_ai_locked]
    assert len(ai_locked) >= 3
    journal_entry = next(
        a for a in ai_locked if "Journal Entry 3" in a.text
    )
    assert journal_entry.level == 2
    assert journal_entry.source == "ai"
    assert journal_entry.confidence == 0.91

    job = FormatJob(
        font_family="Times New Roman",
        font_size_pt=12,
        line_spacing=1.5,
        alignment="left",
        first_line_indent=False,
        space_before_pt=0,
        space_after_pt=0,
        margin_preset="normal",
        page_number_position="none",
        auto_headings=True,
        heading_all_caps=False,
        auto_justify_refs=False,
    )
    debug = format_document_full(
        doc,
        job,
        apply_result.assignments,
        structure_debug=True,
        recovery_mode=apply_result.recovery_mode,
        ai_powered=apply_result.ai_powered,
    )
    assert debug is not None
    assert debug.recovery_mode == "ai_reconstructed"
    assert debug.ai_powered is True

    journal_diag = next(h for h in debug.headings if "Journal Entry 3" in h.paragraph)
    assert journal_diag.source == "ai"
    assert journal_diag.level == 2
    assert journal_diag.applied_style == "Heading 2"

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    final = Document(buf)
    styles = [p.style.name for p in final.paragraphs if "Journal Entry 3" in (p.text or "")]
    assert styles == ["Heading 2"]


def test_heuristic_embedded_journal_entry_splits_and_styles():
    doc = Document()
    doc.add_paragraph(
        "Journal Entry 4: Comparison of Two Historical Concepts\n"
        "When comparing two contrasting concepts such as the marketing mix."
    )
    job = FormatJob(
        font_family="Times New Roman",
        font_size_pt=12,
        line_spacing=1.5,
        alignment="left",
        first_line_indent=False,
        space_before_pt=0,
        space_after_pt=0,
        margin_preset="normal",
        page_number_position="none",
        auto_headings=True,
        heading_all_caps=False,
        auto_justify_refs=False,
    )
    format_document_full(doc, job, None)
    texts = [p.text for p in doc.paragraphs]
    assert texts[0] == "Journal Entry 4: Comparison of Two Historical Concepts"
    assert texts[1].startswith("When comparing")
    assert doc.paragraphs[0].style.name == "Heading 2"
    assert doc.paragraphs[1].style.name == "Normal"
