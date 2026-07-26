"""Document reconstruction engine — structure before formatting."""

from __future__ import annotations

from docx import Document

from formatter.document_reconstruction import (
    build_expected_structure,
    reconstruct_blocks,
    reconstruct_document_before_format,
)
from formatter.requirement_headings import extract_format_section_labels
from tests.conftest import run_format_pipeline
from formatter import FormatJob


def test_build_expected_structure_prioritizes_requirements():
    sections = build_expected_structure(
        document_type="essay",
        required_sections=["Introduction", "Body Paragraph 1", "Conclusion", "References"],
    )
    keys = [s.canonical for s in sections]
    assert "introduction" in keys
    assert "body paragraph 1" in keys
    req = [s for s in sections if s.source == "requirement"]
    assert len(req) == 4
    assert all(s.priority == 1.0 for s in req)


def test_reconstruct_merged_introduction():
    blocks = reconstruct_blocks(
        ["Introduction Artificial Intelligence has become important in education."],
        document_type="essay",
    )
    assert blocks[0].kind == "heading"
    assert blocks[0].text == "Introduction"
    assert blocks[1].kind == "body"
    assert blocks[1].text.startswith("Artificial Intelligence")


def test_reconstruct_merged_journal_entry():
    blocks = reconstruct_blocks(
        ["Journal Entry 1: Reflection on Week 1 Today I learned about entrepreneurship."],
        document_type="learning_journal",
    )
    assert blocks[0].text == "Journal Entry 1: Reflection on Week 1"
    assert blocks[1].text.startswith("Today I learned")


def test_reconstruct_merged_references():
    blocks = reconstruct_blocks(
        ["References Smith, J. (2022). Example article. Journal, 1(1), 1-10."],
        document_type="essay",
    )
    assert blocks[0].text == "References"
    assert "Smith" in blocks[1].text


def test_full_pipeline_with_requirements():
    brief = "Introduction\nBody paragraph 1:\nConclusion\nReferences"
    labels = extract_format_section_labels(brief)
    merged = (
        "Introduction AI is growing in education. Body Paragraph 1 The main argument is "
        "that AI helps students. Conclusion Institutions must adapt. References "
        "Smith, J. (2024). AI in Education. Oxford."
    )
    doc = Document()
    doc.add_paragraph(merged)
    job = FormatJob(
        font_family="Times New Roman",
        font_size_pt=12,
        line_spacing=1.5,
        alignment="left",
        first_line_indent=False,
        space_before_pt=0,
        space_after_pt=12,
        margin_preset="normal",
        page_number_position="none",
        auto_headings=True,
        heading_all_caps=False,
        auto_justify_refs=True,
        requirement_headings=True,
    )
    run_format_pipeline(doc, job, document_type="essay", required_sections=labels)
    texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    assert texts[0] == "Introduction"
    assert any(t.startswith("Body Paragraph 1") for t in texts)
    assert any(t == "References" for t in texts)
