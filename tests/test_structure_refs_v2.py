"""Reference-section latching for Formatter V2 structure extractors."""

from __future__ import annotations

import io

import pytest
from docx import Document

from formatter_v2.pipeline import format_document_v2
from formatter_v2.spec import ParagraphRole, StyleName, UserOverrides
from formatter_v2.structure.from_heuristics import HeuristicsExtractor
from formatter_v2.structure.from_word_styles import (
    WordStylesExtractor,
    implausible_heading_notices,
)
from formatter_v2.structure.references import (
    REFS_HEADING_TITLES,
    is_references_heading,
    looks_like_reference_entry,
    normalize_refs_heading,
)


@pytest.mark.parametrize("title", sorted(REFS_HEADING_TITLES))
def test_each_references_heading_title_is_recognised(title: str) -> None:
    assert is_references_heading(title)
    assert is_references_heading(title.title())
    assert is_references_heading(title.upper())


@pytest.mark.parametrize(
    "raw",
    [
        "6. References",
        "6) References",
        "VII. Bibliography",
        "VII) Works Cited",
        "A. Reference List",
        "B. Literature Cited",
        "3. Sources",
        "IV. Works Consulted",
        "C. List of References",
        "1. References and Bibliography",
        "  12.  Literature  ",
    ],
)
def test_numbered_references_headings_are_recognised(raw: str) -> None:
    assert is_references_heading(raw), normalize_refs_heading(raw)


def test_normalize_strips_punctuation_and_numbering() -> None:
    assert normalize_refs_heading("6. References!") == "references"
    assert normalize_refs_heading("Works Cited:") == "works cited"


def test_heading_latch_with_numbered_title() -> None:
    lines = [
        "Introduction",
        "Body about coastal risk and municipal budgets.",
        "6. References",
        "Smith, J. (2020). Coastal governance. Journal; https://example.com",
        "Doe, A. B. (2019). Flood maps. Publisher.",
    ]
    model = HeuristicsExtractor().extract(lines)
    assert model.references[0].role == ParagraphRole.REFERENCES_HEADING
    assert model.references[0].text == "6. References"
    assert len(model.references) == 3
    assert all(b.role == ParagraphRole.REFERENCES_ENTRY for b in model.references[1:])


def test_numbered_reference_list_becomes_references_not_list_items() -> None:
    lines = [
        "Introduction",
        "The literature is summarised below.",
        "1. Smith, J. (2020). Coastal governance and municipal budgets.",
        "2. Doe, A. (2019). Flood risk mapping for cities.",
        "3. Roe, B. (2021). Adaptation finance in practice.",
    ]
    model = HeuristicsExtractor().extract(lines)
    assert not any(b.role == ParagraphRole.LIST_NUMBER for b in model.body)
    assert len(model.references) >= 3
    assert all(b.role == ParagraphRole.REFERENCES_ENTRY for b in model.references)
    assert [b.text[:2] for b in model.references] == ["1.", "2.", "3."]


def test_bracketed_ieee_reference_list_becomes_references() -> None:
    lines = [
        "Methods",
        "Prior work is summarised in the list below.",
        "[1] Smith, J. (2020). Coastal governance.",
        "[2] Doe, A. (2019). Flood risk mapping.",
        "[3] Roe, B. (2021). Adaptation finance.",
    ]
    model = HeuristicsExtractor().extract(lines)
    assert not any(b.role == ParagraphRole.LIST_NUMBER for b in model.body)
    assert len(model.references) == 3
    assert all(b.role == ParagraphRole.REFERENCES_ENTRY for b in model.references)
    assert model.references[0].text.startswith("[1]")


def test_genuine_numbered_list_in_body_stays_list_number() -> None:
    lines = [
        "Methods",
        "The study proceeded in three steps:",
        "1. Collect municipal budget data for coastal cities.",
        "2. Estimate flood exposure under RCP scenarios carefully.",
        "3. Compare adaptation instruments across jurisdictions.",
        "Results",
        "Exposure rose in every modelled city.",
    ]
    model = HeuristicsExtractor().extract(lines)
    list_items = [b for b in model.body if b.role == ParagraphRole.LIST_NUMBER]
    assert len(list_items) == 3
    assert model.references == []


def test_content_latch_requires_three_consecutive_from_end() -> None:
    lines = [
        "Discussion",
        "Smith, J. (2020). Only two bibliography-like lines is not enough.",
        "Doe, A. (2019). Second bibliography-like line with doi:10.1/xyz.",
        "Conclusion continues as ordinary prose without a third entry.",
    ]
    model = HeuristicsExtractor().extract(lines)
    assert model.references == []


def test_content_latch_stops_at_first_non_matching() -> None:
    lines = [
        "Body paragraph before the bibliography tail.",
        "This ordinary sentence is not a reference entry at all.",
        "Smith, J. (2020). Coastal governance; doi:10.1/abc.",
        "Doe, A. B. (2019). Flood maps. http://example.com/paper",
        "Roe, C. (2021). Adaptation finance; https://example.org",
    ]
    model = HeuristicsExtractor().extract(lines)
    assert len(model.references) == 3
    assert "ordinary sentence" in model.body[-1].text


def test_year_parens_with_semicolon_counts_as_reference() -> None:
    assert looks_like_reference_entry(
        "Smith, J. (2020). Title of article; Journal Name."
    )


def test_word_styles_numbered_list_style_still_latches_ieee_tail() -> None:
    doc = Document()
    doc.add_paragraph("Introduction", style="Heading 1")
    doc.add_paragraph("Body text about sensors.")
    for line in (
        "[1] Smith, J. (2020). Coastal governance.",
        "[2] Doe, A. (2019). Flood risk mapping.",
        "[3] Roe, B. (2021). Adaptation finance.",
    ):
        doc.add_paragraph(line, style="List Number")
    model = WordStylesExtractor().extract(doc)
    assert len(model.references) == 3
    assert all(b.role == ParagraphRole.REFERENCES_ENTRY for b in model.references)
    assert not any(b.role == ParagraphRole.LIST_NUMBER for b in model.body)


def test_implausible_heading_notice_when_too_many_h1() -> None:
    from formatter_v2.render.document import Block
    from formatter_v2.render.model import DocumentModel

    body = [Block(ParagraphRole.HEADING_1, f"Section {i}") for i in range(15)]
    body.extend(Block(ParagraphRole.BODY, f"Body line {i}.") for i in range(10))
    model = DocumentModel(body=body)
    notices = implausible_heading_notices(model)
    assert len(notices) == 1
    assert notices[0].severity == "info"
    assert "headings" in notices[0].message


def _appendix_after_refs_lines() -> list[str]:
    return [
        "Introduction",
        "Body paragraph about coastal risk and municipal budgets.",
        "References",
        "Smith, J. (2020). Coastal governance and municipal budgets.",
        "Doe, A. (2019). Flood risk mapping for delta cities.",
        "Roe, B. (2021). Adaptation finance in practice.",
        "8 Appendix A",
        "Supplementary tables that must not become bibliography entries.",
    ]


def test_appendix_after_references_breaks_the_latch() -> None:
    lines = _appendix_after_refs_lines()
    heuristic = HeuristicsExtractor().extract(lines)
    assert [b.text for b in heuristic.references] == [
        "References",
        "Smith, J. (2020). Coastal governance and municipal budgets.",
        "Doe, A. (2019). Flood risk mapping for delta cities.",
        "Roe, B. (2021). Adaptation finance in practice.",
    ]
    assert not any(
        isinstance(b.text, str) and "Appendix" in b.text for b in heuristic.references
    )
    assert any(
        isinstance(b.text, str) and b.text == "8 Appendix A" for b in heuristic.appendices
    )
    assert any(
        isinstance(b.text, str) and "Supplementary tables" in b.text
        for b in heuristic.appendices
    )

    doc = Document()
    for line in lines:
        doc.add_paragraph(line)
    styled = WordStylesExtractor().extract(doc)
    assert styled.references[0].text == "References"
    assert len(styled.references) == 4
    assert not any(
        isinstance(b.text, str) and "Appendix" in b.text for b in styled.references
    )
    assert any(
        isinstance(b.text, str) and b.text == "8 Appendix A" for b in styled.appendices
    )


def test_appendix_heading_gets_appendix_role() -> None:
    model = HeuristicsExtractor().extract(_appendix_after_refs_lines())
    appendix = next(b for b in model.appendices if b.text == "8 Appendix A")
    assert appendix.role == ParagraphRole.APPENDIX_HEADING
    following = next(
        b
        for b in model.appendices
        if isinstance(b.text, str) and "Supplementary tables" in b.text
    )
    assert following.role == ParagraphRole.BODY


def test_annex_and_glossary_also_break_the_latch() -> None:
    base = [
        "Introduction",
        "Body paragraph about methods and data sources.",
        "References",
        "Smith, J. (2020). Coastal governance and municipal budgets.",
        "Doe, A. (2019). Flood risk mapping for delta cities.",
    ]
    annex_model = HeuristicsExtractor().extract(
        base + ["Annex", "Extra legal materials follow in this annex."]
    )
    assert annex_model.references[0].text == "References"
    assert all("Annex" not in str(b.text) for b in annex_model.references)
    annex = next(b for b in annex_model.body if b.text == "Annex")
    assert annex.role == ParagraphRole.APPENDIX_HEADING

    glossary_model = HeuristicsExtractor().extract(
        base + ["Glossary", "Adaptation: adjustment to actual or expected climate."]
    )
    glossary = next(b for b in glossary_model.body if b.text == "Glossary")
    assert glossary.role == ParagraphRole.HEADING_1
    assert all("Glossary" not in str(b.text) for b in glossary_model.references)


def test_references_at_end_of_document_still_latch_all_entries() -> None:
    lines = [
        "Introduction",
        "Body paragraph about coastal risk and municipal budgets.",
        "References",
        "Smith, J. (2020). Coastal governance and municipal budgets.",
        "Doe, A. (2019). Flood risk mapping for delta cities.",
        "Roe, B. (2021). Adaptation finance in practice.",
    ]
    model = HeuristicsExtractor().extract(lines)
    assert len(model.references) == 4
    assert model.references[0].role == ParagraphRole.REFERENCES_HEADING
    assert all(b.role == ParagraphRole.REFERENCES_ENTRY for b in model.references[1:])
    assert [b.text for b in model.references[1:]] == lines[3:]


def test_pipeline_surfaces_implausible_heading_notice() -> None:
    doc = Document()
    for i in range(15):
        doc.add_paragraph(f"Chapter Heading Number {i}", style="Heading 1")
    for i in range(10):
        doc.add_paragraph(f"A substantial body paragraph {i} about the topic.")
    buf = io.BytesIO()
    doc.save(buf)
    result = format_document_v2(buf.getvalue(), UserOverrides(), StyleName.HARVARD)
    assert result.extractor_name == "word_styles"
    assert any(n.severity == "info" and n.field == "structure.headings" for n in result.notices)


def _word_doc_from_lines(lines: list[str]) -> Document:
    doc = Document()
    for line in lines:
        doc.add_paragraph(line)
    return doc


def test_appendix_after_references_breaks_the_latch() -> None:
    lines = [
        "Introduction",
        "Body paragraph about coastal risk and municipal budgets.",
        "References",
        "Smith, J. (2020). Coastal governance and municipal budgets.",
        "Doe, A. (2019). Flood risk mapping for delta cities.",
        "8 Appendix A",
        "Supplementary tables that must not become bibliography entries.",
    ]
    heuristic = HeuristicsExtractor().extract(lines)
    styled = WordStylesExtractor().extract(_word_doc_from_lines(lines))
    for model in (heuristic, styled):
        assert all(
            "Appendix" not in str(b.text) and "Supplementary" not in str(b.text)
            for b in model.references
        )
        assert any(
            isinstance(b.text, str) and b.text == "8 Appendix A"
            for b in model.appendices
        )
        assert any(
            isinstance(b.text, str) and "Supplementary tables" in b.text
            for b in model.appendices
        )
        assert [b.text for b in model.references if b.role == ParagraphRole.REFERENCES_ENTRY] == [
            "Smith, J. (2020). Coastal governance and municipal budgets.",
            "Doe, A. (2019). Flood risk mapping for delta cities.",
        ]


def test_appendix_heading_gets_appendix_role() -> None:
    lines = [
        "Introduction",
        "Body paragraph about coastal risk and municipal budgets.",
        "References",
        "Smith, J. (2020). Coastal governance and municipal budgets.",
        "8 Appendix A",
        "Supplementary tables that must not become bibliography entries.",
    ]
    heuristic = HeuristicsExtractor().extract(lines)
    styled = WordStylesExtractor().extract(_word_doc_from_lines(lines))
    for model in (heuristic, styled):
        appendix = next(b for b in model.appendices if b.text == "8 Appendix A")
        assert appendix.role == ParagraphRole.APPENDIX_HEADING
        following = next(
            b
            for b in model.appendices
            if isinstance(b.text, str) and "Supplementary" in b.text
        )
        assert following.role == ParagraphRole.BODY


def test_annex_and_glossary_also_break_the_latch() -> None:
    def _run(breaker: str, expected_role: ParagraphRole) -> None:
        lines = [
            "Introduction",
            "Body paragraph about coastal risk and municipal budgets.",
            "References",
            "Smith, J. (2020). Coastal governance and municipal budgets.",
            breaker,
            "Material that belongs after the bibliography section.",
        ]
        heuristic = HeuristicsExtractor().extract(lines)
        styled = WordStylesExtractor().extract(_word_doc_from_lines(lines))
        for model in (heuristic, styled):
            assert all(breaker not in str(b.text) for b in model.references)
            container = (
                model.appendices
                if expected_role == ParagraphRole.APPENDIX_HEADING
                else model.body
            )
            hit = next(b for b in container if b.text == breaker)
            assert hit.role == expected_role
            assert any(
                isinstance(b.text, str) and "Material that belongs" in b.text
                for b in (model.appendices if expected_role == ParagraphRole.APPENDIX_HEADING else model.body)
            )

    _run("Annex", ParagraphRole.APPENDIX_HEADING)
    _run("Glossary", ParagraphRole.HEADING_1)


def test_references_at_end_of_document_still_latch_all_entries() -> None:
    lines = [
        "Introduction",
        "Body paragraph about coastal risk and municipal budgets.",
        "References",
        "Smith, J. (2020). Coastal governance and municipal budgets.",
        "Doe, A. (2019). Flood risk mapping for delta cities.",
        "Roe, B. (2021). Adaptation finance in practice.",
    ]
    heuristic = HeuristicsExtractor().extract(lines)
    styled = WordStylesExtractor().extract(_word_doc_from_lines(lines))
    for model in (heuristic, styled):
        assert model.references[0].role == ParagraphRole.REFERENCES_HEADING
        assert [b.text for b in model.references[1:]] == [
            "Smith, J. (2020). Coastal governance and municipal budgets.",
            "Doe, A. (2019). Flood risk mapping for delta cities.",
            "Roe, B. (2021). Adaptation finance in practice.",
        ]
        assert not any(
            b.role == ParagraphRole.REFERENCES_ENTRY for b in model.body
        )


_MARKS_BLOB = (
    "Donnelly, J. (2007). *Universal Human Rights in Theory and Practice*. Cornell University Press. "
    "Marks, S. P. (2006). *Human Rights: A Brief Introduction*. "
    "(This is a placeholder for a relevant work by Stephen P. Marks). "
    "United Nations. (1948). *Universal Declaration of Human Rights*. "
    "United Nations. (1966). *International Covenant on Civil and Political Rights*. "
    "United Nations. (1966). *International Covenant on Economic, Social and Cultural Rights*."
)


def test_split_concatenated_marks_reading_list() -> None:
    from formatter_v2.structure.references import split_concatenated_reference_entries

    entries = split_concatenated_reference_entries(_MARKS_BLOB)
    assert len(entries) == 5
    assert entries[0].startswith("Donnelly, J. (2007).")
    assert "Universal Human Rights in Theory and Practice" in entries[0]
    assert "*" not in "".join(entries)
    assert "placeholder" not in "".join(entries).lower()
    assert entries[1].startswith("Marks, S. P. (2006).")
    assert entries[2].startswith("United Nations. (1948).")
    assert entries[3].startswith("United Nations. (1966).") and "Civil and Political" in entries[3]
    assert entries[4].startswith("United Nations. (1966).") and "Economic, Social" in entries[4]


def test_concatenated_references_paragraph_explodes_in_extractor() -> None:
    lines = [
        "Question 1",
        "Human rights are entitlements grounded in dignity.",
        "References",
        _MARKS_BLOB,
    ]
    model = HeuristicsExtractor().extract(lines)
    entries = [b.text for b in model.references if b.role == ParagraphRole.REFERENCES_ENTRY]
    assert len(entries) == 5
    assert all("*" not in text for text in entries)
    assert all("placeholder" not in text.lower() for text in entries)
