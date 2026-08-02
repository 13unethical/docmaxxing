"""Tests for DOCX cover page generation."""

from __future__ import annotations

import io

from docx import Document

from formatter.cover_page import CoverPageData, prepend_cover_page


def test_prepend_cover_page_inserts_page_break():
    doc = Document()
    doc.add_paragraph("Body paragraph one.")
    doc.add_paragraph("Body paragraph two.")

    cover = CoverPageData(
        assignment_title="Machine Learning in Healthcare",
        student_name="Jane Smith",
        university="Example University",
        submission_date="2024-06-01",
    )
    prepend_cover_page(doc, cover, font_family="Times New Roman")

    texts = [p.text for p in doc.paragraphs]
    joined = " ".join(texts)
    assert "Machine Learning in Healthcare" in joined
    assert "Jane Smith" in joined
    assert "Body paragraph one." in joined
    assert len(doc.paragraphs) >= 5

    out = io.BytesIO()
    doc.save(out)
    assert out.getvalue()[:2] == b"PK"


def test_prepend_cover_document_keeps_body():
    from formatter.cover_page import prepend_cover_document

    body = Document()
    body.add_paragraph("Main essay body.")

    cover = Document()
    cover.add_paragraph("Ready Cover Title")
    cover.add_paragraph("Student Name")

    prepend_cover_document(body, cover)
    texts = [p.text for p in body.paragraphs if (p.text or "").strip()]
    assert texts[0] == "Ready Cover Title"
    assert "Student Name" in texts
    assert "Main essay body." in texts
