"""Requirement-based heading split for humanizer-merged section labels."""

from __future__ import annotations

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from formatter import FormatJob, format_document_full
from formatter.requirement_headings import (
    extract_format_section_labels,
    expand_requirement_heading_paragraphs,
    split_paragraph_by_requirement_headings,
)
from services.document_structure_engine import detect_heading_level

BRIEF = """
Your summary should include:
Introduction
· the title of the article and name of the author
Body paragraph 1:
· The main idea of the article
Body paragraph 2:
· your opinion or recommendation
Concluding Paragraph
"""

MERGED_TEXT = (
    "Introduction Artificial intelligence has become one of the fastest growing technologies "
    "in education during the past decade. Universities around the world have started integrating "
    "AI-powered tools into classrooms. Body Paragraph 1 one of the main arguments presented in "
    "the article is that AI improves learning outcomes. Body Paragraph 2 another important issue "
    "is academic integrity. Conclusion therefore educational institutions should encourage "
    "students to proofread formatting carefully. References Smith, J. (2024). Artificial "
    "Intelligence in Education. Journal of Learning."
)


def test_extract_sections_from_reading_summary_brief():
    labels = extract_format_section_labels(BRIEF)
    assert "Introduction" in labels
    assert any("body paragraph 1" in x.lower() for x in labels)
    assert any("body paragraph 2" in x.lower() for x in labels)
    assert any("concluding" in x.lower() or "conclusion" in x.lower() for x in labels)


def test_split_introduction_at_paragraph_start():
    labels = extract_format_section_labels(BRIEF)
    segments = split_paragraph_by_requirement_headings(
        "Introduction Artificial intelligence has become popular.",
        labels,
    )
    assert segments[0][0] == "Introduction"
    assert segments[0][1].startswith("Artificial intelligence")


def test_split_conclusion_mid_paragraph():
    labels = extract_format_section_labels(BRIEF)
    segments = split_paragraph_by_requirement_headings(
        "Some text about integrity. Conclusion therefore institutions should act.",
        labels,
    )
    headings = [h for h, _ in segments if h]
    assert any(h.lower().startswith("conclusion") for h in headings)


def test_split_references_before_citation():
    labels = extract_format_section_labels(BRIEF) + ["References"]
    segments = split_paragraph_by_requirement_headings(
        "References Smith, J. (2024). Artificial Intelligence.",
        labels,
    )
    assert segments[0][0] == "References"
    assert "Smith" in segments[0][1]


def test_format_pipeline_splits_merged_headings():
    doc = Document()
    doc.add_paragraph(MERGED_TEXT)
    labels = extract_format_section_labels(BRIEF) + ["References", "Conclusion"]
    job = FormatJob(
        font_family="Times New Roman",
        font_size_pt=12,
        line_spacing=1.5,
        alignment="left",
        first_line_indent=False,
        space_before_pt=0,
        space_after_pt=0,
        margin_preset="normal",
        page_number_position="none",
        auto_headings=True,
        heading_all_caps=False,
        auto_justify_refs=False,
        requirement_headings=True,
        heading_size_pt=18,
    )
    format_document_full(doc, job, None, required_sections=labels)
    texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    assert texts[0] == "Introduction"
    assert any(t.lower().startswith("body paragraph 1") for t in texts)
    assert any(t == "References" for t in texts)
    intro_idx = texts.index("Introduction")
    assert doc.paragraphs[intro_idx].runs[0].font.bold is True
    assert doc.paragraphs[intro_idx].runs[0].font.size.pt == 18


def test_extract_references_when_brief_mentions_reference_list():
    brief = "Submit a reference list in APA format with at least 5 sources."
    labels = extract_format_section_labels(brief)
    assert "References" in labels


def test_split_references_before_citations_mid_paragraph():
    labels = ["References"]
    text = (
        "These findings suggest further research is needed. References Smith, J. (2024). "
        "Artificial Intelligence. Oxford Press. Brown, T. (2023). Digital Learning."
    )
    segments = split_paragraph_by_requirement_headings(text, labels)
    headings = [h for h, _ in segments if h]
    assert headings == ["References"]
    ref_body = next(b for h, b in segments if h == "References")
    assert ref_body.startswith("Smith, J.")


def test_split_discussion_with_lowercase_follow_when_in_brief():
    labels = extract_format_section_labels(BRIEF) + ["Discussion", "Methodology", "Results"]
    text = "Some prior text. Discussion these findings suggest that structure matters."
    segments = split_paragraph_by_requirement_headings(text, labels)
    assert any(h and h.lower() == "discussion" for h, _ in segments)


def test_collapse_double_spaces_in_expand():
    doc = Document()
    doc.add_paragraph("Word1    Word2   Word3")
    labels = ["Introduction"]
    expand_requirement_heading_paragraphs(doc, labels)
    assert doc.paragraphs[0].text == "Word1 Word2 Word3"


def test_no_false_introduction_when_mentioned_in_prose():
    """Do not split when Introduction is mentioned inside a sentence about formatting."""
    labels = extract_format_section_labels(BRIEF)
    text = (
        "Students often submit assignments with hidden formatting errors because these tools "
        "frequently destroy its original formatting Introduction suddenly appears in the middle "
        "of a sentence while paragraphs become merged together."
    )
    segments = split_paragraph_by_requirement_headings(text, labels)
    headings = [h for h, _ in segments if h]
    assert not headings
    assert len(segments) == 1
    assert segments[0][0] is None


def test_introduction_only_once_in_long_paragraph():
    labels = extract_format_section_labels(BRIEF)
    used: set[str] = set()
    text = (
        "Introduction Artificial intelligence is growing. Later the word Introduction "
        "suddenly appears again in prose."
    )
    segments = split_paragraph_by_requirement_headings(text, labels, used_labels=used)
    intro_headings = [h for h, _ in segments if h and h.lower() == "introduction"]
    assert len(intro_headings) == 1
    assert "introduction" in used


def test_detect_heading_level_for_body_paragraph_label():
    labels = frozenset({"introduction", "body paragraph 1", "concluding paragraph"})
    assert detect_heading_level("Body paragraph 1", True, requirement_labels=labels) == 2
