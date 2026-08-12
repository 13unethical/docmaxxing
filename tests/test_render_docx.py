"""XML-level tests for Formatter V2 DOCX rendering (no PDF/PNG)."""

from __future__ import annotations

import io

import pytest
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from formatter_v2.fixtures.sample_document import assert_all_roles_covered, sample_blocks
from formatter_v2.profiles import load_profile
from formatter_v2.render.document import Block, render_document
from formatter_v2.render.styles import style_name_for_role
from formatter_v2.resolve import resolve_format_spec
from formatter_v2.spec import ParagraphRole, StyleName, TextCase, UserOverrides


def _render(style: StyleName):
    profile = load_profile(style)
    result = resolve_format_spec(profile, UserOverrides())
    doc = render_document(sample_blocks(), result.spec)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return Document(buf), result.spec


def _paras_with_style(doc, style_name: str):
    return [p for p in doc.paragraphs if p.style and p.style.name == style_name]


def _style(doc, name: str):
    return doc.styles[name]


@pytest.fixture(scope="module")
def sample() -> list[Block]:
    blocks = sample_blocks()
    assert_all_roles_covered(blocks)
    return blocks


def test_all_roles_produce_paragraph_with_expected_style_name(sample) -> None:
    doc, _spec = _render(StyleName.APA7)
    for role in ParagraphRole:
        expected = style_name_for_role(role)
        matching = _paras_with_style(doc, expected)
        assert matching, f"No paragraph with style {expected!r} for role {role.value}"


@pytest.mark.parametrize(
    "style",
    [StyleName.HARVARD, StyleName.APA7, StyleName.MLA9, StyleName.CHICAGO17, StyleName.IEEE],
)
def test_hanging_indent_is_negative_first_line(style: StyleName) -> None:
    doc, spec = _render(style)
    hanging = spec.roles[ParagraphRole.REFERENCES_ENTRY].hanging_indent_in
    assert hanging > 0
    style_name = style_name_for_role(ParagraphRole.REFERENCES_ENTRY)
    pf = _style(doc, style_name).paragraph_format
    assert pf.first_line_indent is not None
    assert pf.first_line_indent.inches == pytest.approx(-hanging, abs=0.001)


@pytest.mark.parametrize(
    "style",
    [StyleName.HARVARD, StyleName.APA7, StyleName.MLA9, StyleName.CHICAGO17, StyleName.IEEE],
)
def test_reference_entry_left_indent_equals_hanging_value(style: StyleName) -> None:
    doc, spec = _render(style)
    typo = spec.roles[ParagraphRole.REFERENCES_ENTRY]
    style_name = style_name_for_role(ParagraphRole.REFERENCES_ENTRY)
    pf = _style(doc, style_name).paragraph_format
    expected_left = typo.left_indent_in + typo.hanging_indent_in
    assert pf.left_indent is not None
    assert pf.left_indent.inches == pytest.approx(expected_left, abs=0.001)


def test_apa_heading1_line_spacing_is_double() -> None:
    doc, _ = _render(StyleName.APA7)
    pf = _style(doc, "Heading 1").paragraph_format
    assert pf.line_spacing == pytest.approx(2.0)


def test_apa_all_headings_are_12pt() -> None:
    doc, _ = _render(StyleName.APA7)
    for name in ("Title", "Heading 1", "Heading 2", "Heading 3", "Heading 4"):
        size = _style(doc, name).font.size
        assert size == Pt(12)


def test_chicago_reference_entries_are_single_spaced() -> None:
    doc, _ = _render(StyleName.CHICAGO17)
    pf = _style(doc, style_name_for_role(ParagraphRole.REFERENCES_ENTRY)).paragraph_format
    assert pf.line_spacing == pytest.approx(1.0)


def test_ieee_body_is_10pt_and_justified() -> None:
    doc, _ = _render(StyleName.IEEE)
    style = _style(doc, "Normal")
    assert style.font.size == Pt(10)
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    assert style.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY


def test_mla_works_cited_heading_is_not_bold() -> None:
    doc, spec = _render(StyleName.MLA9)
    assert spec.references.heading_text == "Works Cited"
    style = _style(doc, style_name_for_role(ParagraphRole.REFERENCES_HEADING))
    assert style.font.bold is False


def test_heading_font_color_is_black_not_theme_blue() -> None:
    doc, _ = _render(StyleName.APA7)
    for name in ("Heading 1", "Heading 2", "Heading 3"):
        color = _style(doc, name).font.color.rgb
        assert color == RGBColor(0, 0, 0)


def test_toc_heading_style_has_no_outline_level() -> None:
    doc, _ = _render(StyleName.APA7)
    style = _style(doc, style_name_for_role(ParagraphRole.TOC_HEADING))
    ppr = style.element.find(qn("w:pPr"))
    if ppr is None:
        return
    assert ppr.find(qn("w:outlineLvl")) is None


def test_references_heading_style_has_outline_level_zero() -> None:
    doc, _ = _render(StyleName.APA7)
    style = _style(doc, style_name_for_role(ParagraphRole.REFERENCES_HEADING))
    ppr = style.element.find(qn("w:pPr"))
    assert ppr is not None
    outline = ppr.find(qn("w:outlineLvl"))
    assert outline is not None
    assert outline.get(qn("w:val")) == "0"


def test_page_numbering_creates_field_in_header_or_footer() -> None:
    doc, spec = _render(StyleName.APA7)
    assert spec.page_numbering.position.value.startswith("top")
    section = doc.sections[0]
    xml = section.header._element.xml
    assert "PAGE" in xml or "w:fldChar" in xml


def test_apa_does_not_skip_first_page_numbering() -> None:
    doc, spec = _render(StyleName.APA7)
    assert spec.page_numbering.skip_first_page is False
    assert doc.sections[0].different_first_page_header_footer is False


def test_upper_case_role_uses_all_caps_not_transformed_text() -> None:
    """IEEE table captions still use TextCase.UPPER via font.all_caps."""
    profile = load_profile(StyleName.IEEE)
    assert profile.roles[ParagraphRole.TABLE_CAPTION].text_case == TextCase.UPPER
    doc, _ = _render(StyleName.IEEE)
    caption_style = _style(doc, style_name_for_role(ParagraphRole.TABLE_CAPTION))
    assert caption_style.font.all_caps is True
    captions = _paras_with_style(doc, style_name_for_role(ParagraphRole.TABLE_CAPTION))
    assert captions
    # Source text must not be pre-uppercased in the XML
    assert captions[0].text == "Table 1. Adaptation instruments by city"


def test_footnote_text_style_is_created_when_missing_from_template() -> None:
    """Default template lacks 'Footnote Text'; _get_or_create_style must add it."""
    doc, _ = _render(StyleName.CHICAGO17)
    style = doc.styles["Footnote Text"]
    assert style is not None
    assert _paras_with_style(doc, "Footnote Text")


def test_builtin_title_style_has_no_theme_font_attributes() -> None:
    doc, _ = _render(StyleName.APA7)
    rpr = _style(doc, "Title").element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    assert rfonts is not None
    for attr in ("w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"):
        assert rfonts.get(qn(attr)) is None


def test_subtitle_has_no_character_spacing() -> None:
    doc, _ = _render(StyleName.APA7)
    rpr = _style(doc, "Subtitle").element.get_or_add_rPr()
    for junk in ("w:spacing", "w:kern", "w:position", "w:w"):
        assert rpr.find(qn(junk)) is None


def test_builtin_heading_styles_use_configured_font_family() -> None:
    doc, spec = _render(StyleName.APA7)
    expected = spec.roles[ParagraphRole.HEADING_1].font_family.value
    for name in ("Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"):
        style = _style(doc, name)
        assert style.font.name == expected
        rfonts = style.element.get_or_add_rPr().find(qn("w:rFonts"))
        assert rfonts is not None
        assert rfonts.get(qn("w:ascii")) == expected
        assert rfonts.get(qn("w:asciiTheme")) is None


def test_title_style_has_no_paragraph_border() -> None:
    doc, _ = _render(StyleName.APA7)
    ppr = _style(doc, "Title").element.pPr
    if ppr is not None:
        assert ppr.find(qn("w:pBdr")) is None


def test_no_style_carries_paragraph_borders() -> None:
    doc, _ = _render(StyleName.HARVARD)
    for role in ParagraphRole:
        style = _style(doc, style_name_for_role(role))
        ppr = style.element.pPr
        if ppr is None:
            continue
        assert ppr.find(qn("w:pBdr")) is None, style.name


def test_ieee_heading_uses_small_caps() -> None:
    profile = load_profile(StyleName.IEEE)
    assert profile.roles[ParagraphRole.HEADING_1].small_caps is True
    assert profile.roles[ParagraphRole.HEADING_1].text_case == TextCase.PRESERVE
    assert profile.roles[ParagraphRole.REFERENCES_HEADING].small_caps is True
    doc, _ = _render(StyleName.IEEE)
    assert _style(doc, "Heading 1").font.small_caps is True
    assert _style(doc, "Heading 1").font.all_caps is False
    refs = _style(doc, style_name_for_role(ParagraphRole.REFERENCES_HEADING))
    assert refs.font.small_caps is True
    assert refs.font.all_caps is False


def test_small_caps_and_upper_case_together_are_rejected() -> None:
    from formatter_v2.spec import TypographySpec

    with pytest.raises(Exception):
        TypographySpec(small_caps=True, text_case=TextCase.UPPER)


def test_apa_table_caption_renders_as_two_paragraphs() -> None:
    from formatter_v2.render.builder import build_document
    from formatter_v2.render.model import DocumentModel

    profile = load_profile(StyleName.APA7)
    result = resolve_format_spec(profile, UserOverrides())
    assert result.spec.captions.two_line is True
    model = DocumentModel(
        body=[Block(ParagraphRole.TABLE_CAPTION, "Adaptation instruments by city")]
    )
    doc = build_document(model, result.spec)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    reloaded = Document(buf)
    style_name = style_name_for_role(ParagraphRole.TABLE_CAPTION)
    captions = _paras_with_style(reloaded, style_name)
    assert len(captions) == 2
    assert captions[0].text == "Table 1"
    assert any(r.bold for r in captions[0].runs)
    assert captions[1].text == "Adaptation instruments by city"
    assert any(r.italic for r in captions[1].runs)


def test_non_apa_table_caption_renders_as_one_paragraph() -> None:
    from formatter_v2.render.builder import build_document
    from formatter_v2.render.model import DocumentModel

    profile = load_profile(StyleName.HARVARD)
    result = resolve_format_spec(profile, UserOverrides())
    assert result.spec.captions.two_line is False
    model = DocumentModel(
        body=[Block(ParagraphRole.TABLE_CAPTION, "Adaptation instruments by city")]
    )
    doc = build_document(model, result.spec)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    reloaded = Document(buf)
    style_name = style_name_for_role(ParagraphRole.TABLE_CAPTION)
    captions = _paras_with_style(reloaded, style_name)
    assert len(captions) == 1
    assert captions[0].text.startswith("Table 1")
    assert "Adaptation instruments by city" in captions[0].text
