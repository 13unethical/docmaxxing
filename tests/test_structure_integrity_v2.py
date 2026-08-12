"""Tests for refs heading normalisation, content latch with appendices,
Word-style plausibility overrides, document kind, and text integrity.
"""

from __future__ import annotations

import io

import pytest
from docx import Document

from formatter_v2.pipeline import format_document_v2
from formatter_v2.render.document import Block
from formatter_v2.render.model import DocumentModel
from formatter_v2.spec import ParagraphRole, StyleName, UserOverrides
from formatter_v2.structure.document_kind import DocumentKind, detect_kind
from formatter_v2.structure.from_heuristics import HeuristicsExtractor
from formatter_v2.structure.from_word_styles import (
    WordStylesExtractor,
    apply_style_plausibility_overrides,
)
from formatter_v2.structure.references import is_references_heading, normalize_refs_heading
from formatter_v2.structure.text_integrity import normalize_homoglyphs


@pytest.mark.parametrize(
    "raw",
    [
        "7 References",
        "7. References",
        "7) References",
        "VII References",
        "VII. Bibliography",
        "A References",
        "A. Works Cited",
        "12: Literature",
        "3- Sources",
    ],
)
def test_refs_heading_leading_number_without_separator(raw: str) -> None:
    assert is_references_heading(raw), normalize_refs_heading(raw)


def test_references_followed_by_appendices_still_latch() -> None:
    lines = [
        "Introduction",
        "Body paragraph one about the topic under review.",
        "Body paragraph two continues the argument with more detail.",
        "Smith, J. (2020). Coastal governance; doi:10.1/abc.",
        "Doe, A. B. (2019). Flood maps. http://example.com/a",
        "Roe, C. (2021). Adaptation finance; https://example.org/b",
        "Appendix A",
        "Extra materials that are not bibliography entries at all.",
    ]
    # Pad so the bibliography sits in the last third.
    pad = [f"Supporting paragraph number {i} with enough text." for i in range(12)]
    model = HeuristicsExtractor().extract(pad + lines)
    assert len(model.references) == 3
    assert all(b.role == ParagraphRole.REFERENCES_ENTRY for b in model.references)
    assert any(
        isinstance(b.text, str) and b.text.startswith("Appendix") for b in model.appendices
    )
    assert any(
        isinstance(b.text, str) and "Extra materials" in b.text
        for b in model.appendices
    )


def test_long_paragraph_styled_as_heading_becomes_body() -> None:
    long = (
        "This paragraph is deliberately longer than two hundred characters and ends "
        "with a full stop so that a mis-applied Heading 1 style is corrected back to "
        "ordinary body text by the plausibility pass in the Word-styles extractor."
    )
    assert len(long) > 200 and long.endswith(".")
    doc = Document()
    doc.add_paragraph(long, style="Heading 1")
    model = WordStylesExtractor().extract(doc)
    assert len(model.body) == 1
    assert model.body[0].role == ParagraphRole.BODY


def test_numbered_short_line_styled_as_body_becomes_heading() -> None:
    doc = Document()
    doc.add_paragraph("6 Conclusion", style="Normal")
    doc.add_paragraph(
        "A following body paragraph that is long enough to stay body text."
    )
    model = WordStylesExtractor().extract(doc)
    assert model.body[0].role == ParagraphRole.HEADING_1
    assert model.body[0].text == "6 Conclusion"


def test_correctly_styled_document_gets_no_overrides() -> None:
    doc = Document()
    doc.add_paragraph("Introduction", style="Heading 1")
    doc.add_paragraph("A normal body paragraph about coastal adaptation policy.")
    doc.add_paragraph("Methods", style="Heading 1")
    doc.add_paragraph("Another normal body paragraph describing the research design.")
    extractor = WordStylesExtractor()
    model = extractor.extract(doc)
    assert extractor.last_notices == []
    assert [b.role for b in model.body] == [
        ParagraphRole.HEADING_1,
        ParagraphRole.BODY,
        ParagraphRole.HEADING_1,
        ParagraphRole.BODY,
    ]


def test_override_count_reported_in_notices() -> None:
    long = (
        "Another long heading-styled paragraph that exceeds two hundred characters "
        "in total length and finishes with a period so the plausibility override "
        "must reclassify this incorrectly styled block back into ordinary body text."
    )
    assert len(long) > 200 and long.endswith(".")
    model = DocumentModel(
        body=[
            Block(ParagraphRole.HEADING_1, long),
            Block(ParagraphRole.BODY, "3.1 Method"),
            Block(ParagraphRole.BODY, "Short leftover body line without a period"),
        ]
    )
    updated, notices = apply_style_plausibility_overrides(model)
    assert updated.body[0].role == ParagraphRole.BODY
    assert updated.body[1].role == ParagraphRole.HEADING_2
    assert len(notices) == 1
    assert notices[0].severity == "info"
    assert "2" in notices[0].message


def test_mixed_script_word_is_normalised() -> None:
    # Cyrillic 'е' (U+0435) inside an otherwise Latin word.
    raw = "thеory"
    assert "е" in raw
    fixed, n = normalize_homoglyphs(raw)
    assert fixed == "theory"
    assert n == 1


def test_pure_cyrillic_word_is_left_alone() -> None:
    raw = "теория"
    fixed, n = normalize_homoglyphs(raw)
    assert fixed == raw
    assert n == 0


def test_clean_latin_text_produces_no_notice() -> None:
    result = format_document_v2(
        ["Introduction", "Clean Latin body text about sensors."],
        UserOverrides(),
        StyleName.HARVARD,
    )
    assert not any(n.field == "text.homoglyphs" for n in result.notices)


def test_notice_reports_character_count() -> None:
    # >1% of letters substituted.
    lines = [
        "Intrоduction",  # Cyrillic о
        "thеory and prаctice of соastal risk",  # several lookalikes
    ]
    result = format_document_v2(lines, UserOverrides(), StyleName.APA7)
    hits = [n for n in result.notices if n.field == "text.homoglyphs"]
    assert hits
    assert hits[0].severity == "deviation"
    assert any(ch.isdigit() for ch in hits[0].message)


def test_normalisation_preserves_length_and_spacing() -> None:
    raw = "  thеory  and  prаctice\n"
    fixed, n = normalize_homoglyphs(raw)
    assert len(fixed) == len(raw)
    assert fixed.startswith("  ")
    assert "  and  " in fixed
    assert n == 2


def test_slide_script_kind_emits_deviation_notice() -> None:
    lines = [
        "Slide 1",
        "Welcome",
        "Slide 2",
        "Agenda",
        "Slide 3",
        "Summary",
    ]
    result = format_document_v2(lines, UserOverrides(), StyleName.HARVARD)
    assert any(
        n.field == "structure.document_kind" and n.severity == "deviation"
        for n in result.notices
    )
    assert result.docx_bytes[:2] == b"PK"


def test_detect_kind_outline_vs_essay() -> None:
    outline = [
        Block(ParagraphRole.HEADING_1, f"H{i}") for i in range(12)
    ] + [Block(ParagraphRole.BODY, f"Body {i}") for i in range(10)]
    assert detect_kind(outline) == DocumentKind.OUTLINE

    essay = [
        Block(ParagraphRole.HEADING_1, "Intro"),
        *[Block(ParagraphRole.BODY, f"Paragraph {i} with text.") for i in range(25)],
    ]
    assert detect_kind(essay) == DocumentKind.ESSAY
