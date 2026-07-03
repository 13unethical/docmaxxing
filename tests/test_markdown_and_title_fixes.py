"""Markdown cleanup and title detection fixes."""

from __future__ import annotations

from docx import Document

from formatter.markdown_cleanup import clean_markdown_in_document, strip_markdown_text
from services.document_structure_engine import _looks_like_title, detect_heading_level
from tests.conftest import run_format_pipeline
from formatter import FormatJob


def test_strip_orphan_hash_lines():
    assert strip_markdown_text("##") is None
    assert strip_markdown_text("  ###  ") is None


def test_strip_markdown_heading_prefix():
    assert strip_markdown_text("## Introduction") == "Introduction"
    assert strip_markdown_text("### References") == "References"


def test_strip_italic_markers():
    assert strip_markdown_text("*Journal of AI*") == "Journal of AI"
    assert "International Journal" in strip_markdown_text(
        "Crompton, H. *International Journal of Educational Technology in Higher Education*."
    )


def test_title_with_trailing_period_is_detected():
    title = "The Impact of Artificial Intelligence on Higher Education."
    assert _looks_like_title(title)
    assert detect_heading_level(title, True, is_first_nonempty=True) == 1


def test_body_paragraph_heading_keeps_full_subtitle():
    from formatter.document_reconstruction import reconstruct_blocks

    merged = (
        "Body Paragraph 1: Benefits of Artificial Intelligence in Higher Education "
        "Artificial intelligence improves learning outcomes."
    )
    blocks = reconstruct_blocks([merged], document_type="essay")
    assert blocks[0].text == "Body Paragraph 1: Benefits of Artificial Intelligence in Higher Education"
    assert blocks[1].text.startswith("Artificial intelligence")


def test_build_formatted_preview_html_does_not_crash():
    from formatter import FormatJob
    from formatter.preview_html import build_formatted_preview_html

    job = FormatJob(
        font_family="Times New Roman",
        font_size_pt=12,
        line_spacing=1.5,
        alignment="justify",
        first_line_indent=False,
        space_before_pt=0,
        space_after_pt=12,
        margin_preset="normal",
        page_number_position="none",
        auto_headings=True,
        heading_all_caps=False,
        auto_justify_refs=True,
        format_style="harvard",
    )
    html = build_formatted_preview_html(
        "The Impact of AI on Education.\n\nIntroduction\n\nBody text here.",
        job,
    )
    assert "Introduction" in html


def test_format_pipeline_title_and_markdown_cleanup():
    doc = Document()
    doc.add_paragraph("The Impact of Artificial Intelligence on Higher Education.")
    doc.add_paragraph("## Introduction")
    doc.add_paragraph("This is the introduction paragraph with enough text.")
    doc.add_paragraph("##")
    doc.add_paragraph("Body Paragraph 1: Benefits of Artificial Intelligence in Higher Education")
    doc.add_paragraph("Some body text about AI in education continues here.")

    job = FormatJob(
        font_family="Times New Roman",
        font_size_pt=12,
        line_spacing=1.5,
        alignment="justify",
        first_line_indent=False,
        space_before_pt=0,
        space_after_pt=12,
        margin_preset="normal",
        page_number_position="none",
        auto_headings=True,
        heading_all_caps=False,
        auto_justify_refs=True,
        format_style="harvard",
    )
    run_format_pipeline(doc, job, document_type="essay")
    texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    assert texts[0] == "The Impact of Artificial Intelligence on Higher Education."
    assert doc.paragraphs[0].style.name == "Heading 1"
    assert doc.paragraphs[0].paragraph_format.space_after.pt == 24
    from docx.shared import RGBColor

    assert doc.paragraphs[0].runs[0].font.color.rgb == RGBColor(0, 0, 0)
    assert doc.paragraphs[1].runs[0].font.color.rgb == RGBColor(0, 0, 0)
    assert "Introduction" in texts
    assert "##" not in "".join(texts)
    assert any("Benefits of Artificial Intelligence in Higher Education" in t for t in texts)
