"""Tests for StructureConfig.expected_sections matching and pipeline integration."""

from __future__ import annotations

import io
from pathlib import Path

from docx import Document

from formatter_v2.pipeline import format_document_v2
from formatter_v2.profiles import load_profile
from formatter_v2.render.document import Block
from formatter_v2.resolve import resolve_format_spec
from formatter_v2.smartform.prefill import to_user_overrides
from formatter_v2.spec import (
    ExtractedRequirements,
    ParagraphRole,
    StructureConfig,
    StyleName,
    UserOverrides,
)
from formatter_v2.structure.expected_sections import (
    apply_expected_sections,
    normalize_section_key,
)
from formatter_v2.structure.from_heuristics import HeuristicsExtractor

FIXTURES = Path(__file__).resolve().parent / "fixtures"
HUMANIZED_ESSAY = FIXTURES / "humanized_essay.txt"

ESSAY_SECTIONS = [
    "Introduction",
    "Literature Review",
    "Methods",
    "Findings",
    "Discussion",
    "Conclusion",
]


def _humanized_lines() -> list[str]:
    return HUMANIZED_ESSAY.read_text(encoding="utf-8").splitlines()


def test_expected_section_matching_paragraph_becomes_heading() -> None:
    blocks = [Block(ParagraphRole.BODY, "5. Discussion")]
    updated, notices = apply_expected_sections(blocks, ["Discussion"])
    assert updated[0].role == ParagraphRole.HEADING_1
    assert updated[0].text == "Discussion"
    assert any("Found:" in n.message for n in notices)


def test_embedded_section_name_is_split_into_heading_and_body() -> None:
    merged = (
        "4. Findings The results summarized in Table 1 show measurable gains in clarity."
    )
    blocks = [Block(ParagraphRole.BODY, merged)]
    updated, _ = apply_expected_sections(blocks, ["Findings"])
    assert len(updated) == 2
    assert updated[0].role == ParagraphRole.HEADING_1
    assert updated[0].text == "Findings"
    assert updated[1].role == ParagraphRole.BODY
    assert updated[1].text.startswith("The results summarized")


def test_short_trailing_text_is_not_split() -> None:
    blocks = [Block(ParagraphRole.BODY, "Introduction Brief note.")]
    updated, _ = apply_expected_sections(blocks, ["Introduction"])
    assert len(updated) == 1
    assert updated[0].role == ParagraphRole.HEADING_1
    assert updated[0].text == "Introduction"


def test_already_classified_heading_is_not_touched() -> None:
    blocks = [Block(ParagraphRole.HEADING_2, "Custom Methods Overview")]
    updated, notices = apply_expected_sections(blocks, ["Methods"])
    assert updated[0].role == ParagraphRole.HEADING_2
    assert updated[0].text == "Custom Methods Overview"
    assert any("Missing:" in n.message and "Methods" in n.message for n in notices)


def test_expected_section_not_found_adds_warning() -> None:
    blocks = [Block(ParagraphRole.BODY, "Some unrelated paragraph.")]
    _, notices = apply_expected_sections(blocks, ["Abstract"])
    assert any(
        n.field == "structure.expected_sections" and "Missing: Abstract" in n.message
        for n in notices
    )


def test_matching_ignores_case_numbering_and_punctuation() -> None:
    blocks = [Block(ParagraphRole.BODY, "3. METHODS:")]
    updated, _ = apply_expected_sections(blocks, ["methods"])
    assert updated[0].role == ParagraphRole.HEADING_1
    assert normalize_section_key(updated[0].text) == "methods"


def test_pipeline_recovers_headings_from_expected_sections() -> None:
    lines = _humanized_lines()
    overrides = UserOverrides(
        structure=StructureConfig(expected_sections=list(ESSAY_SECTIONS))
    )
    result = format_document_v2(lines, overrides, StyleName.HARVARD)
    doc = Document(io.BytesIO(result.docx_bytes))

    heading_texts = [
        p.text.strip()
        for p in doc.paragraphs
        if p.text.strip() and p.style and p.style.name == "Heading 1"
    ]
    for section in ESSAY_SECTIONS:
        assert any(
            normalize_section_key(section) == normalize_section_key(found)
            for found in heading_texts
        ), f"Missing Heading 1 for {section!r}; got {heading_texts}"


def test_pipeline_without_expected_sections_falls_back_to_heuristics() -> None:
    lines = _humanized_lines()
    with_sections = format_document_v2(
        lines,
        UserOverrides(structure=StructureConfig(expected_sections=list(ESSAY_SECTIONS))),
        StyleName.HARVARD,
    )
    without_sections = format_document_v2(lines, UserOverrides(), StyleName.HARVARD)

    def findings_is_split(docx_bytes: bytes) -> bool:
        doc = Document(io.BytesIO(docx_bytes))
        texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        has_heading_only = any(
            normalize_section_key(t) == "findings" for t in texts
        )
        has_body = any(t.startswith("The results summarized") for t in texts)
        return has_heading_only and has_body

    def findings_still_merged(docx_bytes: bytes) -> bool:
        doc = Document(io.BytesIO(docx_bytes))
        return any("Findings The results" in p.text for p in doc.paragraphs)

    assert findings_is_split(with_sections.docx_bytes)
    assert findings_still_merged(without_sections.docx_bytes)


def test_expected_sections_reach_structure_config_from_brief() -> None:
    extracted = ExtractedRequirements(
        required_sections=[
            "Introduction",
            "Literature Review",
            "Methods",
            "Findings",
            "Discussion",
            "Conclusion",
        ],
        evidence={"required_sections": "Introduction, Literature Review, Methods"},
    )
    profile = load_profile(StyleName.HARVARD)
    prefill = to_user_overrides(extracted, profile)
    assert prefill.overrides.structure is not None
    assert prefill.overrides.structure.expected_sections == extracted.required_sections

    resolved = resolve_format_spec(profile, prefill.overrides)
    assert resolved.spec.structure.expected_sections == extracted.required_sections


def test_heuristics_extractor_applies_expected_sections_before_refs() -> None:
    lines = [
        "Introduction Body starts here with enough characters to split clearly.",
        "References",
        "Smith, J. (2020). Example.",
    ]
    extractor = HeuristicsExtractor()
    model = extractor.extract(
        lines,
        expected_sections=["Introduction"],
    )
    assert model.body[0].role == ParagraphRole.HEADING_1
    assert model.body[0].text == "Introduction"
    assert model.body[1].role == ParagraphRole.BODY
    assert model.body[1].text.startswith("Body starts here")
    assert model.references[0].role == ParagraphRole.REFERENCES_HEADING
