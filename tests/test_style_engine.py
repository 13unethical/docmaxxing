"""Formatting Style Engine tests."""

from __future__ import annotations

from docx import Document

from formatter.format_job import FormatJob
from formatter.style_engine import resolve_active_profile, resolve_contextual_spacing
from styles import load_profile


def test_load_harvard_profile_body_spacing_from_profile_not_hardcoded():
    profile = load_profile("harvard")
    sb, sa = resolve_contextual_spacing(
        profile,
        role="body",
        prev_level=0,
        next_level=0,
        prev_has_text=True,
    )
    assert sa == profile.body.contextual.body_space_after_pt
    assert sa == 12


def test_load_apa_profile_double_spacing():
    profile = load_profile("apa7")
    assert profile.body.paragraph.line_spacing == 2.0
    assert profile.body.paragraph.line_spacing_rule == "double"
    assert profile.body.contextual.body_space_after_pt == 0


def test_load_chicago_profile_body_spacing():
    profile = load_profile("chicago17")
    assert profile.body.paragraph.line_spacing == 1.5
    assert profile.body.contextual.body_space_after_pt == 6


def test_academic_heading_spacing_from_profile():
    profile = load_profile("harvard")
    assert profile.title.space_before_pt == 0
    assert profile.title.space_after_pt == 24
    assert profile.heading2.space_before_pt == 18
    assert profile.heading2.space_after_pt == 6
    assert profile.heading3.space_before_pt == 12
    assert profile.heading3.space_after_pt == 4
    sb, sa = resolve_contextual_spacing(
        profile,
        role="heading2",
        prev_level=0,
        next_level=0,
        prev_has_text=False,
    )
    assert sb == 18
    assert sa == 6


def test_resolve_active_profile_from_job_style():
    job = FormatJob(
        font_family="Times New Roman",
        font_size_pt=12,
        line_spacing=2.0,
        alignment="left",
        first_line_indent=False,
        space_before_pt=0,
        space_after_pt=0,
        margin_preset="normal",
        page_number_position="top_right",
        auto_headings=True,
        heading_all_caps=False,
        auto_justify_refs=True,
        format_style="apa7",
    )
    profile = resolve_active_profile(job)
    assert profile.id == "apa7"
    assert profile.body.paragraph.line_spacing == 2.0


def test_apa_references_hanging_indent_applied():
    from formatter.pipeline import format_document_full

    doc = Document()
    doc.add_paragraph("References")
    doc.add_paragraph("Smith, J. (2020). Example. Publisher.")
    job = FormatJob(
        font_family="Times New Roman",
        font_size_pt=12,
        line_spacing=2.0,
        alignment="left",
        first_line_indent=False,
        space_before_pt=0,
        space_after_pt=0,
        margin_preset="normal",
        page_number_position="none",
        auto_headings=True,
        heading_all_caps=False,
        auto_justify_refs=True,
        format_style="apa7",
    )
    format_document_full(doc, job, None)
    ref_para = doc.paragraphs[1]
    assert ref_para.paragraph_format.left_indent is not None
