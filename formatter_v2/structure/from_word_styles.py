"""Extract DocumentModel from a DOCX that already has Word styles.

Priority path — Word style names are the primary source of truth.
Explicit contradictions (long Heading paragraphs, short numbered Normal
lines) may be overridden with an info notice. Reference latching and
content-based bibliography detection live in ``references``.
"""

from __future__ import annotations

import io
import re
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.text.paragraph import Paragraph

from formatter_v2.render.document import Block
from formatter_v2.render.model import DocumentModel
from formatter_v2.resolve import ResolutionNotice
from formatter_v2.spec import ParagraphRole
from formatter_v2.structure.references import (
    find_heading_latch_index,
    find_heading_latch_end,
    find_content_latch_range,
    is_refs_latch_breaker,
    refs_latch_break_role,
    move_appendices_from_body,
)
from formatter_v2.structure.numbered import classify_numbered_line

_CAPTION_TABLE_RE = re.compile(r"^\s*tables?\s+\d+", re.IGNORECASE)
_CAPTION_FIGURE_RE = re.compile(r"^\s*(figures?|fig\.)\s+\d+", re.IGNORECASE)
_NUMBERED_LIST_RE = re.compile(r"^\s*\d+[.)]\s+")
_SECTION_NUMBER_RE = re.compile(
    r"^(?P<nums>\d+(?:\.\d+)*)(?:[.)])?\s+\S",
)

_BODY_STYLE_NAMES = frozenset(
    {
        "normal",
        "body text",
        "first paragraph",
    }
)

_HEADING_ROLES = frozenset(
    {
        ParagraphRole.HEADING_1,
        ParagraphRole.HEADING_2,
        ParagraphRole.HEADING_3,
        ParagraphRole.HEADING_4,
    }
)

_DEPTH_TO_ROLE = {
    1: ParagraphRole.HEADING_1,
    2: ParagraphRole.HEADING_2,
    3: ParagraphRole.HEADING_3,
    4: ParagraphRole.HEADING_4,
}

_HEURISTIC_BULLET_RE = re.compile(r"^[\u2022\u2023\u25E6\u00B7\u00B7\u2022\u2023\u25E6\u2022\-\*•]\s+")
_HEURISTIC_TABLE_CAP_RE = re.compile(r"^Table\s+\d+", re.IGNORECASE)
_HEURISTIC_FIGURE_CAP_RE = re.compile(r"^(Figure|Fig\.)\s+\d+", re.IGNORECASE)


def _paragraph_indent_inches(paragraph: Paragraph | None) -> float:
    if paragraph is None:
        return 0.0
    pf = paragraph.paragraph_format
    indent = 0.0
    if pf.first_line_indent is not None:
        indent = max(indent, float(pf.first_line_indent.inches))
    if pf.left_indent is not None:
        indent = max(indent, float(pf.left_indent.inches))
    return indent


def _looks_like_block_quote(text: str, paragraph: Paragraph | None) -> bool:
    stripped = text.strip()
    if len(stripped) <= 200:
        return False
    terminal = sum(stripped.count(mark) for mark in ".!?")
    if terminal > 2:
        return False
    return _paragraph_indent_inches(paragraph) > 0.15


def _coerce_document(source: object) -> DocxDocument:
    if isinstance(source, DocxDocument):
        return source
    if isinstance(source, (bytes, bytearray)):
        return Document(io.BytesIO(source))
    if isinstance(source, (str, Path)):
        path = Path(source)
        if path.is_file():
            return Document(str(path))
        raise TypeError(f"Word-styles extractor expects a DOCX path, got {source!r}")
    raise TypeError(f"Unsupported source type for Word-styles extractor: {type(source)!r}")


def document_has_structural_styles(document: DocxDocument) -> bool:
    """True when the DOCX already carries Title / Heading / Quote styles."""
    structural = {
        "Title",
        "Subtitle",
        "Heading 1",
        "Heading 2",
        "Heading 3",
        "Heading 4",
        "Quote",
        "Intense Quote",
        "List Bullet",
        "List Number",
        "Caption",
    }
    for paragraph in document.paragraphs:
        if not (paragraph.text or "").strip():
            continue
        name = getattr(getattr(paragraph, "style", None), "name", None) or ""
        if name in structural:
            return True
        if name.startswith("Heading "):
            return True
    return False


def _style_name(paragraph: Paragraph) -> str:
    return (getattr(getattr(paragraph, "style", None), "name", None) or "").strip()


def _list_role(paragraph: Paragraph, text: str) -> ParagraphRole:
    name = _style_name(paragraph).casefold()
    if "list number" in name or _NUMBERED_LIST_RE.match(text):
        return ParagraphRole.LIST_NUMBER
    return ParagraphRole.LIST_BULLET


def _role_from_style(
    text: str,
    paragraph: Paragraph | None,
    _is_first_nonempty: bool,
) -> ParagraphRole:
    if paragraph is None:
        return ParagraphRole.BODY
    name = _style_name(paragraph)
    lower = name.casefold()

    if name == "Title":
        return ParagraphRole.DOC_TITLE
    if name == "Subtitle":
        return ParagraphRole.SUBTITLE
    if name == "Heading 1":
        return ParagraphRole.HEADING_1
    if name == "Heading 2":
        return ParagraphRole.HEADING_2
    if name == "Heading 3":
        return ParagraphRole.HEADING_3
    if name == "Heading 4":
        return ParagraphRole.HEADING_4
    if name in {"Quote", "Intense Quote"}:
        return ParagraphRole.BLOCK_QUOTE
    if "list" in lower:
        return _list_role(paragraph, text)
    if name == "Caption" or lower == "caption":
        if _CAPTION_FIGURE_RE.match(text):
            return ParagraphRole.FIGURE_CAPTION
        return ParagraphRole.TABLE_CAPTION
    if lower.startswith("heading "):
        try:
            level = int(name.split()[-1])
        except ValueError:
            level = 1
        return _DEPTH_TO_ROLE.get(level, ParagraphRole.BODY)
    return ParagraphRole.BODY


def _plain(block: Block) -> str:
    if isinstance(block.text, str):
        return block.text
    return str(block.text)


def _heading_depth_from_section_number(text: str) -> int | None:
    match = _SECTION_NUMBER_RE.match(text.strip())
    if not match:
        return None
    parts = match.group("nums").split(".")
    return min(len(parts), 4)


def apply_style_plausibility_overrides(
    model: DocumentModel,
) -> tuple[DocumentModel, list[ResolutionNotice]]:
    """Override clear Word-style contradictions; report how many changed."""
    overrides = 0
    new_body: list[Block] = []

    for block in model.body:
        text = _plain(block).strip()
        role = block.role

        if role in _HEADING_ROLES and len(text) > 200 and text.endswith("."):
            new_body.append(Block(ParagraphRole.BODY, block.text))
            overrides += 1
            continue

        if role in {ParagraphRole.BODY, ParagraphRole.BODY_FIRST}:
            if (
                len(text) < 80
                and not text.endswith(".")
                and (depth := _heading_depth_from_section_number(text)) is not None
            ):
                new_body.append(Block(_DEPTH_TO_ROLE[depth], block.text))
                overrides += 1
                continue

        new_body.append(block)

    notices: list[ResolutionNotice] = []
    if overrides:
        notices.append(
            ResolutionNotice(
                field="structure.style_overrides",
                severity="info",
                message=(
                    f"Переразмечено абзацев из-за противоречий стилей Word: {overrides}. "
                    "Длинные «Heading» стали телом; короткие нумерованные Normal — заголовками."
                ),
            )
        )

    updated = DocumentModel(
        cover=model.cover,
        front_matter=list(model.front_matter),
        body=new_body,
        references=list(model.references),
        appendices=list(model.appendices),
    )
    return updated, notices


def implausible_heading_notices(model: DocumentModel) -> list[ResolutionNotice]:
    """Info notices when Word heading density looks unusually high."""
    blocks = list(model.front_matter) + list(model.body) + list(model.references)
    if model.appendices:
        blocks.extend(model.appendices)
    total = len(blocks)
    if total == 0:
        return []

    heading_n = sum(1 for b in blocks if b.role in _HEADING_ROLES)
    body_n = sum(
        1 for b in blocks if b.role in {ParagraphRole.BODY, ParagraphRole.BODY_FIRST}
    )
    if total > 20 and heading_n >= body_n:
        return [
            ResolutionNotice(
                field="structure.headings",
                severity="info",
                message=(
                    "В исходном документе найдено необычно много заголовков "
                    f"(заголовков: {heading_n}, абзацев текста: {body_n}). "
                    "Структуру стоит проверить вручную."
                ),
            )
        ]
    return []


class WordStylesExtractor:
    """Map existing Word paragraph styles onto V2 ``ParagraphRole`` values."""

    def __init__(self) -> None:
        self.last_notices: list[ResolutionNotice] = []

    def extract(self, source: object) -> DocumentModel:
        document = _coerce_document(source)
        texts: list[str] = []
        paragraphs: list[Paragraph | None] = []
        for paragraph in document.paragraphs:
            stripped = (paragraph.text or "").strip()
            if not stripped:
                continue
            texts.append(stripped)
            paragraphs.append(paragraph)

        # References heading latch is based on title matching and stops at a
        # first "section breaker" paragraph (appendix, annex, etc).
        heading_idx = find_heading_latch_index(texts)
        if heading_idx is not None:
            latch_mode = "heading"
            latch_start = heading_idx
            latch_end = find_heading_latch_end(texts, heading_idx)
        else:
            latch_mode = "none"
            latch_start = len(texts)
            latch_end = len(texts)

        body_texts_for_numbering = [
            texts[i] for i in range(len(texts)) if not (latch_start <= i < latch_end)
        ]

        body: list[Block] = []
        body_paras: list[Paragraph | None] = []
        references: list[Block] = []

        seen_nonempty = False
        body_numbering_index = 0

        for i, text in enumerate(texts):
            para = paragraphs[i] if i < len(paragraphs) else None
            if not (text or "").strip():
                continue
            stripped = (text or "").strip()

            if latch_start <= i < latch_end:
                if latch_mode == "heading" and i == latch_start:
                    references.append(Block(ParagraphRole.REFERENCES_HEADING, stripped))
                else:
                    references.append(
                        Block(ParagraphRole.REFERENCES_ENTRY, stripped)
                    )
                continue

            # When the latch is based on a "References" heading, the very first
            # "section breaker" paragraph must be re-typed as appendix heading
            # (or a generic heading). This mirrors split_body_and_references()
            # in formatter_v2.structure.references.
            if (
                latch_mode == "heading"
                and i == latch_end
                and is_refs_latch_breaker(stripped)
            ):
                body.append(Block(refs_latch_break_role(stripped), stripped))
                body_paras.append(para)
                seen_nonempty = True
                body_numbering_index += 1
                continue

            numbered_role = classify_numbered_line(
                body_texts_for_numbering, body_numbering_index
            )
            body_numbering_index += 1
            if numbered_role is not None:
                body.append(Block(numbered_role, stripped))
                body_paras.append(para)
                seen_nonempty = True
                continue

            is_first = not seen_nonempty
            seen_nonempty = True
            role = _role_from_style(stripped, para, is_first)
            body.append(Block(role, stripped))
            body_paras.append(para)

        model = DocumentModel(body=body, references=references)
        model, override_notices = apply_style_plausibility_overrides(model)

        # Content latch: only when no references heading exists.
        # Do it with indices so we can keep body_paras aligned for post-heuristics.
        if not any(b.role == ParagraphRole.REFERENCES_ENTRY for b in model.references):
            body_texts = [_plain(b) for b in model.body]
            found = find_content_latch_range(body_texts, body_paras)
            if found is not None:
                start, end = found
                latched = model.body[start:end]
                model = DocumentModel(
                    cover=model.cover,
                    front_matter=list(model.front_matter),
                    body=list(model.body[:start]) + list(model.body[end:]),
                    references=list(model.references)
                    + [
                        Block(ParagraphRole.REFERENCES_ENTRY, _plain(b))
                        for b in latched
                    ],
                    appendices=list(model.appendices),
                )
                body_paras = body_paras[:start] + body_paras[end:]

        # Post heuristics: run over paragraphs that remained BODY by Word styles.
        # Do not override explicit Word roles (Heading/Title/Quote/List/Captions).
        body_texts = [_plain(b) for b in model.body]
        for i, (block, para) in enumerate(zip(model.body, body_paras)):
            if block.role not in {ParagraphRole.BODY, ParagraphRole.BODY_FIRST}:
                continue
            text = _plain(block).strip()

            # Numbered section/list classifier should keep working even when
            # Word paragraph style is Normal.
            numbered_role = classify_numbered_line(body_texts, i)
            if numbered_role is not None:
                model.body[i] = Block(numbered_role, block.text)
                continue

            if _HEURISTIC_TABLE_CAP_RE.match(text):
                model.body[i] = Block(ParagraphRole.TABLE_CAPTION, block.text)
                continue
            if _HEURISTIC_FIGURE_CAP_RE.match(text):
                model.body[i] = Block(ParagraphRole.FIGURE_CAPTION, block.text)
                continue
            if _HEURISTIC_BULLET_RE.match(text):
                model.body[i] = Block(ParagraphRole.LIST_BULLET, block.text)
                continue
            if _looks_like_block_quote(text, para):
                model.body[i] = Block(ParagraphRole.BLOCK_QUOTE, block.text)
                continue

        self.last_notices = list(override_notices)
        return move_appendices_from_body(model)
