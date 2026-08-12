"""Homoglyph / mixed-script cleanup for Formatter V2.

Normalises lookalike Cyrillic/Greek letters inside otherwise-Latin words.
Pure Cyrillic (or pure Greek) words are left untouched.
"""

from __future__ import annotations

import io
import re
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument

from formatter_v2.resolve import ResolutionNotice

# Cyrillic lookalikes → Latin
_HOMOGLYPHS: dict[str, str] = {
    "а": "a",
    "е": "e",
    "о": "o",
    "с": "c",
    "р": "p",
    "х": "x",
    "у": "y",
    "А": "A",
    "В": "B",
    "Е": "E",
    "К": "K",
    "М": "M",
    "Н": "H",
    "О": "O",
    "Р": "P",
    "С": "C",
    "Т": "T",
    "Х": "X",
    # Greek lookalikes
    "ο": "o",
    "α": "a",
    "ε": "e",
    "ρ": "p",
    "ν": "v",
    "Ο": "O",
    "Α": "A",
    "Ε": "E",
    "Ρ": "P",
    "Ν": "N",
}

_WORD_RE = re.compile(r"\S+", re.UNICODE)


def _is_ascii_latin(ch: str) -> bool:
    return ("a" <= ch <= "z") or ("A" <= ch <= "Z")


def _is_cyrillic(ch: str) -> bool:
    return "\u0400" <= ch <= "\u04FF"


def _is_greek(ch: str) -> bool:
    return "\u0370" <= ch <= "\u03FF"


def _scripts(word: str) -> set[str]:
    found: set[str] = set()
    for ch in word:
        if not ch.isalpha():
            continue
        if _is_ascii_latin(ch):
            found.add("latin")
        elif _is_cyrillic(ch):
            found.add("cyrillic")
        elif _is_greek(ch):
            found.add("greek")
    return found


def _normalize_word(word: str) -> tuple[str, int]:
    scripts = _scripts(word)
    if "latin" not in scripts:
        return word, 0
    if "cyrillic" not in scripts and "greek" not in scripts:
        return word, 0
    replaced = 0
    out: list[str] = []
    for ch in word:
        if ch in _HOMOGLYPHS:
            out.append(_HOMOGLYPHS[ch])
            replaced += 1
        else:
            out.append(ch)
    return "".join(out), replaced


def normalize_homoglyphs(text: str) -> tuple[str, int]:
    """Return ``(fixed_text, substituted_letter_count)``.

    Only mixed-script words are touched. Length and spacing are preserved
    character-for-character aside from the 1:1 homoglyph substitutions.
    """
    if not text:
        return text, 0

    total = 0
    parts: list[str] = []
    last = 0
    for match in _WORD_RE.finditer(text):
        parts.append(text[last : match.start()])
        fixed, n = _normalize_word(match.group(0))
        parts.append(fixed)
        total += n
        last = match.end()
    parts.append(text[last:])
    return "".join(parts), total


def _count_letters(text: str) -> int:
    return sum(1 for ch in text if ch.isalpha())


def integrity_notices(substituted: int, letter_count: int) -> list[ResolutionNotice]:
    if letter_count <= 0 or substituted <= 0:
        return []
    if substituted / letter_count <= 0.01:
        return []
    return [
        ResolutionNotice(
            field="text.homoglyphs",
            severity="deviation",
            message=(
                "В тексте найдены подменённые буквы (кириллица/греческий вместо "
                f"латиницы): исправлено символов — {substituted}. "
                "Стоит проверить исходник на копирование из PDF/скана."
            ),
        )
    ]


def normalize_plain_lines(lines: list[str]) -> tuple[list[str], list[ResolutionNotice]]:
    fixed: list[str] = []
    substituted = 0
    letters = 0
    for line in lines:
        new, n = normalize_homoglyphs(line)
        fixed.append(new)
        substituted += n
        letters += _count_letters(line)
    return fixed, integrity_notices(substituted, letters)


def normalize_docx_document(document: DocxDocument) -> list[ResolutionNotice]:
    """Rewrite paragraph text in-place; return integrity notices."""
    substituted = 0
    letters = 0
    for paragraph in document.paragraphs:
        original = paragraph.text or ""
        letters += _count_letters(original)
        fixed, n = normalize_homoglyphs(original)
        if n == 0:
            continue
        substituted += n
        if paragraph.runs:
            paragraph.runs[0].text = fixed
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.add_run(fixed)
    return integrity_notices(substituted, letters)


def normalize_source(source: object) -> tuple[object, list[ResolutionNotice]]:
    """Normalise homoglyphs before structure extraction.

    Returns a possibly-new source object plus integrity notices.
    """
    if isinstance(source, DocxDocument):
        notices = normalize_docx_document(source)
        return source, notices
    if isinstance(source, (bytes, bytearray)):
        doc = Document(io.BytesIO(source))
        notices = normalize_docx_document(doc)
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue(), notices
    if isinstance(source, list):
        lines, notices = normalize_plain_lines([str(x) for x in source])
        return lines, notices
    if isinstance(source, (str, Path)):
        path = Path(source)
        if path.is_file() and path.suffix.lower() == ".docx":
            doc = Document(str(path))
            notices = normalize_docx_document(doc)
            return doc, notices
        text = path.read_text(encoding="utf-8") if path.is_file() else str(source)
        lines, notices = normalize_plain_lines(text.splitlines())
        return lines, notices
    return source, []
