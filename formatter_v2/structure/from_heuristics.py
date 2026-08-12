"""Heuristic structure extraction — thin adapter over V1 heading detection.

Detection of heading *levels* is delegated to
``services.document_structure_engine.detect_heading_level`` (not rewritten).
Only roles V1 never produced (block quotes, lists, captions) are added here.

Reference latching (heading + content) runs *before* list detection — see
``formatter_v2.structure.references``.
"""

from __future__ import annotations

import io
import re
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.text.paragraph import Paragraph

from services.document_structure_engine import detect_heading_level

from formatter_v2.render.document import Block
from formatter_v2.render.model import DocumentModel
from formatter_v2.spec import ParagraphRole
from formatter_v2.structure.references import (
    move_appendices_from_body,
    split_body_and_references,
)

_BULLET_RE = re.compile(r"^[\u2022\u2023\u25E6\u00B7\-\*•]\s+")
_TABLE_CAP_RE = re.compile(r"^Table\s+\d+", re.IGNORECASE)
_FIGURE_CAP_RE = re.compile(r"^(Figure|Fig\.)\s+\d+", re.IGNORECASE)

_LEVEL_TO_ROLE = {
    0: ParagraphRole.BODY,
    1: ParagraphRole.DOC_TITLE,
    2: ParagraphRole.HEADING_1,
    3: ParagraphRole.HEADING_2,
}


def _coerce_paragraphs(source: object) -> list[tuple[str, Paragraph | None]]:
    """Return (text, optional paragraph) pairs for role detection."""
    if isinstance(source, DocxDocument):
        return [(p.text or "", p) for p in source.paragraphs]
    if isinstance(source, (bytes, bytearray)):
        doc = Document(io.BytesIO(source))
        return [(p.text or "", p) for p in doc.paragraphs]
    if isinstance(source, list):
        return [(str(line), None) for line in source]
    if isinstance(source, (str, Path)):
        path = Path(source)
        if path.is_file() and path.suffix.lower() == ".docx":
            doc = Document(str(path))
            return [(p.text or "", p) for p in doc.paragraphs]
        text = path.read_text(encoding="utf-8") if path.is_file() else str(source)
        return [(line, None) for line in text.splitlines()]
    raise TypeError(f"Unsupported source type for heuristics extractor: {type(source)!r}")


def _paragraph_indent_inches(paragraph: Paragraph | None) -> float:
    if paragraph is None:
        return 0.0
    pf = paragraph.paragraph_format
    indent = 0.0
    if pf.first_line_indent is not None:
        indent = max(indent, float(pf.first_line_indent.inches))
    if pf.left_indent is not None:
        indent = max(indent, float(pf.left_indent.inches))
    return indent


def _looks_like_block_quote(text: str, paragraph: Paragraph | None) -> bool:
    stripped = text.strip()
    if len(stripped) <= 200:
        return False
    terminal = sum(stripped.count(mark) for mark in ".!?")
    if terminal > 2:
        return False
    return _paragraph_indent_inches(paragraph) > 0.15


def _role_for_paragraph(
    text: str,
    paragraph: Paragraph | None,
    is_first_nonempty: bool,
) -> ParagraphRole:
    stripped = text.strip()
    if not stripped:
        return ParagraphRole.BODY

    if _TABLE_CAP_RE.match(stripped):
        return ParagraphRole.TABLE_CAPTION
    if _FIGURE_CAP_RE.match(stripped):
        return ParagraphRole.FIGURE_CAPTION
    if _BULLET_RE.match(stripped):
        return ParagraphRole.LIST_BULLET
    if _looks_like_block_quote(stripped, paragraph):
        return ParagraphRole.BLOCK_QUOTE

    level = detect_heading_level(
        stripped,
        True,
        is_first_nonempty=is_first_nonempty,
    )
    return _LEVEL_TO_ROLE.get(level, ParagraphRole.BODY)


class HeuristicsExtractor:
    """V1 heading levels + a few V2-only role detectors; unknown → BODY."""

    def extract(self, source: object) -> DocumentModel:
        pairs = _coerce_paragraphs(source)
        texts: list[str] = []
        paragraphs: list[Paragraph | None] = []
        for text, paragraph in pairs:
            stripped = (text or "").strip()
            if not stripped:
                continue
            texts.append(stripped)
            paragraphs.append(paragraph)

        body, references = split_body_and_references(
            texts,
            paragraphs,
            _role_for_paragraph,
        )
        model = DocumentModel(body=body, references=references)
        return move_appendices_from_body(model)
