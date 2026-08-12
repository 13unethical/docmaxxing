"""
Formatter V2 — слой вёрстки, часть 1: стили документа.

Принцип: рендерер НЕ красит абзацы напрямую. Он создаёт в документе
настоящие именованные стили Word и назначает их абзацам.

Зачем именно так:
  * поле оглавления работает только по стилям заголовков с outline level;
  * документ остаётся редактируемым — студент видит нормальную панель стилей;
  * роль абзаца живёт на уровне разметки и переживает переписывание текста.

Ловушки python-docx, из-за которых этот файл выглядит сложнее, чем ожидаешь:
  1. Висячий отступ = left_indent + ОТРИЦАТЕЛЬНЫЙ first_line_indent.
     Свойства "hanging" не существует.
  2. Встроенные Heading 1-4 в стандартном шаблоне синие и Calibri Light
     через тему документа. Цвет и шрифт надо задавать принудительно.
  3. Номера страниц — это поля Word (PAGE/NUMPAGES), их нет в API,
     нужен ручной OxmlElement.
  4. Старт нумерации и формат номера живут в w:pgNumType внутри sectPr,
     тоже мимо API.
  5. TITLE_CASE не существует как свойство Word — применяется к строке
     при сборке документа, не здесь.
"""

from __future__ import annotations

from docx.document import Document as DocxDocument
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor
from docx.styles.style import ParagraphStyle

from ..spec import (
    Alignment,
    FormatSpec,
    NumberFormat,
    PageNumberPosition,
    PageSize,
    ParagraphRole,
    TextCase,
    TypographySpec,
)

BLACK = RGBColor(0x00, 0x00, 0x00)

# --------------------------------------------------------------------------
# Соответствие ролей и стилей Word
# --------------------------------------------------------------------------

# Роли, которые ложатся на ВСТРОЕННЫЕ стили Word. Встроенные используются там,
# где это даёт функциональность: outline level для оглавления, корректное
# поведение списков, распознавание сносок.
BUILTIN_STYLE_FOR_ROLE: dict[ParagraphRole, str] = {
    ParagraphRole.DOC_TITLE: "Title",
    ParagraphRole.SUBTITLE: "Subtitle",
    ParagraphRole.HEADING_1: "Heading 1",
    ParagraphRole.HEADING_2: "Heading 2",
    ParagraphRole.HEADING_3: "Heading 3",
    ParagraphRole.HEADING_4: "Heading 4",
    ParagraphRole.BODY: "Normal",
    ParagraphRole.LIST_BULLET: "List Bullet",
    ParagraphRole.LIST_NUMBER: "List Number",
    ParagraphRole.FOOTNOTE: "Footnote Text",
}

# Роли со своими стилями. Префикс DM исключает столкновение с чужими стилями,
# если документ пришёл от пользователя с собственной разметкой.
CUSTOM_STYLE_FOR_ROLE: dict[ParagraphRole, str] = {
    ParagraphRole.ABSTRACT_HEADING: "DM Abstract Heading",
    ParagraphRole.ABSTRACT: "DM Abstract",
    ParagraphRole.KEYWORDS: "DM Keywords",
    ParagraphRole.BODY_FIRST: "DM Body First",
    ParagraphRole.BLOCK_QUOTE: "DM Block Quote",
    ParagraphRole.TABLE_CAPTION: "DM Table Caption",
    ParagraphRole.TABLE_HEADER: "DM Table Header",
    ParagraphRole.TABLE_CELL: "DM Table Cell",
    ParagraphRole.FIGURE_CAPTION: "DM Figure Caption",
    ParagraphRole.TOC_HEADING: "DM TOC Heading",
    ParagraphRole.TOC_ENTRY: "DM TOC Entry",
    ParagraphRole.ABBREVIATION_ENTRY: "DM Abbreviation",
    ParagraphRole.APPENDIX_HEADING: "DM Appendix Heading",
    ParagraphRole.REFERENCES_HEADING: "DM References Heading",
    ParagraphRole.REFERENCES_ENTRY: "DM Reference Entry",
    ParagraphRole.COVER_TITLE: "DM Cover Title",
    ParagraphRole.COVER_FIELD: "DM Cover Field",
}

# Роли, которые ДОЛЖНЫ попадать в оглавление, но не являются Heading-стилями.
# Им вручную проставляется outline level.
OUTLINE_LEVEL_FOR_ROLE: dict[ParagraphRole, int] = {
    ParagraphRole.REFERENCES_HEADING: 0,
    ParagraphRole.APPENDIX_HEADING: 0,
    ParagraphRole.ABSTRACT_HEADING: 0,
}

# ВАЖНО: TOC_HEADING намеренно отсутствует в OUTLINE_LEVEL_FOR_ROLE.
# Если заголовок оглавления получит outline level, оглавление включит само себя.

ALIGNMENT_MAP: dict[Alignment, WD_ALIGN_PARAGRAPH] = {
    Alignment.LEFT: WD_ALIGN_PARAGRAPH.LEFT,
    Alignment.CENTER: WD_ALIGN_PARAGRAPH.CENTER,
    Alignment.RIGHT: WD_ALIGN_PARAGRAPH.RIGHT,
    Alignment.JUSTIFY: WD_ALIGN_PARAGRAPH.JUSTIFY,
}

PAGE_DIMENSIONS = {
    PageSize.A4: (Mm(210), Mm(297)),
    PageSize.LETTER: (Inches(8.5), Inches(11)),
}

NUMBER_FORMAT_CODE = {
    NumberFormat.ARABIC: "decimal",
    NumberFormat.ROMAN_LOWER: "lowerRoman",
    NumberFormat.ROMAN_UPPER: "upperRoman",
}


def style_name_for_role(role: ParagraphRole) -> str:
    """Единственный источник правды о том, какой стиль носит роль.
    Слой сборки документа обращается только сюда."""
    if role in BUILTIN_STYLE_FOR_ROLE:
        return BUILTIN_STYLE_FOR_ROLE[role]
    return CUSTOM_STYLE_FOR_ROLE[role]


# --------------------------------------------------------------------------
# Применение типографики к стилю
# --------------------------------------------------------------------------


def _apply_typography(style: ParagraphStyle, typo: TypographySpec) -> None:
    # Снять тематические шрифты ДО установки своих: при наличии обоих
    # Word отдаёт приоритет *Theme, и Title/Heading остаются Calibri.
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for theme_attr in ("w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"):
        rfonts.attrib.pop(qn(theme_attr), None)

    font = style.font
    font.name = typo.font_family.value
    font.size = Pt(typo.font_size_pt)
    font.bold = typo.bold
    font.italic = typo.italic
    font.underline = typo.underline

    # Ловушка 2: без явного цвета встроенные Heading берут синий из темы.
    font.color.rgb = BLACK

    # UPPER — свойство отображения, текст в файле не меняется.
    # TITLE_CASE и SENTENCE_CASE в Word отсутствуют и применяются к строке
    # на этапе сборки документа.
    font.all_caps = typo.text_case == TextCase.UPPER
    font.small_caps = typo.small_caps

    # Шрифт надо продублировать в восточноазиатский слот, иначе Word
    # подставляет тему для части символов. Повторно снять theme-атрибуты
    # на случай, если font.name их вернул.
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for theme_attr in ("w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"):
        rfonts.attrib.pop(qn(theme_attr), None)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(attr), typo.font_family.value)

    # Strip inherited character metrics from the template (Subtitle letter-spacing, etc.).
    for junk in ("w:spacing", "w:kern", "w:position", "w:w"):
        node = rpr.find(qn(junk))
        if node is not None:
            rpr.remove(node)

    pf = style.paragraph_format
    pf.alignment = ALIGNMENT_MAP[typo.alignment]
    pf.line_spacing = typo.line_spacing
    pf.space_before = Pt(typo.space_before_pt)
    pf.space_after = Pt(typo.space_after_pt)
    pf.right_indent = Inches(typo.right_indent_in)
    pf.keep_with_next = typo.keep_with_next
    pf.page_break_before = typo.page_break_before
    pf.widow_control = typo.widow_control

    # Ловушка 1: висячий отступ — это отрицательная первая строка.
    if typo.hanging_indent_in > 0:
        pf.left_indent = Inches(typo.left_indent_in + typo.hanging_indent_in)
        pf.first_line_indent = Inches(-typo.hanging_indent_in)
    else:
        pf.left_indent = Inches(typo.left_indent_in)
        pf.first_line_indent = Inches(typo.first_line_indent_in)


def _clear_paragraph_borders(style: ParagraphStyle) -> None:
    """Remove ``w:pBdr`` from the style — academic profiles never draw rules."""
    ppr = style.element.pPr
    if ppr is None:
        return
    borders = ppr.find(qn("w:pBdr"))
    if borders is not None:
        ppr.remove(borders)


def _set_outline_level(style: ParagraphStyle, level: int) -> None:
    """Делает абзац видимым для поля оглавления. level 0 = верхний уровень."""
    ppr = style.element.get_or_add_pPr()
    existing = ppr.find(qn("w:outlineLvl"))
    if existing is not None:
        ppr.remove(existing)
    element = OxmlElement("w:outlineLvl")
    element.set(qn("w:val"), str(level))
    ppr.append(element)


def _get_or_create_style(document: DocxDocument, name: str) -> ParagraphStyle:
    try:
        return document.styles[name]
    except KeyError:
        return document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


# --------------------------------------------------------------------------
# Публичное API
# --------------------------------------------------------------------------


def build_styles(document: DocxDocument, spec: FormatSpec) -> None:
    """Создаёт и настраивает стили для всех ролей из спеки.

    Вызывается один раз на документ, до добавления содержимого."""
    for role, typo in spec.roles.items():
        style = _get_or_create_style(document, style_name_for_role(role))

        # Пользовательские стили наследуются от Normal, чтобы не тянуть
        # неожиданные свойства из шаблона.
        if role in CUSTOM_STYLE_FOR_ROLE:
            try:
                style.base_style = document.styles["Normal"]
            except KeyError:
                pass

        _apply_typography(style, typo)
        _clear_paragraph_borders(style)

        if role in OUTLINE_LEVEL_FOR_ROLE:
            _set_outline_level(style, OUTLINE_LEVEL_FOR_ROLE[role])

        # Стили заголовков не должны отрываться от следующего абзаца.
        if role in (
            ParagraphRole.HEADING_1,
            ParagraphRole.HEADING_2,
            ParagraphRole.HEADING_3,
            ParagraphRole.HEADING_4,
            ParagraphRole.REFERENCES_HEADING,
            ParagraphRole.APPENDIX_HEADING,
        ):
            style.paragraph_format.keep_with_next = True


def apply_page_setup(document: DocxDocument, spec: FormatSpec) -> None:
    width, height = PAGE_DIMENSIONS[spec.page.size]
    margins = spec.page.margins

    for section in document.sections:
        section.page_width = width
        section.page_height = height
        section.top_margin = Inches(margins.top_in)
        section.bottom_margin = Inches(margins.bottom_in)
        section.left_margin = Inches(margins.left_in)
        section.right_margin = Inches(margins.right_in)
        if spec.page.binding_offset_in:
            section.gutter = Inches(spec.page.binding_offset_in)


def _add_field(paragraph, instruction: str, placeholder: str = "1") -> None:
    """Вставляет поле Word. API python-docx полей не умеет вовсе."""
    run = paragraph.add_run()

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")

    placeholder_el = OxmlElement("w:t")
    placeholder_el.set(qn("xml:space"), "preserve")
    placeholder_el.text = placeholder

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    for element in (begin, instr, separate, placeholder_el, end):
        run._r.append(element)


def _set_page_number_start(section, start_at: int, number_format: NumberFormat) -> None:
    """Старт и формат нумерации живут в w:pgNumType, минуя API."""
    sect_pr = section._sectPr
    existing = sect_pr.find(qn("w:pgNumType"))
    if existing is not None:
        sect_pr.remove(existing)

    pg_num = OxmlElement("w:pgNumType")
    pg_num.set(qn("w:start"), str(start_at))
    pg_num.set(qn("w:fmt"), NUMBER_FORMAT_CODE[number_format])
    sect_pr.append(pg_num)


def apply_page_numbering(document: DocxDocument, spec: FormatSpec) -> None:
    numbering = spec.page_numbering
    if numbering.position == PageNumberPosition.NONE:
        return

    for section in document.sections:
        apply_page_numbering_to_section(section, numbering)


def apply_page_numbering_to_section(
    section,
    numbering,
    *,
    number_format: NumberFormat | None = None,
    start_at: int | None = None,
    running_head_prefix: str | None = None,
) -> None:
    """Apply PAGE field (and optional surname prefix) to one section."""
    if numbering.position == PageNumberPosition.NONE:
        return

    position = numbering.position.value
    at_top = position.startswith("top")
    if position.endswith("left"):
        alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif position.endswith("center"):
        alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        alignment = WD_ALIGN_PARAGRAPH.RIGHT

    fmt = number_format if number_format is not None else numbering.number_format
    start = start_at if start_at is not None else numbering.start_at

    section.different_first_page_header_footer = numbering.skip_first_page

    container = section.header if at_top else section.footer
    container.is_linked_to_previous = False
    paragraph = container.paragraphs[0] if container.paragraphs else container.add_paragraph()
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    paragraph.alignment = alignment

    if running_head_prefix:
        paragraph.add_run(f"{running_head_prefix} ")
    _add_field(paragraph, "PAGE", placeholder="1")
    if numbering.include_total:
        paragraph.add_run(" of ")
        _add_field(paragraph, "NUMPAGES", placeholder="1")

    _set_page_number_start(section, start, fmt)


TOC_FIELD_PLACEHOLDER = (
    "Right-click here and choose Update Field to build the table of contents."
)


def add_table_of_contents_field(paragraph, max_depth: int) -> None:
    """Word TOC field. Refresh on open (updateFields) or via Update Field."""
    _add_field(
        paragraph,
        rf'TOC \o "1-{max_depth}" \h \z \u',
        placeholder=TOC_FIELD_PLACEHOLDER,
    )


def start_new_section(document: DocxDocument, restart_numbering_at: int | None = None):
    """Новый раздел нужен, когда меняется схема нумерации —
    например римские цифры во front matter и арабские в теле."""
    section = document.add_section(WD_SECTION_START.NEW_PAGE)
    if restart_numbering_at is not None:
        _set_page_number_start(section, restart_numbering_at, NumberFormat.ARABIC)
    return section


def enable_field_update(document: DocxDocument) -> None:
    """Force Word to refresh fields (TOC, PAGE) on open.

    Writes ``<w:updateFields w:val="true"/>`` into ``word/settings.xml``.
    Not exposed by the python-docx API — must use OxmlElement.
    """
    settings = document.settings.element
    existing = settings.find(qn("w:updateFields"))
    if existing is not None:
        settings.remove(existing)
    element = OxmlElement("w:updateFields")
    element.set(qn("w:val"), "true")
    settings.append(element)
