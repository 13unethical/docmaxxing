"""Assemble a python-docx Document from typed content blocks + FormatSpec."""

from __future__ import annotations

from dataclasses import dataclass

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.text import WD_BREAK
from docx.text.paragraph import Paragraph

from formatter_v2.citations.renderer import FormattedText
from formatter_v2.render.rich_text import apply_formatted_text, is_formatted_text
from formatter_v2.render.styles import (
    apply_page_numbering,
    apply_page_setup,
    build_styles,
    style_name_for_role,
)
from formatter_v2.render.text_case import academic_title_case
from formatter_v2.spec import FormatSpec, ParagraphRole, TextCase, TypographySpec


@dataclass(frozen=True)
class Block:
    """One logical paragraph. Newlines are not allowed inside plain ``text`` —
    use separate Block instances instead.

    ``REFERENCES_ENTRY`` may carry ``FormattedText`` (CSL rich bibliography)
    instead of a plain string.
    """

    role: ParagraphRole
    text: str | FormattedText


def _apply_text_case(text: str, text_case: TextCase) -> str:
    """TITLE_CASE / SENTENCE_CASE must be applied to the string — Word has no such properties.
    UPPER is handled via font.all_caps in styles; do not transform the string again.
    """
    if text_case == TextCase.TITLE_CASE:
        return academic_title_case(text)
    if text_case == TextCase.SENTENCE_CASE:
        stripped = text.strip()
        if not stripped:
            return text
        lower = stripped.lower()
        return lower[0].upper() + lower[1:]
    return text


def add_paragraph(
    document: DocxDocument,
    role: ParagraphRole,
    content: str | FormattedText,
    spec: FormatSpec,
) -> Paragraph:
    """Single entry point for writing a styled paragraph.

    Applies the role style, ``academic_title_case`` / sentence case when needed,
    and rich-text runs. Both ``render_document`` and ``build_document`` must
    call only this helper — never ``document.add_paragraph`` directly.
    """
    typo: TypographySpec = spec.typography(role)
    style_name = style_name_for_role(role)

    if is_formatted_text(content):
        paragraph = document.add_paragraph(style=style_name)
        apply_formatted_text(paragraph, content)  # type: ignore[arg-type]
        return paragraph

    plain = content if isinstance(content, str) else ""
    text = _apply_text_case(plain, typo.text_case)
    paragraph = document.add_paragraph(style=style_name)
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)
    return paragraph


def add_page_break(document: DocxDocument) -> None:
    """Insert a page break without going through role styling."""
    paragraph = document.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def _clear_default_empty_paragraph(document: DocxDocument) -> None:
    """Remove the blank paragraph Document() inserts by default."""
    if len(document.paragraphs) == 1 and document.paragraphs[0].text == "":
        element = document.paragraphs[0]._element
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)


def render_document(blocks: list[Block], spec: FormatSpec) -> DocxDocument:
    """Build a DOCX from content blocks using the resolved FormatSpec styles."""
    document = Document()
    apply_page_setup(document, spec)
    build_styles(document, spec)
    apply_page_numbering(document, spec)
    _clear_default_empty_paragraph(document)

    for block in blocks:
        plain = block.text if isinstance(block.text, str) else block.text
        if isinstance(plain, str):
            # Never embed "\n" in a paragraph — split into separate paragraphs.
            lines = plain.replace("\r\n", "\n").replace("\r", "\n").split("\n")
            for line in lines:
                if line == "" and len(lines) > 1:
                    continue
                add_paragraph(document, block.role, line, spec)
        else:
            add_paragraph(document, block.role, plain, spec)

    return document
