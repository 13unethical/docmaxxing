"""Assemble a full academic DOCX from DocumentModel + FormatSpec."""

from __future__ import annotations

import re
import string

from docx import Document
from docx.document import Document as DocxDocument
from docx.shared import Inches
from docx.text.paragraph import Paragraph

from formatter_v2.render.dates import format_cover_date
from formatter_v2.render.document import Block, add_page_break, add_paragraph
from formatter_v2.render.model import DocumentModel
from formatter_v2.render.styles import (
    add_table_of_contents_field,
    apply_page_numbering_to_section,
    apply_page_setup,
    build_styles,
    enable_field_update,
    start_new_section,
)
from formatter_v2.spec import (
    CoverPage,
    FormatSpec,
    NumberFormat,
    ParagraphRole,
    StyleName,
)

_HEADING_DEPTH: dict[ParagraphRole, int] = {
    ParagraphRole.HEADING_1: 1,
    ParagraphRole.HEADING_2: 2,
    ParagraphRole.HEADING_3: 3,
    ParagraphRole.HEADING_4: 4,
}

_TOC_LEVEL_INDENT_IN = 0.25


def _clear_default_empty_paragraph(document: DocxDocument) -> None:
    if len(document.paragraphs) == 1 and document.paragraphs[0].text == "":
        element = document.paragraphs[0]._element
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)


def _write_block(document: DocxDocument, block: Block, spec: FormatSpec) -> Paragraph | None:
    """Write a block via the shared ``add_paragraph`` helper. Returns last para."""
    last: Paragraph | None = None
    if isinstance(block.text, str):
        lines = block.text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
        for line in lines:
            if line == "" and len(lines) > 1:
                continue
            last = add_paragraph(document, block.role, line, spec)
        return last
    return add_paragraph(document, block.role, block.text, spec)


def _strip_existing_caption_prefix(text: str, label: str) -> str:
    """Remove a leading 'Table 1. ' / 'Fig. 2: ' style prefix if present."""
    pattern = re.compile(
        rf"^\s*{re.escape(label)}\s*\d+\s*[.:\-–—]?\s*",
        re.IGNORECASE,
    )
    return pattern.sub("", text).strip() or text.strip()


def format_caption_text(
    *,
    label: str,
    number: int,
    separator: str,
    description: str,
) -> str:
    """``{label} {number}{separator}{text}`` per CaptionConfig."""
    desc = _strip_existing_caption_prefix(description, label)
    return f"{label} {number}{separator}{desc}"


class _CaptionCounters:
    def __init__(self) -> None:
        self.table = 0
        self.figure = 0

    def next_table(self) -> int:
        self.table += 1
        return self.table

    def next_figure(self) -> int:
        self.figure += 1
        return self.figure


def _caption_block(
    role: ParagraphRole,
    description: str,
    spec: FormatSpec,
    counters: _CaptionCounters,
) -> Block:
    cfg = spec.captions
    if role == ParagraphRole.TABLE_CAPTION:
        number = counters.next_table() if cfg.auto_number else 1
        label = cfg.table_label
    else:
        number = counters.next_figure() if cfg.auto_number else 1
        label = cfg.figure_label
    text = format_caption_text(
        label=label,
        number=number,
        separator=cfg.separator,
        description=description if isinstance(description, str) else "",
    )
    return Block(role, text)


def _surname_from_cover(cover: CoverPage | None) -> str:
    if cover is None or not cover.student_name.strip():
        return "Author"
    parts = cover.student_name.strip().split()
    return parts[-1]


def _flush_left_no_indent(paragraph: Paragraph) -> None:
    """MLA identity lines: no first-line indent (BODY would indent 0.5\")."""
    paragraph.paragraph_format.first_line_indent = Inches(0)
    paragraph.paragraph_format.left_indent = Inches(0)


def _mla_header_blocks(cover: CoverPage, date_format: str) -> list[Block]:
    """Four flush-left identity lines MLA requires instead of a title page."""
    date_text = ""
    if cover.submission_date is not None:
        date_text = format_cover_date(cover.submission_date, date_format)  # type: ignore[arg-type]
    course = cover.course or cover.module_title or cover.module_code
    lines = [
        cover.student_name or "",
        cover.lecturer or "",
        course or "",
        date_text,
    ]
    return [Block(ParagraphRole.BODY_FIRST, line) for line in lines]


def _should_render_cover(spec: FormatSpec, model: DocumentModel) -> bool:
    if spec.style == StyleName.MLA9:
        return False
    if model.cover is None:
        return False
    return bool(spec.cover_page.enabled and model.cover.enabled)


def _render_cover(document: DocxDocument, cover: CoverPage, spec: FormatSpec) -> None:
    for _ in range(max(0, cover.top_spacer_lines)):
        add_paragraph(document, ParagraphRole.BODY, "", spec)

    if cover.title:
        _write_block(document, Block(ParagraphRole.COVER_TITLE, cover.title), spec)
    if cover.subtitle:
        _write_block(document, Block(ParagraphRole.SUBTITLE, cover.subtitle), spec)

    field_rows: list[str] = []
    for value in (
        cover.student_name,
        cover.student_id,
        cover.university,
        cover.faculty,
        cover.course,
        cover.module_code,
        cover.module_title,
        cover.lecturer,
    ):
        if value and value.strip():
            field_rows.append(value.strip())
    if cover.submission_date is not None:
        field_rows.append(format_cover_date(cover.submission_date, spec.date_format))
    if cover.word_count is not None:
        field_rows.append(f"Word count: {cover.word_count}")
    for extra in cover.extra_fields:
        if extra.show_label and extra.label:
            field_rows.append(f"{extra.label}: {extra.value}".strip())
        elif extra.value:
            field_rows.append(extra.value)

    for row in field_rows:
        _write_block(document, Block(ParagraphRole.COVER_FIELD, row), spec)

    if cover.page_break_after:
        add_page_break(document)


def _heading_plain_text(block: Block) -> str:
    if isinstance(block.text, str):
        return block.text.strip()
    # FormattedText bibliography-style — unlikely for headings; stringify.
    return str(block.text).strip()


def _render_static_toc(
    document: DocxDocument,
    spec: FormatSpec,
    body_blocks: list[Block],
) -> None:
    toc = spec.table_of_contents
    for block in body_blocks:
        depth = _HEADING_DEPTH.get(block.role)
        if depth is None or depth > toc.max_depth:
            continue
        text = _heading_plain_text(block)
        if not text:
            continue
        para = add_paragraph(document, ParagraphRole.TOC_ENTRY, text, spec)
        para.paragraph_format.first_line_indent = Inches(0)
        para.paragraph_format.left_indent = Inches(_TOC_LEVEL_INDENT_IN * (depth - 1))


def _render_toc(
    document: DocxDocument,
    spec: FormatSpec,
    body_blocks: list[Block],
) -> None:
    toc = spec.table_of_contents
    _write_block(
        document,
        Block(ParagraphRole.TOC_HEADING, toc.heading_text),
        spec,
    )
    if toc.field_based:
        field_para = add_paragraph(document, ParagraphRole.TOC_ENTRY, "", spec)
        for run in list(field_para.runs):
            run._element.getparent().remove(run._element)
        add_table_of_contents_field(field_para, toc.max_depth)
    else:
        _render_static_toc(document, spec, body_blocks)
    if toc.page_break_after:
        add_page_break(document)


def _render_abbreviations(document: DocxDocument, spec: FormatSpec) -> None:
    abbr = spec.abbreviations
    _write_block(
        document,
        Block(ParagraphRole.HEADING_1, abbr.heading_text),
        spec,
    )
    items = list(abbr.entries.items())
    if abbr.sort_alphabetically:
        items.sort(key=lambda kv: kv[0].casefold())
    for key, meaning in items:
        _write_block(
            document,
            Block(ParagraphRole.ABBREVIATION_ENTRY, f"{key} — {meaning}"),
            spec,
        )
    if abbr.page_break_after:
        add_page_break(document)


def _write_two_line_table_caption(
    document: DocxDocument,
    *,
    label: str,
    number: int,
    description: str,
    spec: FormatSpec,
) -> None:
    """APA-style caption: bold label line, italic title line."""
    label_para = add_paragraph(
        document, ParagraphRole.TABLE_CAPTION, f"{label} {number}", spec
    )
    for run in label_para.runs:
        run.bold = True

    title_para = add_paragraph(document, ParagraphRole.TABLE_CAPTION, description, spec)
    for run in title_para.runs:
        run.italic = True


def _render_body_blocks(
    document: DocxDocument,
    blocks: list[Block],
    spec: FormatSpec,
    counters: _CaptionCounters,
) -> None:
    """Render body blocks; number table/figure captions via CaptionConfig.

    Caption *position* (above/below) is the caller's responsibility when
    assembling ``DocumentModel.body`` — the builder does not reorder blocks.
    """
    cfg = spec.captions
    for block in blocks:
        if (
            cfg.enabled
            and block.role == ParagraphRole.TABLE_CAPTION
            and isinstance(block.text, str)
        ):
            number = counters.next_table() if cfg.auto_number else 1
            description = _strip_existing_caption_prefix(block.text, cfg.table_label)
            if cfg.two_line:
                _write_two_line_table_caption(
                    document,
                    label=cfg.table_label,
                    number=number,
                    description=description,
                    spec=spec,
                )
                continue
            numbered = Block(
                ParagraphRole.TABLE_CAPTION,
                format_caption_text(
                    label=cfg.table_label,
                    number=number,
                    separator=cfg.separator,
                    description=description,
                ),
            )
            _write_block(document, numbered, spec)
            continue

        if (
            cfg.enabled
            and block.role == ParagraphRole.FIGURE_CAPTION
            and isinstance(block.text, str)
        ):
            block = _caption_block(block.role, block.text, spec, counters)
        _write_block(document, block, spec)


def _split_appendices(blocks: list[Block]) -> list[list[Block]]:
    groups: list[list[Block]] = []
    current: list[Block] = []
    for block in blocks:
        if block.role == ParagraphRole.APPENDIX_HEADING:
            if current:
                groups.append(current)
            current = [block]
        else:
            if not current:
                current = [Block(ParagraphRole.APPENDIX_HEADING, "")]
            current.append(block)
    if current:
        groups.append(current)
    return groups


def _appendix_title(spec: FormatSpec, index: int) -> str:
    prefix = spec.appendices.heading_prefix or "Appendix"
    if spec.appendices.lettered:
        letter = string.ascii_uppercase[index % 26]
        return f"{prefix} {letter}"
    return f"{prefix} {index + 1}"


def _render_appendices(
    document: DocxDocument,
    blocks: list[Block],
    spec: FormatSpec,
) -> None:
    if not blocks:
        return
    groups = _split_appendices(blocks)
    cfg = spec.appendices
    for index, group in enumerate(groups):
        if cfg.enabled and cfg.page_break_before_each:
            add_page_break(document)
        if cfg.enabled:
            title = _appendix_title(spec, index)
            rest = (
                group[1:]
                if group and group[0].role == ParagraphRole.APPENDIX_HEADING
                else group
            )
        else:
            if group and group[0].role == ParagraphRole.APPENDIX_HEADING:
                head = group[0].text if isinstance(group[0].text, str) else ""
                title = head or "Appendix"
                rest = group[1:]
            else:
                title = "Appendix"
                rest = group
        _write_block(document, Block(ParagraphRole.APPENDIX_HEADING, title), spec)
        for block in rest:
            _write_block(document, block, spec)


def _render_front_matter_bundle(
    document: DocxDocument,
    model: DocumentModel,
    spec: FormatSpec,
    *,
    is_mla: bool,
) -> None:
    """Cover (non-MLA) + MLA header/title + abstract/keywords + TOC + abbreviations."""
    if _should_render_cover(spec, model):
        assert model.cover is not None
        _render_cover(document, model.cover, spec)

    if is_mla and model.cover is not None:
        for block in _mla_header_blocks(model.cover, spec.date_format):
            para = _write_block(document, block, spec)
            if para is not None:
                _flush_left_no_indent(para)
        title = (model.cover.title or "").strip()
        if title:
            _write_block(document, Block(ParagraphRole.DOC_TITLE, title), spec)

    for block in model.front_matter:
        _write_block(document, block, spec)

    if spec.table_of_contents.enabled:
        _render_toc(document, spec, model.body)

    if spec.abbreviations.enabled and spec.abbreviations.entries:
        _render_abbreviations(document, spec)


def _document_has_fields(spec: FormatSpec) -> bool:
    if spec.table_of_contents.enabled and spec.table_of_contents.field_based:
        return True
    if spec.page_numbering.position.value != "none":
        return True
    return False


def build_document(model: DocumentModel, spec: FormatSpec) -> DocxDocument:
    """Build a full DOCX in fixed section order.

    Order: cover → abstract/keywords → TOC → abbreviations → body →
    references → appendices. Disabled ``FormatSpec`` parts are skipped.
    """
    document = Document()
    apply_page_setup(document, spec)
    build_styles(document, spec)
    _clear_default_empty_paragraph(document)

    is_mla = spec.style == StyleName.MLA9
    counters = _CaptionCounters()
    numbering = spec.page_numbering
    restart = numbering.restart_after_front_matter and not is_mla
    mla_surname = _surname_from_cover(model.cover) if is_mla else None

    if restart:
        # Section 0 — front matter with roman numerals.
        apply_page_numbering_to_section(
            document.sections[0],
            numbering,
            number_format=NumberFormat.ROMAN_LOWER,
            start_at=1,
        )
        _render_front_matter_bundle(document, model, spec, is_mla=False)

        # Section 1 — body onwards, arabic from 1.
        body_section = start_new_section(document, restart_numbering_at=1)
        apply_page_numbering_to_section(
            body_section,
            numbering,
            number_format=NumberFormat.ARABIC,
            start_at=1,
        )
        body_section.different_first_page_header_footer = False

        _render_body_blocks(document, model.body, spec, counters)
        if spec.references.enabled:
            for block in model.references:
                _write_block(document, block, spec)
        _render_appendices(document, model.appendices, spec)
    else:
        apply_page_numbering_to_section(
            document.sections[0],
            numbering,
            running_head_prefix=mla_surname if is_mla else None,
        )
        _render_front_matter_bundle(document, model, spec, is_mla=is_mla)
        _render_body_blocks(document, model.body, spec, counters)
        if spec.references.enabled:
            for block in model.references:
                _write_block(document, block, spec)
        _render_appendices(document, model.appendices, spec)

    if _document_has_fields(spec):
        enable_field_update(document)

    return document
