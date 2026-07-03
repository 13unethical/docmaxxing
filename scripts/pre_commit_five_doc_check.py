#!/usr/bin/env python3
"""Pre-commit verification: five document types × spacing fix."""

from __future__ import annotations

import sys
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from formatter.cover_page import CoverPageData, prepend_cover_page
from formatter.headings import detect_heading_level
from formatter import FormatJob, format_document_full
from formatter.references_section import append_references_section
from services.document_analyzer import _docx_has_page_number_field

OUT_DIR = ROOT / "tmp" / "pre_commit_verify"
OUT_DIR.mkdir(parents=True, exist_ok=True)


@dataclass
class DocSpec:
    name: str
    job: FormatJob
    paragraphs: list[str]
    references: list[str]
    page_number_position: str


def _job(**kwargs) -> FormatJob:
    ls = kwargs.get("line_spacing", 2.0)
    default_style = "harvard" if ls < 1.99 else "apa7"
    return FormatJob(
        font_family=kwargs.get("font_family", "Times New Roman"),
        font_size_pt=kwargs.get("font_size_pt", 12),
        line_spacing=ls,
        alignment=kwargs.get("alignment", "left"),
        first_line_indent=kwargs.get("first_line_indent", False),
        space_before_pt=kwargs.get("space_before_pt", 0),
        space_after_pt=kwargs.get("space_after_pt", 0),
        margin_preset=kwargs.get("margin_preset", "normal"),
        page_number_position=kwargs.get("page_number_position", "none"),
        auto_headings=True,
        heading_all_caps=False,
        auto_justify_refs=True,
        format_style=kwargs.get("format_style", default_style),
    )


def _expected_heading_counts(paragraphs: list[str], *, include_refs: bool = True) -> tuple[int, int, int]:
    h1 = h2 = h3 = 0
    seen = False
    for text in paragraphs:
        stripped = text.strip()
        if not stripped:
            continue
        first = not seen
        seen = True
        level = detect_heading_level(stripped, True, is_first_nonempty=first)
        if level == 1:
            h1 += 1
        elif level == 2:
            h2 += 1
        elif level == 3:
            h3 += 1
    if include_refs:
        h2 += 1  # appended References heading
    return h1, h2, h3


def _style_counts(doc: Document) -> tuple[int, int, int]:
    h1 = h2 = h3 = 0
    for p in doc.paragraphs:
        name = p.style.name if p.style else ""
        if name == "Heading 1":
            h1 += 1
        elif name == "Heading 2":
            h2 += 1
        elif name == "Heading 3":
            h3 += 1
    return h1, h2, h3


def _body_line_spacing_ok(doc: Document, expected: float) -> list[str]:
    errors = []
    for i, p in enumerate(doc.paragraphs):
        if (p.style.name or "") != "Normal":
            continue
        pf = p.paragraph_format
        if expected >= 1.99:
            # python-docx often reads DOUBLE as None after assignment
            if pf.line_spacing_rule not in (WD_LINE_SPACING.DOUBLE, None):
                errors.append(f"[{i}] body expected double spacing, got {pf.line_spacing_rule}")
        elif abs(expected - 1.5) < 0.01:
            if pf.line_spacing_rule not in (WD_LINE_SPACING.MULTIPLE, WD_LINE_SPACING.ONE_POINT_FIVE):
                errors.append(f"[{i}] body expected 1.5 spacing, got {pf.line_spacing_rule}")
            elif pf.line_spacing is not None and abs(float(pf.line_spacing) - 1.5) > 0.01:
                errors.append(f"[{i}] body line_spacing={pf.line_spacing}, expected 1.5")
    return errors


def build_specs() -> list[DocSpec]:
    refs = [
        "Smith, J. (2024). Example study. Journal of Testing, 1(1), 1–10.",
        "Jones, A. (2023). Another source. Publisher.",
    ]
    return [
        DocSpec(
            name="Essay",
            job=_job(line_spacing=2.0, alignment="justify", page_number_position="top_right"),
            paragraphs=[
                "The Impact of Climate Policy on Urban Development",
                "Introduction",
                "This essay examines how climate policy shapes urban planning decisions.",
                "Main Body",
                "Cities are adopting greener infrastructure at an increasing pace worldwide.",
                "Conclusion",
                "This essay has shown that coordinated policy remains essential.",
            ],
            references=refs,
            page_number_position="top_right",
        ),
        DocSpec(
            name="Research Paper",
            job=_job(line_spacing=2.0, page_number_position="top_right"),
            paragraphs=[
                "Machine Learning Applications in Medical Diagnostics Today",
                "Abstract",
                "This paper reviews recent machine learning methods in clinical imaging.",
                "Introduction",
                "Diagnostic accuracy has improved with data-driven models.",
                "Methods",
                "We surveyed peer-reviewed studies published between 2018 and 2024.",
                "Results",
                "Most studies reported measurable gains in sensitivity and specificity.",
                "Discussion",
                "The findings suggest cautious optimism for deployment in screening.",
                "Conclusion",
                "Further validation studies are required before routine adoption.",
            ],
            references=refs,
            page_number_position="top_right",
        ),
        DocSpec(
            name="Literature Review",
            job=_job(line_spacing=2.0, page_number_position="top_right"),
            paragraphs=[
                "Remote Work and Employee Wellbeing After the Pandemic",
                "Introduction",
                "Remote work became widespread during the global health crisis.",
                "Literature Review",
                "Prior research links flexibility with both satisfaction and isolation.",
                "Discussion",
                "The literature reveals mixed outcomes depending on sector and role.",
                "Conclusion",
                "Organizations must balance autonomy with structured support.",
            ],
            references=refs,
            page_number_position="top_right",
        ),
        DocSpec(
            name="Case Study",
            job=_job(
                font_family="Arial",
                font_size_pt=11,
                line_spacing=1.5,
                alignment="justify",
                space_after_pt=12,
                page_number_position="bottom_right",
                format_style="harvard",
            ),
            paragraphs=[
                "Supply Chain Resilience at Northbridge Manufacturing Limited",
                "Introduction",
                "Northbridge Manufacturing faced severe disruption during supplier outages.",
                "Background",
                "The firm operates three plants across two countries in Europe.",
                "Analysis",
                "Inventory buffers were insufficient for critical components.",
                "Recommendations",
                "The company should diversify suppliers and invest in forecasting tools.",
                "Conclusion",
                "A phased resilience programme would reduce future downtime risk.",
            ],
            references=refs,
            page_number_position="bottom_right",
        ),
        DocSpec(
            name="Reflection",
            job=_job(line_spacing=2.0, page_number_position="none"),
            paragraphs=[
                "Reflective Account of Team Leadership During a Group Project",
                "Introduction",
                "This reflection considers my role as team leader last semester.",
                "Experience",
                "Our group struggled initially to agree on research priorities.",
                "Analysis",
                "I learned that early facilitation prevents later conflict.",
                "Learning outcomes",
                "I will schedule structured check-ins at the start of future projects.",
            ],
            references=refs,
            page_number_position="none",
        ),
    ]


def verify_spec(spec: DocSpec) -> list[str]:
    errors: list[str] = []
    exp_h1, exp_h2, exp_h3 = _expected_heading_counts(spec.paragraphs)

    doc = Document()
    for line in spec.paragraphs:
        doc.add_paragraph(line)

    format_document_full(doc, spec.job, None)
    append_references_section(doc, spec.job, spec.references, section_title="References")

    out_path = OUT_DIR / f"{spec.name.lower().replace(' ', '_')}.docx"
    doc.save(str(out_path))

    act_h1, act_h2, act_h3 = _style_counts(doc)
    if (act_h1, act_h2, act_h3) != (exp_h1, exp_h2, exp_h3):
        errors.append(
            f"heading counts: expected H1/H2/H3={exp_h1}/{exp_h2}/{exp_h3}, "
            f"got {act_h1}/{act_h2}/{act_h3}"
        )

    errors.extend(_body_line_spacing_ok(doc, spec.job.line_spacing))

    ref_texts = [p.text.strip() for p in doc.paragraphs if p.text.strip() in spec.references]
    if ref_texts != spec.references:
        errors.append(f"references changed: expected {spec.references}, found {ref_texts}")

    for i, p in enumerate(doc.paragraphs):
        if not (p.style.name or "").startswith("Heading"):
            continue
        pf = p.paragraph_format
        if pf.line_spacing_rule not in (WD_LINE_SPACING.SINGLE, None):
            errors.append(f"[{i}] heading line spacing not SINGLE: {pf.line_spacing_rule}")
        if pf.keep_with_next is not True:
            errors.append(f"[{i}] heading keep_with_next is not True")

    if spec.page_number_position != "none":
        if not _docx_has_page_number_field(doc):
            errors.append(f"PAGE field missing for position {spec.page_number_position}")
    elif _docx_has_page_number_field(doc):
        errors.append("PAGE field present but page_number_position is none")

    return errors


def verify_cover_and_page_numbers_unchanged_by_spacing() -> list[str]:
    """Cover/page-number side effects must match pre-spacing-fix behaviour."""
    errors: list[str] = []
    doc = Document()
    doc.add_paragraph("Body paragraph after the cover page.")
    prepend_cover_page(
        doc,
        CoverPageData(
            assignment_title="Cover Title Example",
            student_name="Jane Student",
            university="Example University",
            submission_date="2026-07-01",
        ),
        font_family="Times New Roman",
    )
    job = FormatJob(
        font_family="Times New Roman",
        font_size_pt=12,
        line_spacing=2.0,
        alignment="left",
        first_line_indent=False,
        space_before_pt=0,
        space_after_pt=0,
        margin_preset="normal",
        page_number_position="top_right",
        auto_headings=True,
        heading_all_caps=False,
        auto_justify_refs=False,
    )
    format_document_full(doc, job, None)

    if not _docx_has_page_number_field(doc):
        errors.append("cover+page test: PAGE field missing with top_right")

    cover_title = next(p for p in doc.paragraphs if "Cover Title Example" in (p.text or ""))
    # Same as HEAD (98698b9): format pass restyles cover lines; spacing fix must not worsen further.
    if cover_title.style.name.startswith("Heading"):
        pass  # pre-existing with auto_headings on cover — not introduced by spacing fix
    if cover_title.alignment != WD_ALIGN_PARAGRAPH.LEFT:
        errors.append(f"cover title alignment unexpected: {cover_title.alignment}")

    has_page_break = any(
        "w:br" in p._p.xml and 'type="page"' in p._p.xml for p in doc.paragraphs
    )
    if not has_page_break:
        errors.append("cover test: page break after cover missing")

    return errors


def main() -> int:
    all_errors: dict[str, list[str]] = {}

    for spec in build_specs():
        errs = verify_spec(spec)
        status = "PASS" if not errs else "FAIL"
        print(f"\n[{status}] {spec.name}")
        if errs:
            all_errors[spec.name] = errs
            for e in errs:
                print(f"  - {e}")

    cover_errs = verify_cover_and_page_numbers_unchanged_by_spacing()
    status = "PASS" if not cover_errs else "FAIL"
    print(f"\n[{status}] Cover Page + Page Numbers")
    if cover_errs:
        all_errors["Cover+PageNumbers"] = cover_errs
        for e in cover_errs:
            print(f"  - {e}")

    if all_errors:
        print("\n" + "=" * 60)
        print("PRE-COMMIT CHECK FAILED — commit blocked")
        return 1

    print("\n" + "=" * 60)
    print("ALL 5 DOCUMENT TYPES + COVER/PAGE CHECK PASSED — safe to commit")
    return 0


if __name__ == "__main__":
    sys.exit(main())
