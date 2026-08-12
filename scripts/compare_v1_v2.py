#!/usr/bin/env python3
"""Side-by-side V1 / V2 formatting for manual comparison.

Usage:
  PYTHONPATH=. python3 scripts/compare_v1_v2.py path/to/input.docx [--style apa7]

Writes build/compare/<stem>_v1.docx and <stem>_v2.docx. Not a test.
"""

from __future__ import annotations

import argparse
import io
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from docx import Document  # noqa: E402

from formatter import FormatJob, format_document_full  # noqa: E402
from formatter_v2.pipeline import format_document_v2  # noqa: E402
from formatter_v2.spec import UserOverrides  # noqa: E402


def main() -> int:
    parser = argparse.ArgumentParser(description="Compare Formatter V1 vs V2 output")
    parser.add_argument("input", type=Path, help="Input .docx (or plain .txt)")
    parser.add_argument(
        "--style",
        default="harvard",
        help="Style id for both pipelines (default: harvard)",
    )
    parser.add_argument(
        "--out",
        type=Path,
        default=ROOT / "build" / "compare",
        help="Output directory (default: build/compare)",
    )
    args = parser.parse_args()

    if not args.input.exists():
        print(f"input not found: {args.input}", file=sys.stderr)
        return 1

    args.out.mkdir(parents=True, exist_ok=True)
    stem = args.input.stem

    # --- V1 ---
    if args.input.suffix.lower() == ".docx":
        v1_doc = Document(str(args.input))
    else:
        v1_doc = Document()
        for line in args.input.read_text(encoding="utf-8").splitlines():
            v1_doc.add_paragraph(line)

    job = FormatJob(
        font_family="Times New Roman",
        font_size_pt=12,
        line_spacing=2.0 if args.style.lower().startswith("apa") else 1.5,
        alignment="left",
        first_line_indent=True,
        space_before_pt=0,
        space_after_pt=0,
        margin_preset="normal",
        page_number_position="top_right",
        auto_headings=True,
        heading_all_caps=False,
        auto_justify_refs=False,
        format_style=str(args.style),
    )
    format_document_full(v1_doc, job, None)
    v1_path = args.out / f"{stem}_v1.docx"
    v1_doc.save(v1_path)
    print(f"wrote {v1_path}")

    # --- V2 ---
    source: object
    if args.input.suffix.lower() == ".docx":
        source = Document(str(args.input))
    else:
        source = args.input.read_text(encoding="utf-8").splitlines()

    result = format_document_v2(source, UserOverrides(), args.style)
    v2_path = args.out / f"{stem}_v2.docx"
    v2_path.write_bytes(result.docx_bytes)
    print(f"wrote {v2_path}")
    print(f"extractor={result.extractor_name} notices={len(result.notices)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
