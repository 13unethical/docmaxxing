"""
Rules-based academic document analysis: structure from text; formatting from .docx when provided.

Does not rewrite content. Only reports checks we can ground in the input (no invented problems).
"""

from __future__ import annotations

import re
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from formatter.headings import (
    COMMON_HEADINGS,
    REFS_HEADINGS,
    normalize_paragraph_text,
)


def _paragraphs_from_text(text: str) -> list[str]:
    blocks = re.split(r"\n\s*\n", (text or "").strip())
    return [b.strip() for b in blocks if b.strip()]


def _has_section_heading(paragraphs: list[str], labels: set[str]) -> bool:
    for p in paragraphs:
        if normalize_paragraph_text(p) in labels:
            return True
    return False


def _approx_line_spacing_multiple(pf) -> float | None:
    """Best-effort numeric line spacing for comparison (None if unknown)."""
    try:
        rule = pf.line_spacing_rule
    except (AttributeError, TypeError):
        return None
    if rule == WD_LINE_SPACING.DOUBLE:
        return 2.0
    if rule == WD_LINE_SPACING.SINGLE:
        return 1.0
    if rule == WD_LINE_SPACING.MULTIPLE and pf.line_spacing is not None:
        try:
            return float(pf.line_spacing)
        except (TypeError, ValueError):
            pass
    if pf.line_spacing is not None:
        try:
            return float(pf.line_spacing)
        except (TypeError, ValueError):
            pass
    return None


def _alignment_label(pf) -> str | None:
    a = pf.alignment
    if a == WD_ALIGN_PARAGRAPH.JUSTIFY:
        return "justify"
    if a == WD_ALIGN_PARAGRAPH.LEFT:
        return "left"
    return None


def _collect_explicit_fonts(doc: Document) -> tuple[set[int], set[str]]:
    """Font sizes (pt) and names explicitly set on runs with text."""
    sizes: set[int] = set()
    names: set[str] = set()
    for p in doc.paragraphs:
        for r in p.runs:
            if not (r.text or "").strip():
                continue
            if r.font.size:
                try:
                    sizes.add(int(round(r.font.size.pt)))
                except (AttributeError, TypeError, ValueError):
                    pass
            if r.font.name:
                names.add(r.font.name.strip())
    return sizes, names


def _docx_has_page_number_field(doc: Document) -> bool:
    """True if any header/footer XML suggests a PAGE field."""
    for section in doc.sections:
        for part in (section.header, section.footer):
            try:
                xml = part._element.xml
            except Exception:
                continue
            if "PAGE" in xml and ("fldChar" in xml or "fldSimple" in xml):
                return True
    return False


def _references_body_looks_sparse(paragraphs: list[str]) -> bool:
    """After a refs heading, flag if there is no plausible citation line."""
    idx = None
    for i, p in enumerate(paragraphs):
        if normalize_paragraph_text(p) in REFS_HEADINGS:
            idx = i
            break
    if idx is None:
        return False
    tail = paragraphs[idx + 1 : idx + 4]
    if not tail:
        return True
    yearish = re.compile(r"\(?(19|20)\d{2}[a-z]?\)?|n\.d\.")
    for line in tail:
        if yearish.search(line):
            return False
    return True


def normalize_expected(raw: dict[str, Any] | None) -> dict[str, Any]:
    """Defaults for strict checks when keys are omitted."""
    e = dict(raw or {})
    out: dict[str, Any] = {
        "font_family": e.get("font_family"),
        "font_size": e.get("font_size"),
        "line_spacing": e.get("line_spacing"),
        "alignment": (e.get("alignment") or "").lower() or None,
        "first_line_indent": e.get("first_line_indent"),
        "require_page_numbers": bool(e.get("require_page_numbers")),
        "check_intro_conclusion": e.get("check_intro_conclusion", True),
        "expect_references_section": e.get("expect_references_section", True),
        "page_number_position": e.get("page_number_position") or "top_right",
    }
    # Coerce line_spacing to float if string
    ls = out["line_spacing"]
    if isinstance(ls, str):
        try:
            out["line_spacing"] = float(ls)
        except ValueError:
            out["line_spacing"] = None
    elif isinstance(ls, int):
        out["line_spacing"] = float(ls)
    return out


def analyze_document(
    *,
    text: str,
    doc: Document | None,
    expected: dict[str, Any] | None = None,
) -> dict[str, Any]:
    """
    Return { "issues": [...], "summary": {...}, "notes": [...] }.

    When `doc` is None, only structural checks run; formatting checks are skipped
    (not reported as failures — we cannot verify fonts from plain text).
    """
    exp = normalize_expected(expected)
    issues: list[dict[str, str]] = []
    notes: list[str] = []

    paras_text = _paragraphs_from_text(text)
    if not paras_text and doc:
        paras_text = [p.text.strip() for p in doc.paragraphs if (p.text or "").strip()]

    if not paras_text:
        issues.append(
            {
                "type": "empty_document",
                "message": "No readable paragraphs were found to analyze.",
                "severity": "high",
            }
        )
        return _finalize(issues, notes)

    # --- Structure (text) ---
    if exp.get("expect_references_section", True) and not _has_section_heading(paras_text, REFS_HEADINGS):
        issues.append(
            {
                "type": "missing_references",
                "message": "No References / Bibliography / Works Cited section heading was detected.",
                "severity": "high",
            }
        )
    elif exp.get("expect_references_section", True) and _references_body_looks_sparse(paras_text):
        issues.append(
            {
                "type": "references_incomplete",
                "message": "A references heading exists, but following paragraphs do not look like dated citations.",
                "severity": "medium",
            }
        )

    if exp.get("check_intro_conclusion", True):
        intro_labels = frozenset({"introduction", "intro"})
        concl_labels = frozenset({"conclusion"})
        if not _has_section_heading(paras_text, intro_labels):
            issues.append(
                {
                    "type": "missing_introduction",
                    "message": "No clear Introduction heading or section title was found.",
                    "severity": "medium",
                }
            )
        if not _has_section_heading(paras_text, concl_labels):
            issues.append(
                {
                    "type": "missing_conclusion",
                    "message": "No Conclusion section heading was found.",
                    "severity": "medium",
                }
            )

    if len(paras_text) < 3:
        issues.append(
            {
                "type": "thin_structure",
                "message": "Very few paragraphs; paper may lack typical section structure.",
                "severity": "low",
            }
        )

    # --- Formatting (.docx only) ---
    if doc is None:
        notes.append(
            "Upload a .docx to verify fonts, paragraph spacing, alignment, line spacing, and page numbers."
        )
        return _finalize(issues, notes)

    exp_font = exp.get("font_family")
    if isinstance(exp_font, str) and exp_font.strip():
        names = _collect_explicit_fonts(doc)[1]
        if names:
            bad = {n for n in names if exp_font.lower() not in n.lower()}
            if bad:
                issues.append(
                    {
                        "type": "font_mismatch",
                        "message": f"Run-level font includes {sorted(bad)[:3]}; expected {exp_font} where set.",
                        "severity": "medium",
                    }
                )

    exp_size = exp.get("font_size")
    if exp_size is not None:
        try:
            want = int(exp_size)
        except (TypeError, ValueError):
            want = None
        if want is not None:
            sizes = _collect_explicit_fonts(doc)[0]
            if sizes and all(abs(s - want) > 0 for s in sizes):
                issues.append(
                    {
                        "type": "font_size_mismatch",
                        "message": f"Explicit run font sizes {sorted(sizes)} do not match expected {want} pt.",
                        "severity": "medium",
                    }
                )

    exp_ls = exp.get("line_spacing")
    if isinstance(exp_ls, (int, float)) and exp_ls > 0:
        mismatched = 0
        checked = 0
        for p in doc.paragraphs:
            if not (p.text or "").strip():
                continue
            got = _approx_line_spacing_multiple(p.paragraph_format)
            if got is None:
                continue
            checked += 1
            if abs(got - float(exp_ls)) > 0.09 and not (
                exp_ls >= 1.9 and got >= 1.9
            ) and not (exp_ls <= 1.6 and got <= 1.6 and exp_ls >= 1.4 and got >= 1.4):
                mismatched += 1
        if checked and mismatched > checked * 0.25:
            issues.append(
                {
                    "type": "spacing_error",
                    "message": f"Many paragraphs do not appear to use line spacing {exp_ls} (sampled from document).",
                    "severity": "medium",
                }
            )

    exp_align = exp.get("alignment")
    if exp_align == "justify":
        off = 0
        n = 0
        for p in doc.paragraphs:
            if not (p.text or "").strip():
                continue
            n += 1
            if _alignment_label(p.paragraph_format) != "justify":
                off += 1
        if n and off > n * 0.25:
            issues.append(
                {
                    "type": "alignment_error",
                    "message": "Many body paragraphs are not fully justified.",
                    "severity": "medium",
                }
            )

    exp_indent = exp.get("first_line_indent")
    if exp_indent is True:
        missing = 0
        checked = 0
        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if not t or len(t) < 80:
                continue
            if normalize_paragraph_text(t) in COMMON_HEADINGS | REFS_HEADINGS:
                continue
            checked += 1
            fi = p.paragraph_format.first_line_indent
            try:
                fi_pt = float(fi.pt) if fi is not None else 0.0
            except (AttributeError, TypeError):
                fi_pt = 0.0
            # Typical first-line indent 0.5" ≈ 36 pt; flag if essentially absent
            if fi is None or fi_pt < 12:
                missing += 1
        if checked and missing > checked * 0.35:
            issues.append(
                {
                    "type": "indentation_error",
                    "message": "Many long body paragraphs lack a clear first-line indent.",
                    "severity": "medium",
                }
            )

    if exp.get("require_page_numbers"):
        if not _docx_has_page_number_field(doc):
            issues.append(
                {
                    "type": "missing_page_numbers",
                    "message": "No PAGE field detected in headers or footers.",
                    "severity": "medium",
                }
            )

    # Heading styles for obvious section titles
    heading_style_count = 0
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if not t:
            continue
        n = normalize_paragraph_text(t)
        if n in COMMON_HEADINGS | REFS_HEADINGS:
            st = getattr(p.style, "name", "") or ""
            if st and not st.startswith("Heading"):
                heading_style_count += 1
                if heading_style_count <= 5:
                    issues.append(
                        {
                            "type": "heading_style",
                            "message": f"Section “{t.strip()[:60]}” does not use a Word Heading style.",
                            "severity": "low",
                        }
                    )

    # Extra space before/after (simple heuristic)
    weird = 0
    checked = 0
    for p in doc.paragraphs:
        if not (p.text or "").strip():
            continue
        checked += 1
        pf = p.paragraph_format
        sb = pf.space_before
        sa = pf.space_after
        b_pt = sb.pt if sb else 0
        a_pt = sa.pt if sa else 0
        if b_pt > 24 or a_pt > 24:
            weird += 1
    if checked and weird > checked * 0.2:
        issues.append(
            {
                "type": "paragraph_spacing",
                "message": "Unusually large space before/after on many paragraphs — check spacing settings.",
                "severity": "low",
            }
        )

    return _finalize(issues, notes)


def _finalize(issues: list[dict[str, str]], notes: list[str]) -> dict[str, Any]:
    high = sum(1 for i in issues if i.get("severity") == "high")
    med = sum(1 for i in issues if i.get("severity") == "medium")
    low = sum(1 for i in issues if i.get("severity") == "low")
    return {
        "issues": issues,
        "summary": {
            "total_issues": len(issues),
            "high_priority": high,
            "medium_priority": med,
            "low_priority": low,
            "ready_to_submit": high == 0,
        },
        "notes": notes,
    }
