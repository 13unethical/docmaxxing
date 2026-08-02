"""Real delivery packager: client gets only the final assignment file (docx/pdf)."""

from __future__ import annotations

import io
import json
import re
import uuid
from pathlib import Path
from typing import Any, Protocol

from docx import Document

from services.assignment_pipeline.models import utc_now
from services.delivery_engine.models import (
    DeliveryEngineInput,
    DeliveryFile,
    DeliveryPackage,
    DeliveryStatus,
    ProjectSummary,
)


def _storage_root() -> Path:
    # Inline to avoid circular import via services.assignment_project.__init__.
    import os

    override = (os.environ.get("PROJECT_STORAGE_DIR") or "").strip()
    if override:
        return Path(override).expanduser().resolve()
    return (Path(__file__).resolve().parents[2] / "data" / "projects").resolve()


class DeliveryPackager(Protocol):
    def package(self, payload: DeliveryEngineInput) -> DeliveryPackage: ...


class RealDeliveryPackager:
    VERSION = "real-2.0"

    def package(self, payload: DeliveryEngineInput) -> DeliveryPackage:
        project_id = payload.project_id or "local"
        draft = payload.final_draft
        requirement = payload.requirement_json
        research = payload.research_plan
        blueprint = payload.blueprint
        review = payload.review_report
        detection = payload.detection_report

        title = _safe_filename(
            str(draft.get("title") or requirement.get("title") or requirement.get("assignment_type") or "Assignment")
        )
        root = _storage_root() / project_id / "delivery"
        root.mkdir(parents=True, exist_ok=True)
        debug_root = root / "debug"
        debug_root.mkdir(parents=True, exist_ok=True)

        summary = _build_summary(payload, title, review, detection)
        now = utc_now()
        package_id = str(uuid.uuid4())
        client_format = resolve_client_format(requirement)

        # Always materialize formatted DOCX bytes (source of truth when Format Engine ran).
        formatted_path = (payload.formatted_document_path or "").strip()
        docx_bytes: bytes | None = None
        if formatted_path:
            src = Path(formatted_path)
            if src.is_file():
                docx_bytes = src.read_bytes()
        if docx_bytes is None:
            docx_bytes = _build_docx_bytes(str(draft.get("title") or title), str(draft.get("content") or ""))

        docx_path = root / f"{title}.docx"
        docx_path.write_bytes(docx_bytes)

        files: list[DeliveryFile] = []
        if client_format == "pdf":
            pdf_bytes = _build_pdf_bytes(str(draft.get("title") or title), str(draft.get("content") or ""))
            pdf_name = f"{title}.pdf"
            pdf_path = root / pdf_name
            pdf_path.write_bytes(pdf_bytes)
            files.append(
                _file(
                    "Final Assignment",
                    pdf_name,
                    "final_assignment_pdf",
                    "application/pdf",
                    pdf_path,
                    len(pdf_bytes),
                    package_id=package_id,
                )
            )
            primary_path = pdf_path
        else:
            files.append(
                _file(
                    "Formatted Assignment" if formatted_path else "Final Assignment",
                    f"{title}.docx",
                    "final_assignment_docx",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    docx_path,
                    len(docx_bytes),
                    package_id=package_id,
                )
            )
            primary_path = docx_path

        # Server-side debug artifacts — never included in client download.
        debug_specs = [
            ("requirement.json", requirement),
            ("research.json", research),
            ("blueprint.json", blueprint),
            ("review.json", review),
            ("ai_detection.json", detection),
            ("project_summary.json", summary.to_dict()),
        ]
        for filename, payload_obj in debug_specs:
            blob = json.dumps(payload_obj, ensure_ascii=False, indent=2).encode("utf-8")
            (debug_root / filename).write_bytes(blob)

        package = DeliveryPackage(
            id=package_id,
            project_id=payload.project_id,
            status=DeliveryStatus.READY,
            files=files,
            project_summary=summary,
            package_download_url=f"/api/assignment/projects/{project_id}/download",
            package_size_bytes=primary_path.stat().st_size,
            final_draft_id=str(draft.get("id") or ""),
            engine_version=self.VERSION,
            prepared_at=now,
            ready_at=now,
            client_format=client_format,
            client_filename=primary_path.name,
        )
        return package


def resolve_client_format(requirement: dict[str, Any] | None) -> str:
    """Prefer PDF only when the brief clearly asks for PDF; otherwise DOCX."""
    req = requirement or {}
    parts = [
        str(req.get("submission_format") or ""),
        str(req.get("submission_medium") or ""),
        str(req.get("assignment_medium") or ""),
        str(req.get("format") or ""),
    ]
    blob = " ".join(parts).lower()
    if re.search(r"\bpdf\b", blob) and not re.search(r"\b(word|docx|doc)\b", blob):
        return "pdf"
    if re.search(r"\bpdf\b", blob) and "word-processed" not in blob and "word processed" not in blob:
        # e.g. "PDF preferred" without Word — still pdf
        if "word" not in blob and "docx" not in blob:
            return "pdf"
    return "docx"


def _build_docx_bytes(title: str, content: str) -> bytes:
    doc = Document()
    doc.add_heading(title or "Assignment", level=1)
    for block in (content or "").split("\n\n"):
        text = block.strip()
        if text:
            doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_pdf_bytes(title: str, content: str) -> bytes:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
    except Exception:
        return _build_docx_bytes(title, content)

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    width, height = A4
    y = height - 40
    c.setFont("Helvetica-Bold", 14)
    c.drawString(40, y, title or "Assignment")
    y -= 24
    c.setFont("Helvetica", 10)
    for paragraph in (content or "").split("\n"):
        line = paragraph.strip()
        if not line:
            y -= 8
            continue
        chunks = [line[i : i + 110] for i in range(0, len(line), 110)]
        for chunk in chunks:
            if y < 40:
                c.showPage()
                c.setFont("Helvetica", 10)
                y = height - 40
            c.drawString(40, y, chunk)
            y -= 14
    c.showPage()
    c.save()
    return buf.getvalue()


def _build_summary(
    payload: DeliveryEngineInput,
    title: str,
    review: dict[str, Any],
    detection: dict[str, Any],
) -> ProjectSummary:
    req = payload.requirement_json
    draft = payload.final_draft
    completion_time = payload.completion_time if payload.completion_time and payload.completion_time != "—" else "Not available"
    review_score = int(review.get("overall_score") or 0)
    ai_score = float(detection.get("overall_ai_score") or detection.get("average_score") or 0)
    return ProjectSummary(
        project_name=title,
        assignment_type=str(req.get("assignment_type") or "Not available"),
        word_count=int(draft.get("total_words") or req.get("word_count") or 0),
        citation_style=str(req.get("citation_style") or "Not available"),
        difficulty=str(req.get("difficulty") or "Not available"),
        completion_time=completion_time,
        total_revisions=int(payload.revision_attempts),
        total_humanization_attempts=int(payload.humanization_attempts),
        overall_review_score=review_score,
        final_ai_score=ai_score,
        pipeline_completion_date=utc_now().date().isoformat(),
        overall_quality_score=review_score,
    )


def _file(
    label: str,
    filename: str,
    file_type: str,
    mime_type: str,
    path: Path,
    size: int,
    *,
    package_id: str = "",
) -> DeliveryFile:
    file_id = str(uuid.uuid4())
    return DeliveryFile(
        id=file_id,
        label=label,
        filename=filename,
        file_type=file_type,
        mime_type=mime_type,
        size_bytes=size,
        storage_path=str(path),
        download_url=f"/api/delivery/files/{file_id}",
        ready=True,
    )


def _safe_filename(value: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9._-]+", "-", value.strip()).strip("-")
    return cleaned or "Assignment"
