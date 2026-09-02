"""Format Engine stage for assignment projects — Formatter V2 pipeline."""

from __future__ import annotations

import io
import os
import re
import uuid
from pathlib import Path
from typing import Any

from docx import Document

from formatter_v2.pipeline import format_document_v2, resolve_style_name
from formatter_v2.profiles import load_profile
from formatter_v2.resolve import resolve_format_spec
from formatter_v2.spec import (
    Alignment,
    FontFamily,
    Margins,
    PageNumberPosition,
    PageNumbering,
    ParagraphRole,
    StyleName,
    UserOverrides,
)
from services.assignment_pipeline.models import utc_now
from services.assignment_spec.validate import count_body_words, count_words
from services.assignment_spec.word_count_statement import (
    LEADING_WORD_COUNT_LINE,
    requirement_asks_to_state_word_count,
)

_REPO_ROOT = Path(__file__).resolve().parents[2]


def _storage_root() -> Path:
    override = (os.environ.get("PROJECT_STORAGE_DIR") or "").strip()
    if override:
        return Path(override).expanduser().resolve()
    return (_REPO_ROOT / "data" / "projects").resolve()


def _parse_line_spacing(value: Any) -> float | None:
    """Accept numeric or Word-style labels like Double / Single / 1.5.

    Missing / unparseable values return None so the style profile supplies
    the default — Assignment no longer hardcodes 2.0.
    """
    if value is None or value == "":
        return None
    if isinstance(value, (int, float)):
        return float(value)
    text = str(value).strip().lower().replace("-", " ").replace("_", " ")
    aliases = {
        "single": 1.0,
        "single spacing": 1.0,
        "1.0": 1.0,
        "1": 1.0,
        "1.15": 1.15,
        "1.5": 1.5,
        "1.5 lines": 1.5,
        "one and a half": 1.5,
        "double": 2.0,
        "double spaced": 2.0,
        "double spacing": 2.0,
        "2.0": 2.0,
        "2": 2.0,
        "2.0 lines": 2.0,
        "triple": 3.0,
    }
    if text in aliases:
        return aliases[text]
    for key, num in aliases.items():
        if key in text and key.isalpha():
            return num
    try:
        return float(text.split()[0])
    except (TypeError, ValueError, IndexError):
        return None


def _parse_optional_float(value: Any) -> float | None:
    if value is None or value == "":
        return None
    if isinstance(value, (int, float)):
        return float(value)
    text = str(value).strip().split()[0]
    try:
        return float(text)
    except (TypeError, ValueError):
        return None


def _parse_font_family(value: Any) -> FontFamily | None:
    if not value:
        return None
    text = str(value).strip()
    for fam in FontFamily:
        if fam.value.lower() == text.lower():
            return fam
        if fam.name.lower() == text.lower().replace(" ", "_"):
            return fam
    return None


def _parse_alignment(value: Any) -> Alignment | None:
    if not value:
        return None
    text = str(value).strip().lower()
    if text in {"centre", "center"}:
        return Alignment.CENTER
    try:
        return Alignment(text)
    except ValueError:
        return None


def _parse_margins(fmt: dict[str, Any]) -> Margins | None:
    raw_margins = fmt.get("margins")
    if isinstance(raw_margins, dict):
        try:
            return Margins.model_validate(raw_margins)
        except Exception:  # noqa: BLE001
            return None
    preset = str(fmt.get("margin_preset") or "").strip().lower()
    if preset in {"normal", "narrow", "wide"}:
        return Margins.preset(preset)  # type: ignore[arg-type]
    raw = fmt.get("margins_inches")
    if raw is None:
        raw = raw_margins
    if raw is None or raw == "":
        return None
    if isinstance(raw, (int, float)):
        value = float(raw)
        return Margins(top_in=value, bottom_in=value, left_in=value, right_in=value)
    match = re.search(r"(\d+(?:\.\d+)?)", str(raw))
    if not match:
        return None
    value = float(match.group(1))
    return Margins(top_in=value, bottom_in=value, left_in=value, right_in=value)


def _parse_page_numbering(value: Any) -> PageNumbering | None:
    if not value:
        return None
    text = str(value).strip().lower().replace("-", "_").replace(" ", "_")
    try:
        return PageNumbering(position=PageNumberPosition(text))
    except ValueError:
        return None


def _margin_preset_name(margins: Margins) -> str:
    for name in ("normal", "narrow", "wide"):
        preset = Margins.preset(name)  # type: ignore[arg-type]
        if (
            preset.top_in == margins.top_in
            and preset.bottom_in == margins.bottom_in
            and preset.left_in == margins.left_in
            and preset.right_in == margins.right_in
        ):
            return name
    return "custom"


def _style_from_requirement(requirement_json: dict[str, Any], fmt: dict[str, Any]) -> StyleName:
    raw = (
        fmt.get("style")
        or requirement_json.get("format_style")
        or requirement_json.get("citation_style")
        or "harvard"
    )
    return resolve_style_name(str(raw))


def _overrides_from_requirement(requirement_json: dict[str, Any]) -> tuple[UserOverrides, StyleName]:
    """Map brief fields that were actually specified onto UserOverrides.

    Unspecified fields stay None so StyleProfile defaults apply — including
    line spacing and page-number position.
    """
    fmt = requirement_json.get("formatting") if isinstance(requirement_json.get("formatting"), dict) else {}
    style = _style_from_requirement(requirement_json, fmt)
    data: dict[str, Any] = {}

    font_family = _parse_font_family(fmt.get("font_family"))
    if font_family is not None:
        data["font_family"] = font_family

    font_size = _parse_optional_float(fmt.get("font_size_pt") if fmt.get("font_size_pt") is not None else fmt.get("font_size"))
    if font_size is not None:
        data["font_size_pt"] = font_size

    spacing = _parse_line_spacing(
        fmt.get("line_spacing") if fmt.get("line_spacing") is not None else requirement_json.get("line_spacing")
    )
    if spacing is not None:
        data["line_spacing"] = spacing

    alignment = _parse_alignment(fmt.get("alignment"))
    if alignment is not None:
        data["alignment"] = alignment

    if "first_line_indent" in fmt:
        data["first_line_indent"] = bool(fmt.get("first_line_indent"))

    margins = _parse_margins(fmt)
    if margins is not None:
        data["margins"] = margins

    page_numbering = _parse_page_numbering(fmt.get("page_number_position"))
    if page_numbering is not None:
        data["page_numbering"] = page_numbering

    heading_size = _parse_optional_float(fmt.get("heading_size_pt"))
    if heading_size is not None:
        data["heading_size_pt"] = heading_size

    return UserOverrides.model_validate(data), style


def _profile_summary(spec) -> dict[str, Any]:
    body = spec.roles[ParagraphRole.BODY]
    return {
        "font_family": body.font_family.value,
        "font_size_pt": float(body.font_size_pt),
        "line_spacing": float(body.line_spacing),
        "alignment": body.alignment.value,
        "margin_preset": _margin_preset_name(spec.page.margins),
        "page_number_position": spec.page_numbering.position.value,
    }


def _heading_level_and_title(line: str) -> tuple[int, str] | None:
    if line.startswith("## "):
        return 2, line[3:].strip()
    if line.startswith("# "):
        return 1, line[2:].strip()
    return None


def _is_references_heading_title(title: str) -> bool:
    from formatter_v2.structure.references import is_references_heading

    return is_references_heading(title)


def _add_reference_paragraphs(doc: Document, blob: str) -> None:
    from formatter_v2.structure.references import split_concatenated_reference_entries

    for entry in split_concatenated_reference_entries(blob):
        if entry:
            doc.add_paragraph(entry)


def _docx_from_markdown(title: str, content: str) -> Document:
    """Build a docx from draft text.

    Paragraph boundaries are blank lines (\\n\\n). Single newlines are soft wraps
    and must stay inside one paragraph — never create a paragraph per wrapped line.

    Critical: a block that starts with ``## Heading`` followed by body on the next
    line must become Heading + Normal — never one giant Heading paragraph.

    References are the exception: each bibliographic entry must be its own
    paragraph so hanging indent can apply.
    """
    from docx.enum.text import WD_BREAK

    from formatter.markdown_cleanup import clean_markdown_in_document
    from formatter_v2.structure.references import is_refs_latch_breaker

    doc = Document()
    if title.strip():
        doc.add_heading(title.strip(), level=1)

    text = re.sub(r"<!--\s*pagebreak\s*-->", "\n\n[[PAGEBREAK]]\n\n", content or "", flags=re.I)
    blocks = re.split(r"\n\s*\n", text)
    in_references = False
    for block in blocks:
        raw = (block or "").strip()
        if not raw:
            continue
        if raw == "[[PAGEBREAK]]":
            p = doc.add_paragraph()
            run = p.add_run()
            run.add_break(WD_BREAK.PAGE)
            continue
        lines = [ln.strip() for ln in raw.split("\n") if ln.strip()]
        if not lines:
            continue

        heading = _heading_level_and_title(lines[0])
        if heading:
            level, heading_title = heading
            in_references = _is_references_heading_title(heading_title)
            doc.add_heading(heading_title, level=level)
            rest = "\n".join(lines[1:]).strip()
            if rest:
                if in_references:
                    _add_reference_paragraphs(doc, rest)
                else:
                    doc.add_paragraph(" ".join(lines[1:]).strip())
            continue

        if in_references:
            if is_refs_latch_breaker(" ".join(lines)):
                in_references = False
                para_text = " ".join(lines)
                doc.add_paragraph(para_text)
            else:
                _add_reference_paragraphs(doc, "\n".join(lines))
            continue

        para_text = " ".join(lines)
        doc.add_paragraph(para_text)

    clean_markdown_in_document(doc)
    return doc


class AssignmentFormatEngine:
    VERSION = "format-engine-2.0"

    def format_draft(
        self,
        *,
        draft: dict[str, Any],
        requirement_json: dict[str, Any],
        project_id: str,
        citation_pack: dict[str, Any] | None = None,
    ) -> dict[str, Any]:
        del citation_pack  # references already embedded in draft content when available
        title = str(draft.get("title") or requirement_json.get("title") or "Assignment")
        content = str(draft.get("content") or "")
        body_words = int(draft.get("total_words") or count_body_words(content))
        # Only print "Word count:" when the brief asks to state it on the paper.
        if requirement_asks_to_state_word_count(requirement_json):
            if body_words > 0 and not re.search(r"(?im)^\s*word\s*count\s*:", content):
                content = f"Word count: {body_words}\n\n{content.lstrip()}"
        else:
            content = LEADING_WORD_COUNT_LINE.sub("", content, count=1)

        overrides, style_name = _overrides_from_requirement(requirement_json)
        profile = load_profile(style_name)
        resolution = resolve_format_spec(profile, overrides)

        document = _docx_from_markdown(title, content)
        source = io.BytesIO()
        document.save(source)
        result = format_document_v2(source.getvalue(), overrides, style_name)

        out_dir = _storage_root() / project_id / "formatted"
        out_dir.mkdir(parents=True, exist_ok=True)
        filename = "formatted.docx"
        path = out_dir / filename
        path.write_bytes(result.docx_bytes)

        notices = [n.model_dump(mode="json") for n in result.notices]
        return {
            "id": str(uuid.uuid4()),
            "project_id": project_id,
            "path": str(path),
            "filename": filename,
            "style_id": style_name.value,
            "word_count": int(draft.get("total_words") or count_body_words(content)),
            "body_word_count": int(draft.get("total_words") or count_body_words(content)),
            "document_word_count": count_words(content),
            "profile_summary": _profile_summary(resolution.spec),
            "notices": notices,
            "applied_rules": [
                "v2_pipeline",
                result.extractor_name,
                "resolve_format_spec",
                "build_document",
            ],
            "engine_version": self.VERSION,
            "formatted_at": utc_now().isoformat(),
            "source_draft_id": str(draft.get("id") or ""),
            "plain_text": content,
        }
