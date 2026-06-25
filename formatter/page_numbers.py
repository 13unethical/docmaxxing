"""
Insert a live Word PAGE field in a header or footer paragraph.

python-docx has no high-level API for fields, so we append the OOXML pieces
Word expects: begin → instruction text → separate → (placeholder text) → end.
"""

from __future__ import annotations

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def add_page_field(paragraph) -> None:
    """Append a PAGE field to an existing paragraph (after any existing runs)."""
    run = paragraph.add_run()
    r = run._r

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")

    # Placeholder; Word replaces this when the document opens.
    placeholder = OxmlElement("w:t")
    placeholder.text = "1"

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    r.append(fld_begin)
    r.append(instr)
    r.append(fld_separate)
    r.append(placeholder)
    r.append(fld_end)


def _strip_part_paragraphs(part) -> None:
    """Clear a header/footer part so page numbers never duplicate legacy content."""
    root = part._element
    for child in list(root):
        root.remove(child)
    part.add_paragraph()


def _align_from_position(position: str) -> WD_ALIGN_PARAGRAPH:
    if position.endswith("left"):
        return WD_ALIGN_PARAGRAPH.LEFT
    return WD_ALIGN_PARAGRAPH.RIGHT


def apply_page_numbers_to_document(document, position: str) -> None:
    """
    position: 'none' | 'top_left' | 'top_right' | 'bottom_left' | 'bottom_right'

    Clears headers/footers for every section and inserts the field once per section
    in the chosen location. Other header/footer areas stay empty so numbering
    does not duplicate.
    """
    pos = (position or "none").lower().strip()
    if pos == "none" or pos == "":
        return

    use_header = pos.startswith("top")
    for section in document.sections:
        _strip_part_paragraphs(section.header)
        _strip_part_paragraphs(section.footer)

        if use_header:
            p = section.header.paragraphs[0]
        else:
            p = section.footer.paragraphs[0]

        p.alignment = _align_from_position(pos)
        add_page_field(p)
