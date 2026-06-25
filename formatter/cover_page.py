"""Academic cover / title page generation for DOCX output."""

from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Pt


@dataclass
class CoverPageData:
    assignment_title: str = ""
    student_name: str = ""
    student_id: str = ""
    university: str = ""
    module: str = ""
    lecturer: str = ""
    submission_date: str = ""

    def has_content(self) -> bool:
        return any(
            str(getattr(self, field) or "").strip()
            for field in (
                "assignment_title",
                "student_name",
                "student_id",
                "university",
                "module",
                "lecturer",
                "submission_date",
            )
        )


def _display_date(raw: str) -> str:
    text = (raw or "").strip()
    if not text:
        return ""
    try:
        parsed = datetime.strptime(text[:10], "%Y-%m-%d")
        return parsed.strftime("%d %B %Y")
    except ValueError:
        return text


def prepend_cover_page(document: Document, cover: CoverPageData, *, font_family: str = "Times New Roman") -> None:
    """Insert a centered title page before existing body content, then a page break."""
    if not cover.has_content():
        return

    body = document.element.body
    existing = list(body)
    for element in existing:
        body.remove(element)

    def add_spacer(lines: int = 1) -> None:
        for _ in range(lines):
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(10)

    def add_centered(
        text: str,
        *,
        size_pt: int = 12,
        bold: bool = False,
        space_after_pt: int = 8,
    ) -> None:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(space_after_pt)
        run = paragraph.add_run(text)
        run.font.name = font_family
        run.font.size = Pt(size_pt)
        run.bold = bold

    add_spacer(5)

    title = (cover.assignment_title or "Assignment").strip()
    add_centered(title, size_pt=18, bold=True, space_after_pt=28)

    detail_lines: list[str] = []
    if cover.student_name.strip():
        detail_lines.append(cover.student_name.strip())
    if cover.student_id.strip():
        detail_lines.append(f"Student ID: {cover.student_id.strip()}")
    if cover.university.strip():
        detail_lines.append(cover.university.strip())
    if cover.module.strip():
        detail_lines.append(cover.module.strip())
    if cover.lecturer.strip():
        detail_lines.append(f"Lecturer: {cover.lecturer.strip()}")
    date_text = _display_date(cover.submission_date)
    if date_text:
        detail_lines.append(date_text)

    for line in detail_lines:
        add_centered(line, size_pt=12, space_after_pt=6)

    add_spacer(2)
    break_paragraph = document.add_paragraph()
    break_paragraph.add_run().add_break(WD_BREAK.PAGE)

    for element in existing:
        body.append(element)
