"""
Append a References block to an already formatted document and style new paragraphs.
"""

from __future__ import annotations

from docx import Document

from formatter.format_job import FormatJob
from formatter.headings import (
    apply_heading_caps,
    detect_heading_level,
    is_references_heading,
)
from formatter.paragraph_style import format_paragraph
from formatter.style_engine import (
    resolve_active_profile,
    resolve_contextual_spacing,
    role_for_paragraph,
)


def append_references_section(
    document: Document,
    job: FormatJob,
    citations: list[str],
    *,
    section_title: str = "References",
) -> None:
    """
    Add a section heading and one paragraph per citation using the active profile.
    """
    cleaned = [c.strip() for c in citations if c and str(c).strip()]
    if not cleaned:
        return

    profile = resolve_active_profile(job)
    heading = (section_title or "References").strip() or "References"
    n_before = len(document.paragraphs)
    document.add_paragraph(heading)
    for c in cleaned:
        document.add_paragraph(c)

    in_refs_section = False
    for paragraph in document.paragraphs[n_before:]:
        text = paragraph.text
        refs_title = is_references_heading(text)
        if refs_title:
            in_refs_section = True

        level = detect_heading_level(text, job.auto_headings, is_first_nonempty=False)
        apply_heading_caps(paragraph, job.heading_all_caps, level)

        role = role_for_paragraph(
            level=level,
            in_refs_section=in_refs_section,
            is_refs_title=refs_title,
        )
        spec = profile.paragraph_spec_for_role(role)
        space_before, space_after = resolve_contextual_spacing(
            profile,
            role=role,
            prev_level=0,
            next_level=0,
            prev_has_text=n_before > 0,
        )

        format_paragraph(
            paragraph,
            document,
            spec=spec,
            space_before_pt=space_before,
            space_after_pt=space_after,
            heading_level=level,
        )
