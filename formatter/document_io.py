"""Create documents from pasted text or uploaded .docx / .pdf files."""

from __future__ import annotations

import io
import os
import re

from docx import Document

from formatter.text_cleaning import clean_pasted_raw, paragraphs_snapshot_apply_cleaning

UPLOAD_DOCUMENT_EXTENSIONS = frozenset({".docx", ".pdf"})


def upload_extension(filename: str, mimetype: str | None = None) -> str:
    """Resolve .docx / .pdf from filename or MIME type."""
    ext = os.path.splitext((filename or "").lower())[1]
    if ext in UPLOAD_DOCUMENT_EXTENSIONS:
        return ext
    mt = (mimetype or "").lower()
    if "pdf" in mt:
        return ".pdf"
    if "wordprocessingml" in mt or mt == "application/msword":
        return ".docx"
    return ext


def is_supported_document_upload(filename: str, mimetype: str | None = None) -> bool:
    return upload_extension(filename, mimetype) in UPLOAD_DOCUMENT_EXTENSIONS


def extract_text_from_docx_bytes(raw: bytes) -> str:
    """Extract body text; each paragraph (including headings) stays a separate block."""
    doc = Document(io.BytesIO(raw))
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        text = (paragraph.text or "").strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts).strip()


def extract_text_from_pdf_bytes(raw: bytes) -> str:
    from services.requirements_ocr import extract_text_from_pdf_bytes as _pdf_extract

    return _pdf_extract(raw)


def extract_text_from_document_bytes(raw: bytes, filename: str, mimetype: str | None = None) -> str:
    if not raw:
        raise ValueError("The uploaded file is empty.")
    ext = upload_extension(filename, mimetype)
    if ext == ".docx":
        text = extract_text_from_docx_bytes(raw)
    elif ext == ".pdf":
        text = extract_text_from_pdf_bytes(raw)
    else:
        raise ValueError("Unsupported file type. Upload a .docx or .pdf file.")
    if not text.strip():
        raise ValueError("No readable text found in the uploaded file.")
    return text


def document_has_visible_content(doc: Document) -> bool:
    if any((p.text or "").strip() for p in doc.paragraphs):
        return True
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if (cell.text or "").strip():
                    return True
    return False


def document_from_pasted_text(pasted: str) -> Document:
    """One Word paragraph per blank-line-separated block (\\n\\n)."""
    doc = Document()
    if doc.paragraphs:
        body = doc.paragraphs[0]._element
        body.getparent().remove(body)

    if not pasted.strip():
        doc.add_paragraph("")
        return doc

    blocks = re.split(r"\n\s*\n", pasted)
    for block in blocks:
        text = block.strip("\n")
        doc.add_paragraph(text)
    return doc


def build_document_from_inputs(
    *,
    pasted_raw: str | None,
    file_bytes: bytes | None,
    cleaning_spaces: bool,
    cleaning_breaks: bool,
) -> Document:
    """
    Load from upload or paste. Cleaning (optional) runs on paste as a whole string,
    or per-paragraph after load for .docx uploads.
    """
    if file_bytes is not None:
        doc = Document(io.BytesIO(file_bytes))
        if cleaning_spaces or cleaning_breaks:
            paragraphs_snapshot_apply_cleaning(
                doc,
                extra_spaces=cleaning_spaces,
                extra_linebreaks=cleaning_breaks,
            )
        return doc

    # Paste path
    text = pasted_raw or ""
    text = clean_pasted_raw(
        text,
        extra_spaces=cleaning_spaces,
        extra_linebreaks=cleaning_breaks,
    )
    return document_from_pasted_text(text)


def build_document_from_upload(
    file_bytes: bytes,
    filename: str,
    *,
    mimetype: str | None = None,
    cleaning_spaces: bool,
    cleaning_breaks: bool,
) -> Document:
    """Load .docx directly; convert PDF text into a new Word document."""
    ext = upload_extension(filename, mimetype)
    if ext == ".docx":
        return build_document_from_inputs(
            pasted_raw=None,
            file_bytes=file_bytes,
            cleaning_spaces=cleaning_spaces,
            cleaning_breaks=cleaning_breaks,
        )
    if ext == ".pdf":
        text = extract_text_from_pdf_bytes(file_bytes)
        return build_document_from_inputs(
            pasted_raw=text,
            file_bytes=None,
            cleaning_spaces=cleaning_spaces,
            cleaning_breaks=cleaning_breaks,
        )
    raise ValueError("Unsupported file type. Upload a .docx or .pdf file.")
