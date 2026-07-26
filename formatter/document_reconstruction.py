"""
Academic document reconstruction engine.

Rebuilds intended structure BEFORE formatting. Paragraph boundaries from the input
are treated as unreliable; headings may appear inside paragraphs after humanizers,
OCR, or PDF extraction.

Pipeline position:
  load document → reconstruct_document_before_format() → format_document_full()
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from typing import Any

from docx import Document

from formatter.heading_plan import ParagraphHeadingAssignment, StructureApplyResult
from formatter.requirement_headings import (
    _canonical_label_key,
    _collect_heading_splits,
    _looks_like_citation_start,
    _match_variants,
    _normalize_spaces,
    _valid_split_at,
    display_section_heading,
    normalize_document_internal_spaces,
)
from services.document_structure_engine import (
    SECTION_BLUEPRINTS,
    _infer_document_type,
    _looks_like_title,
    _normalize_doc_type,
    _paragraph_is_reference_line,
    detect_heading_level,
    is_heading_like,
    is_references_heading,
    normalize_paragraph_text,
    recover_structure,
)
# Data model
# ---------------------------------------------------------------------------

_BODY_START_WORDS = frozenset(
    {
        "today",
        "this",
        "the",
        "one",
        "even",
        "many",
        "some",
        "these",
        "there",
        "in",
        "our",
        "we",
        "they",
        "it",
        "as",
        "when",
        "while",
        "universities",
        "students",
        "governments",
        "another",
        "finally",
        "artificial",
        "smith",
        "university",
        "during",
        "over",
        "although",
        "however",
        "therefore",
        "because",
    }
)
_STOPWORDS = frozenset(
    {"a", "an", "and", "as", "at", "by", "for", "from", "in", "of", "on", "or", "the", "to", "with"}
)


@dataclass
class ExpectedSection:
    """A section the document is expected to contain."""

    label: str
    canonical: str
    level: int  # 1 = title, 2 = major section, 3 = subsection
    priority: float  # 0–1; assignment requirements = 1.0
    source: str  # requirement | blueprint | vocabulary | journal


@dataclass
class ReconstructedBlock:
    """One paragraph after reconstruction."""

    kind: str  # title | heading | body | reference
    text: str
    level: int
    source: str
    confidence: float
    section_key: str = ""


@dataclass
class ReconstructionResult:
    assignments: list[ParagraphHeadingAssignment]
    recovery_mode: str
    ai_powered: bool
    blocks: list[ReconstructedBlock] = field(default_factory=list)


# ---------------------------------------------------------------------------
# Expected structure
# ---------------------------------------------------------------------------


def _journal_entry_labels(limit: int = 6) -> list[str]:
    return [f"Journal Entry {i}" for i in range(1, limit + 1)]


def _body_paragraph_labels(limit: int = 8) -> list[str]:
    return [f"Body Paragraph {i}" for i in range(1, limit + 1)]


def build_expected_structure(
    *,
    document_type: str | None,
    required_sections: list[str] | None,
    paragraphs: list[str] | None = None,
) -> list[ExpectedSection]:
    """
    Build ordered expected sections from assignment requirements, document type,
    and heading vocabulary. Requirements always win on priority.
    """
    doc_type = _normalize_doc_type(document_type)
    if doc_type == "other" and paragraphs:
        doc_type, _ = _infer_document_type(paragraphs, None)

    sections: list[ExpectedSection] = []
    seen: set[str] = set()

    def add(label: str, level: int, priority: float, source: str) -> None:
        key = _canonical_label_key(label)
        if not key or key in seen:
            return
        seen.add(key)
        sections.append(
            ExpectedSection(
                label=label.strip(),
                canonical=key,
                level=level,
                priority=priority,
                source=source,
            )
        )

    for label in required_sections or []:
        lvl = 1 if _canonical_label_key(label) == "title" else 2
        add(label, lvl, 1.0, "requirement")

    blueprint = SECTION_BLUEPRINTS.get(doc_type, SECTION_BLUEPRINTS["other"])
    for label in blueprint:
        key = _canonical_label_key(label)
        if key in {"title", "main body", "preamble"}:
            continue
        add(label, 2, 0.72, "blueprint")

    for label in _body_paragraph_labels():
        add(label, 2, 0.68, "vocabulary")

    for label in _journal_entry_labels():
        add(label, 2, 0.72, "journal")

    if doc_type in {"learning_journal", "reflection"}:
        add("Reflection", 2, 0.7, "vocabulary")

    for label in ("Introduction", "Conclusion", "Concluding Paragraph", "References", "Bibliography", "Works Cited"):
        add(label, 2, 0.65, "vocabulary")

    sections.sort(key=lambda s: (-s.priority, s.label))
    return sections


# ---------------------------------------------------------------------------
# Multi-signal scoring
# ---------------------------------------------------------------------------


def _label_pattern(label: str) -> re.Pattern[str]:
    words = label.strip().split()
    inner = r"\s+".join(re.escape(w) for w in words)
    return re.compile(rf"(?i)\b{inner}\b")


def _score_body_likelihood(text: str, offset: int) -> float:
    remainder = text[offset:].lstrip()
    if not remainder or len(remainder) < 6:
        return 0.0
    score = 0.0
    if _looks_like_citation_start(remainder):
        return 0.92
    words = remainder.split()
    first = words[0]
    if len(words) >= 2 and words[1] and words[1][0].islower():
        score += 0.42
    if first.lower() in _BODY_START_WORDS:
        score += 0.4
    if len(remainder) > 90:
        score += 0.12
    if re.search(r"[.!?]", remainder[:140]):
        score += 0.08
    if first[0].isupper() and len(words) >= 2 and words[1][0].islower():
        score += 0.15
    return min(1.0, score)


def _score_subtitle_likelihood(text: str, offset: int) -> float:
    remainder = text[offset:].lstrip()
    if not remainder:
        return 0.0
    words = remainder.split()
    if len(words) > 14:
        return 0.0
    if len(words) >= 2 and words[1][0].islower():
        return 0.05
    title_like = sum(
        1 for w in words[:10] if w[0].isupper() or w.lower() in _STOPWORDS or w.isdigit()
    )
    return min(1.0, title_like / max(1, min(len(words), 10)))


_SUBTITLE_STOP_WORDS = frozenset(
    {"a", "an", "and", "at", "by", "for", "from", "in", "of", "on", "or", "the", "to", "with"}
)


def _remainder_is_title_subtitle_only(remainder: str) -> bool:
    """True when text after a colon is only a title phrase (no body sentence)."""
    text = remainder.strip()
    if not text or text.rstrip().endswith("."):
        return False
    words = text.split()
    if len(words) > 16:
        return False
    lowercase_content = sum(
        1
        for w in words[1:]
        if w and w[0].islower() and w.lower() not in _SUBTITLE_STOP_WORDS
    )
    return lowercase_content == 0


def _infer_label_start(text: str, label_end: int) -> int:
    """Locate the start of the heading label that ends at ``label_end``."""
    chunk = text[:label_end]
    patterns = (
        r"(?:Journal\s+Entry|Body\s+Paragraph|Section|Part|Chapter|Unit|Module|Week|Entry)\s+\d+\s*$",
        r"(?:Learning\s+Outcomes|Literature\s+Review|Executive\s+Summary|Concluding\s+Paragraph|"
        r"Works\s+Cited|Reflection|References|Bibliography|Introduction|Conclusion|"
        r"Abstract|Appendix|Discussion|Methodology|Methods|Results)\s*$",
    )
    for pat in patterns:
        m = re.search(pat, chunk, re.IGNORECASE)
        if m:
            return m.start()
    m = re.search(r"\S+\s*$", chunk)
    return m.start() if m else 0


def _extend_heading_text(text: str, label_end: int, label_start: int | None = None) -> int:
    """
    After a matched label, return the index where body text begins.

    Uses HeadingDetector for labeled journal / section headings so the first
    sentence is never absorbed into the heading. Never promotes stopwords.
    Scoped to the current label occurrence — mid-document labels must not
    reuse an earlier heading's body boundary.
    """
    from services.heading_detector import DEFAULT_HEADING_DETECTOR

    if label_start is None:
        label_start = _infer_label_start(text, label_end)
    label_start = max(0, min(label_start, label_end))

    scoped = text[label_start:]
    split = DEFAULT_HEADING_DETECTOR.split_embedded(scoped.strip())
    if split is not None and split.body:
        body = split.body
        body_lstrip = body.lstrip()
        for candidate in (
            body,
            body_lstrip,
            body_lstrip[:1].lower() + body_lstrip[1:] if body_lstrip else body,
        ):
            if not candidate:
                continue
            idx = text.find(candidate, label_end)
            if idx >= label_end:
                return idx

    pos = label_end
    remainder = text[pos:].lstrip()
    if not remainder:
        return pos

    label_prefix = text[label_start:label_end].strip()
    label_key = normalize_paragraph_text(label_prefix.split(":")[0])
    simple_labels = {
        "introduction",
        "conclusion",
        "references",
        "bibliography",
        "works cited",
        "abstract",
        "discussion",
        "methods",
        "results",
        "reflection",
    }
    if label_key in simple_labels:
        after = re.match(r"\s*\.?\s*", text[pos:])
        if after:
            return pos + after.end()
        return pos

    if re.match(r"^(body paragraph \d+|journal entry \d+)$", label_key):
        colon = re.match(r"\s*:\s*", text[pos:])
        if not colon:
            return pos
        after_colon_at = pos + colon.end()
        rest = text[after_colon_at:]
        for match in re.finditer(r"\.(?:\s+)(?=[A-Z“\"‘'])", rest):
            return after_colon_at + match.end()
        return len(text)

    colon = re.match(r"\s*:\s*", text[pos:])
    if colon:
        after_colon_at = pos + colon.end()
        rest = text[after_colon_at:]
        for match in re.finditer(r"\.(?:\s+)(?=[A-Z“\"‘'])", rest):
            return after_colon_at + match.end()
        return len(text)

    return pos



def _score_heading_candidate(
    text: str,
    start: int,
    heading_end: int,
    section: ExpectedSection,
    *,
    para_index: int,
    para_count: int,
    explicit_keys: set[str],
    word_heading: bool,
) -> float:
    label_key = section.canonical
    label_end = heading_end
    for variant in _match_variants(section.label):
        m = _label_pattern(variant).search(text, start)
        if m and m.start() == start:
            label_end = m.end()
            break

    heading_end = _extend_heading_text(text, label_end, label_start=start)
    if not _valid_split_at(text, start, label_end, label_key, explicit_labels=explicit_keys):
        return 0.0

    score = section.priority * 0.45
    if section.source == "requirement":
        score += 0.25
    if start == 0:
        score += 0.12
    elif start > 0 and text[start - 1] in ".!?\n":
        score += 0.1
    elif start > 0 and text[start - 1].isspace():
        score += 0.06

    body_score = _score_body_likelihood(text, heading_end)
    score += body_score * 0.22

    if word_heading:
        score += 0.18

    rel = para_index / max(1, para_count - 1)
    if label_key in {"introduction", "title"} and rel < 0.25:
        score += 0.08
    if label_key in {"references", "bibliography", "works cited"} and rel > 0.55:
        score += 0.1
    if label_key == "conclusion" and 0.45 < rel < 0.92:
        score += 0.06

    return min(1.0, score)


def _find_best_heading_in_text(
    text: str,
    sections: list[ExpectedSection],
    used: set[str],
    *,
    para_index: int,
    para_count: int,
    explicit_keys: set[str],
    word_heading: bool,
    search_from: int = 0,
) -> tuple[int, int, str, ExpectedSection, float] | None:
    best: tuple[int, int, str, ExpectedSection, float] | None = None
    for section in sections:
        if section.canonical in used:
            continue
        for variant in _match_variants(section.label):
            for m in _label_pattern(variant).finditer(text, search_from):
                start = m.start()
                label_end = m.end()
                heading_end = _extend_heading_text(text, label_end, label_start=start)
                conf = _score_heading_candidate(
                    text,
                    start,
                    heading_end,
                    section,
                    para_index=para_index,
                    para_count=para_count,
                    explicit_keys=explicit_keys,
                    word_heading=word_heading,
                )
                if conf < 0.42:
                    continue
                heading_text = text[start:heading_end].strip()
                if not heading_text:
                    continue
                if best is None or conf > best[4] or (conf == best[4] and start < best[0]):
                    best = (start, heading_end, heading_text, section, conf)
                break
    return best


def _segment_paragraph(
    text: str,
    sections: list[ExpectedSection],
    used: set[str],
    *,
    para_index: int,
    para_count: int,
    explicit_keys: set[str],
    word_heading: bool,
) -> list[tuple[str | None, str, ExpectedSection | None, float]]:
    """Split one input paragraph into (heading, body, section, confidence) segments."""
    stripped = text.strip()
    if not stripped:
        return []

    from services.heading_detector import DEFAULT_HEADING_DETECTOR

    # Deterministic peel first — identify heading, then body.
    # Soft wraps must not create fake paragraph boundaries mid-sentence.
    peel = DEFAULT_HEADING_DETECTOR.split_embedded(stripped)
    if peel is not None and peel.body:
        sec = _match_section_for_text(peel.heading, sections)
        if sec:
            used.add(sec.canonical)
        return [(peel.heading, peel.body, sec, 0.95)]

    stripped = DEFAULT_HEADING_DETECTOR.collapse_soft_linebreaks(stripped)

    if "\n" in stripped:
        first_line, rest = stripped.split("\n", 1)
        first_line = first_line.strip()
        rest = rest.strip()
        if first_line and rest and DEFAULT_HEADING_DETECTOR.is_heading_only_line(first_line):
            sec = _match_section_for_text(first_line, sections)
            if sec or is_heading_like(first_line):
                if sec:
                    used.add(sec.canonical)
                return [(first_line, rest, sec, 0.9)]

    # Standalone short heading line — never promote forbidden tokens.
    if DEFAULT_HEADING_DETECTOR.is_forbidden_heading(stripped):
        return [(None, stripped, None, 0.0)]

    if word_heading or (
        is_heading_like(stripped)
        and len(stripped.split()) <= 14
        and not _score_body_likelihood(stripped, len(stripped.split()[0]) + 1 if stripped.split() else 0) > 0.5
    ):
        sec = _match_section_for_text(stripped, sections)
        if sec and sec.canonical not in used:
            used.add(sec.canonical)
            return [(stripped, "", sec, 0.88)]
        lvl = detect_heading_level(stripped, True, is_first_nonempty=para_index == 0)
        if lvl > 0:
            return [(stripped, "", None, 0.8)]

    segments: list[tuple[str | None, str, ExpectedSection | None, float]] = []
    pos = 0
    while pos < len(stripped):
        remainder = stripped[pos:].lstrip()
        if not remainder:
            break
        offset = pos + (len(stripped[pos:]) - len(remainder))

        hit = _find_best_heading_in_text(
            stripped,
            sections,
            used,
            para_index=para_index,
            para_count=para_count,
            explicit_keys=explicit_keys,
            word_heading=word_heading,
            search_from=offset,
        )
        if hit is None:
            tail = stripped[pos:].strip()
            if tail:
                segments.append((None, tail, None, 0.0))
            break

        start, heading_end, heading_text, section, conf = hit
        if start > pos:
            pre = stripped[pos:start].strip()
            if pre:
                segments.append((None, pre, None, 0.0))

        used.add(section.canonical)
        body = stripped[heading_end:].strip()
        segments.append((heading_text, body, section, conf))
        if not body:
            break
        pos = heading_end

    if not segments:
        return [(None, stripped, None, 0.0)]
    return segments


def _segments_from_expected_labels(
    text: str,
    sections: list[ExpectedSection],
    used: set[str],
) -> list[tuple[str | None, str, ExpectedSection | None, float]]:
    """Split text at requirement / blueprint labels, including mid-paragraph."""
    stripped = text.strip()
    if not stripped:
        return []

    labels = sorted([s.label for s in sections], key=lambda label: len(label.split()), reverse=True)
    splits = _collect_heading_splits(stripped, labels, used)
    if splits:
        filtered: list[tuple[int, int, str, str]] = []
        for start, end, matched, label_key in sorted(splits, key=lambda x: x[0]):
            dominated = False
            for fstart, fend, _, _ in filtered:
                if start >= fstart and start < _extend_heading_text(stripped, fend, label_start=fstart):
                    dominated = True
                    break
            if not dominated:
                filtered.append((start, end, matched, label_key))
        splits = filtered
    if splits:
        segments: list[tuple[str | None, str, ExpectedSection | None, float]] = []
        pos = 0
        for i, (start, end, matched, label_key) in enumerate(splits):
            if start > pos:
                pre = stripped[pos:start].strip()
                if pre:
                    segments.append((None, pre, None, 0.0))
            next_start = splits[i + 1][0] if i + 1 < len(splits) else len(stripped)
            heading_end = _extend_heading_text(stripped, end, label_start=start)
            body = _normalize_spaces(stripped[heading_end:next_start].strip())
            from services.heading_detector import DEFAULT_HEADING_DETECTOR

            matched_label = stripped[start:end].strip()
            raw_heading = stripped[start:heading_end].strip()
            if heading_end > end and raw_heading.lower().startswith(matched_label.lower()):
                suffix = raw_heading[len(matched_label) :]
                heading_text = display_section_heading(matched_label, label_key) + suffix
            else:
                heading_text = display_section_heading(matched_label, label_key)
            peel = DEFAULT_HEADING_DETECTOR.split_embedded(stripped[start:next_start].strip())
            if peel is not None and peel.body:
                peel_key = normalize_paragraph_text(peel.heading.split(":")[0])
                if peel_key == label_key or peel_key.startswith(label_key):
                    heading_text = peel.heading
                    body = _normalize_spaces(peel.body)
            section = next((s for s in sections if s.canonical == label_key), None)
            conf = 0.95 if section and section.source == "requirement" else 0.85
            if heading_text:
                used.add(label_key)
                segments.append((heading_text, body, section, conf))
            pos = next_start
        if pos < len(stripped):
            tail = stripped[pos:].strip()
            if tail:
                if segments and segments[-1][0] and not segments[-1][1]:
                    h, _, sec, c = segments[-1]
                    segments[-1] = (h, tail, sec, c)
                else:
                    segments.append((None, tail, None, 0.0))
        return segments

    return []


def _match_section_for_text(text: str, sections: list[ExpectedSection]) -> ExpectedSection | None:
    norm = normalize_paragraph_text(text)
    for section in sections:
        if norm == section.canonical:
            return section
        for variant in _match_variants(section.label):
            if norm == _canonical_label_key(variant):
                return section
    return None


def _split_reference_paragraphs(blocks: list[ReconstructedBlock]) -> list[ReconstructedBlock]:
    """Ensure References heading is separate and each citation is its own paragraph."""
    out: list[ReconstructedBlock] = []
    in_refs = False
    for block in blocks:
        if block.kind == "heading" and is_references_heading(block.text):
            in_refs = True
            out.append(block)
            continue

        if in_refs and block.kind == "body":
            text = block.text.strip()
            if not text:
                continue
            m = re.match(r"^(references|bibliography|works cited)\s+", text, re.I)
            if m:
                out.append(
                    ReconstructedBlock(
                        kind="heading",
                        text=text[: m.end()].strip(),
                        level=2,
                        source="reconstructed",
                        confidence=0.9,
                        section_key="references",
                    )
                )
                text = text[m.end() :].strip()
                if not text:
                    continue
            parts = [p.strip() for p in text.split("\n") if p.strip()]
            if len(parts) == 1:
                parts = [text]
            for part in parts:
                part = part.strip()
                if not part:
                    continue
                kind = "reference" if _paragraph_is_reference_line(part) or _looks_like_citation_start(part) else "body"
                out.append(
                    ReconstructedBlock(
                        kind=kind,
                        text=part,
                        level=0,
                        source="reconstructed",
                        confidence=0.75 if kind == "reference" else 0.5,
                    )
                )
            continue

        out.append(block)
    return out


def _segments_to_blocks(
    segments: list[tuple[str | None, str, ExpectedSection | None, float]],
    *,
    para_index: int,
    is_title_candidate: bool,
) -> list[ReconstructedBlock]:
    blocks: list[ReconstructedBlock] = []
    for heading, body, section, conf in segments:
        if heading:
            level = section.level if section else 2
            heading_key = normalize_paragraph_text(heading)
            if section and section.canonical == "title":
                level = 1
            elif heading_key.startswith("body paragraph") or heading_key.startswith("journal entry"):
                level = 2
            elif para_index == 0 and is_title_candidate and _looks_like_title(heading) and not section:
                level = 1
            blocks.append(
                ReconstructedBlock(
                    kind="title" if level == 1 else "heading",
                    text=heading,
                    level=level,
                    source="reconstructed",
                    confidence=conf,
                    section_key=section.canonical if section else "",
                )
            )
        if body:
            from services.heading_detector import DEFAULT_HEADING_DETECTOR, HeadingDetector

            # Do not title-case orphan mid-sentence tokens.
            if not DEFAULT_HEADING_DETECTOR.is_forbidden_heading(body):
                body = HeadingDetector._ensure_capital_start(body)
            blocks.append(
                ReconstructedBlock(
                    kind="body",
                    text=body,
                    level=0,
                    source="reconstructed",
                    confidence=max(0.4, conf * 0.5),
                )
            )
    return blocks


def reconstruct_blocks(
    paragraphs: list[str],
    *,
    document_type: str | None = None,
    required_sections: list[str] | None = None,
    meta: list[dict[str, Any]] | None = None,
) -> list[ReconstructedBlock]:
    """Recover academic structure from unreliable paragraph boundaries."""
    expected = build_expected_structure(
        document_type=document_type,
        required_sections=required_sections,
        paragraphs=paragraphs,
    )
    label_list = [s.label for s in expected]
    explicit_keys = {_canonical_label_key(s.label) for s in expected}
    used: set[str] = set()
    blocks: list[ReconstructedBlock] = []

    para_count = len(paragraphs)
    for idx, text in enumerate(paragraphs):
        if not (text or "").strip():
            continue
        m = meta[idx] if meta and idx < len(meta) else {}
        word_heading = bool(m.get("is_word_heading"))
        word_level = m.get("heading_level")

        stripped = text.strip()
        from services.heading_detector import DEFAULT_HEADING_DETECTOR

        # Document titles stay intact — never scan them for in-sentence keywords.
        # Never classify section labels (Journal Entry / Body Paragraph / …) as the doc title.
        embedded = DEFAULT_HEADING_DETECTOR.split_embedded(stripped)
        sectionish = bool(
            re.match(
                r"^(?:Journal\s+Entry|Body\s+Paragraph|Introduction|Conclusion|Discussion|"
                r"References|Reflection|Abstract|Methodology|Methods|Results)\b",
                stripped,
                re.I,
            )
        )
        is_title = (
            idx == 0
            and embedded is None
            and not sectionish
            and (
                (word_heading and word_level == 1)
                or _looks_like_title(stripped)
            )
        )
        if is_title:
            blocks.append(
                ReconstructedBlock(
                    kind="title",
                    text=stripped,
                    level=1,
                    source="word_style" if word_heading else "reconstructed",
                    confidence=0.95 if word_heading else 0.85,
                    section_key="title",
                )
            )
            continue

        # Drop markdown leftovers like "## Document" that are not academic sections.
        from formatter.markdown_cleanup import strip_markdown_text

        cleaned = strip_markdown_text(stripped)
        if cleaned is None:
            continue
        if cleaned != stripped:
            stripped = cleaned
            text = cleaned
        if normalize_paragraph_text(stripped) == "document":
            continue

        req_segments = _segments_from_expected_labels(text, expected, used)
        if req_segments:
            blocks.extend(
                _segments_to_blocks(
                    req_segments,
                    para_index=idx,
                    is_title_candidate=False,
                )
            )
            continue

        segments = _segment_paragraph(
            text,
            expected,
            used,
            para_index=idx,
            para_count=para_count,
            explicit_keys=explicit_keys,
            word_heading=word_heading,
        )
        blocks.extend(
            _segments_to_blocks(
                segments,
                para_index=idx,
                is_title_candidate=False,
            )
        )

    return _merge_sentence_fragments(_split_reference_paragraphs(blocks))


_INCOMPLETE_TAIL = re.compile(
    r"(?i)\b(?:a|an|and|or|the|that|than|then|to|of|for|with|by|from|in|on|at|as|are|is|was|were|be|been|being)\s*$"
)


def _is_sentence_fragment(text: str, *, previous: str | None = None) -> bool:
    """True when ``text`` is an orphan mid-sentence piece, not a real paragraph."""
    from services.heading_detector import DEFAULT_HEADING_DETECTOR, _NEVER_HEADING_WORDS

    t = (text or "").strip()
    if not t:
        return True
    if DEFAULT_HEADING_DETECTOR.is_forbidden_heading(t):
        return True
    words = t.split()
    if len(words) == 1 and words[0][:1].islower():
        return True
    if previous:
        prev = previous.rstrip()
        if not prev:
            return False
        if _INCOMPLETE_TAIL.search(prev):
            return True
        if prev[-1] not in ".!?:;\"'”’)" and (
            t[:1].islower() or words[0].lower().strip(".,;:") in _NEVER_HEADING_WORDS
        ):
            return True
    return False


def _merge_sentence_fragments(blocks: list[ReconstructedBlock]) -> list[ReconstructedBlock]:
    """Rejoin mid-sentence orphans created by soft wraps or bad splits.

    A single sentence must never become multiple paragraphs. Isolated function
    words (are / the / as / …) are never kept as headings or standalone bodies.
    """
    from services.heading_detector import DEFAULT_HEADING_DETECTOR, _NEVER_HEADING_WORDS

    out: list[ReconstructedBlock] = []
    for block in blocks:
        text = (block.text or "").strip()
        if not text:
            continue

        kind = block.kind
        level = block.level
        if kind in {"heading", "title"} and DEFAULT_HEADING_DETECTOR.is_forbidden_heading(text):
            kind = "body"
            level = 0

        if out and kind in {"body", "reference"} and _is_sentence_fragment(text, previous=out[-1].text):
            prev = out[-1]
            piece = text
            if _INCOMPLETE_TAIL.search(prev.text.rstrip()):
                words = piece.split()
                if words and words[0][:1].isupper():
                    first = words[0]
                    bare = first.lower().strip(".,;:\"'“”‘’")
                    if bare in _NEVER_HEADING_WORDS:
                        piece = first.lower() + piece[len(first) :]
                    else:
                        piece = first[:1].lower() + piece[1:]
            elif words := piece.split():
                first = words[0]
                if first[:1].isupper() and first.lower().strip(".,;:") in _NEVER_HEADING_WORDS:
                    piece = first.lower() + piece[len(first) :]

            if prev.kind in {"heading", "title"} and _INCOMPLETE_TAIL.search(prev.text.rstrip()):
                # Prior block was misclassified as a heading mid-sentence.
                out[-1] = ReconstructedBlock(
                    kind="body",
                    text=_normalize_spaces(f"{prev.text.rstrip()} {piece}"),
                    level=0,
                    source="reconstructed",
                    confidence=0.2,
                )
                continue

            if prev.kind in {"body", "reference"}:
                out[-1] = ReconstructedBlock(
                    kind="reference" if prev.kind == "reference" else "body",
                    text=_normalize_spaces(f"{prev.text.rstrip()} {piece}"),
                    level=prev.level if prev.kind == "reference" else 0,
                    source=prev.source,
                    confidence=min(prev.confidence, 0.5),
                    section_key=prev.section_key,
                )
                continue

        out.append(
            ReconstructedBlock(
                kind=kind,
                text=text,
                level=level,
                source=block.source,
                confidence=block.confidence,
                section_key=block.section_key,
            )
        )
    return out


def _clear_document_body(doc: Document) -> None:
    body = doc.element.body
    for child in list(body):
        body.remove(child)


def apply_blocks_to_document(doc: Document, blocks: list[ReconstructedBlock]) -> list[ParagraphHeadingAssignment]:
    """Replace document body with reconstructed paragraphs."""
    _clear_document_body(doc)
    assignments: list[ParagraphHeadingAssignment] = []
    for block in blocks:
        text = block.text.strip()
        if not text:
            continue
        doc.add_paragraph(text)
        level = block.level if block.kind in {"title", "heading"} else None
        assignments.append(
            ParagraphHeadingAssignment(
                text=text,
                level=level,
                source=block.source if level else "none",
                confidence=block.confidence if level else None,
            )
        )
    return assignments


def _meta_from_doc(doc: Document) -> list[dict[str, Any]]:
    from services.document_structure_engine import _paragraph_meta_from_doc

    return _paragraph_meta_from_doc(doc)


def _document_needs_reconstruction(
    paragraphs: list[str],
    meta: list[dict[str, Any]] | None,
    *,
    document_type: str | None,
    required_sections: list[str] | None,
) -> bool:
    """True when paragraphs likely contain merged headings that must be split."""
    from services.document_structure_engine import headings_exist

    nonempty = [t for t in paragraphs if (t or "").strip()]
    if not nonempty:
        return False

    if headings_exist(nonempty, meta=meta) and all(
        (meta[i] if meta and i < len(meta) else {}).get("is_word_heading")
        or not _segments_from_expected_labels(
            text,
            build_expected_structure(
                document_type=document_type,
                required_sections=required_sections,
                paragraphs=nonempty,
            ),
            set(),
        )
        for i, text in enumerate(paragraphs)
        if (text or "").strip()
    ):
        return False

    expected = build_expected_structure(
        document_type=document_type,
        required_sections=required_sections,
        paragraphs=nonempty,
    )
    used: set[str] = set()
    for text in paragraphs:
        stripped = (text or "").strip()
        if not stripped:
            continue
        if _segments_from_expected_labels(stripped, expected, used):
            return True
        hit = _find_best_heading_in_text(
            stripped,
            expected,
            used,
            para_index=0,
            para_count=len(nonempty),
            explicit_keys={s.canonical for s in expected},
            word_heading=False,
        )
        if hit is not None:
            return True
    return False


def reconstruct_document_before_format(
    doc: Document,
    *,
    document_type: str | None = None,
    required_sections: list[str] | None = None,
    prefer_ai: bool = True,
) -> ReconstructionResult:
    """
    Full reconstruction pass — runs before formatting.

    1. Normalize spaces
    2. Try AI structure recovery when enabled
    3. Otherwise run multi-signal reconstruction engine
    4. Physically rebuild the document body
    """
    normalize_document_internal_spaces(doc)
    from formatter.markdown_cleanup import clean_markdown_in_document

    clean_markdown_in_document(doc)
    meta = _meta_from_doc(doc)
    paragraphs = [p.text for p in doc.paragraphs]

    if not _document_needs_reconstruction(
        paragraphs,
        meta,
        document_type=document_type,
        required_sections=required_sections,
    ):
        assignments = []
        for p in doc.paragraphs:
            text = (p.text or "").strip()
            if not text:
                continue
            from formatter.headings import heading_level_from_word_style

            lvl = heading_level_from_word_style(p)
            assignments.append(
                ParagraphHeadingAssignment(
                    text=text,
                    level=lvl,
                    source="word_style" if lvl else "none",
                )
            )
        return ReconstructionResult(
            assignments=assignments,
            recovery_mode="preserved",
            ai_powered=False,
        )

    if prefer_ai:
        recovery = recover_structure(
            doc=doc,
            document_type=document_type,
            prefer_ai=True,
        )
        if not recovery.get("error") and recovery.get("recovery_mode") == "ai_reconstructed":
            from formatter.structure_rebuild import rebuild_document_from_recovery

            apply_result = rebuild_document_from_recovery(doc, recovery)
            if apply_result:
                return ReconstructionResult(
                    assignments=apply_result.assignments,
                    recovery_mode=apply_result.recovery_mode,
                    ai_powered=apply_result.ai_powered,
                )

    blocks = reconstruct_blocks(
        paragraphs,
        document_type=document_type,
        required_sections=required_sections,
        meta=meta,
    )
    if not blocks:
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        blocks = [
            ReconstructedBlock(kind="body", text=t, level=0, source="none", confidence=0.0)
            for t in paragraphs
        ]

    assignments = apply_blocks_to_document(doc, blocks)
    return ReconstructionResult(
        assignments=assignments,
        recovery_mode="reconstructed",
        ai_powered=False,
        blocks=blocks,
    )
