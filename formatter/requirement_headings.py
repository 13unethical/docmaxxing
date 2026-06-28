"""
Split humanizer-merged section labels using assignment requirement headings.

When a brief lists Introduction, Body paragraph 1, etc., merged text like
"Introduction Artificial intelligence…" is split into separate heading + body
paragraphs without rewriting content. Each required section is used at most once.
"""

from __future__ import annotations

import re
from typing import Iterable

from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

from formatter.headings import split_embedded_heading_paragraph
from formatter.text_cleaning import collapse_internal_spaces
from services.requirements_parser import _extract_required_sections

_REF_LIST_SIGNAL = re.compile(
    r"\b(?:reference\s+list|bibliography|works\s+cited|references\s+section|"
    r"(?:include|provide|submit|add)\s+(?:a\s+)?(?:reference\s+list|references|bibliography)|"
    r"peer[-\s]?reviewed\s+(?:sources|articles|references))\b",
    re.I,
)
_CITATION_START = re.compile(
    r"^(?:[A-Z][A-Za-z\-']+(?:\s+et\s+al\.)?,\s*[A-Z]\.?|\[\d+\]|\d+\.\s+[A-Z])",
)

_BULLET_SECTION = re.compile(
    r"^[\s·•\-\*\u2022]+"
    r"(introduction|body\s+paragraph\s+\d+|concluding\s+paragraph|conclusion|"
    r"references|abstract|literature\s+review|methodology|discussion|results|"
    r"counterargument|counter-argument|background|analysis|summary|methods|"
    r"recommendations|appendix|executive\s+summary)\s*:?\s*$",
    re.I | re.M,
)
_STANDALONE_SECTION = re.compile(
    r"^(introduction|body\s+paragraph\s+\d+|concluding\s+paragraph|conclusion|"
    r"references|abstract|literature\s+review|methodology|discussion|results|"
    r"counterargument|counter-argument|background|analysis|summary|methods|"
    r"recommendations|appendix|executive\s+summary)\s*:?\s*$",
    re.I | re.M,
)
_INCLUDE_BLOCK = re.compile(
    r"\b(?:should|must)\s+include\b|\b(?:summary|essay|report|paper)\s+should\s+include\b",
    re.I,
)
_SECTION_LIST = re.compile(
    r"\b(?:required\s+sections?|section\s+headings?|structure|organiz(?:e|ation)|"
    r"must\s+contain|should\s+contain)\s*:?\s*([^\n.]{6,200})",
    re.I,
)

# Single-word labels that often appear inside body sentences when discussing structure.
_PROSE_SENSITIVE = frozenset(
    {
        "introduction",
        "conclusion",
        "conclusions",
        "discussion",
        "analysis",
        "background",
        "summary",
        "abstract",
        "methodology",
        "methods",
        "results",
        "recommendations",
    }
)

# Distinctive multi-word patterns — safe to split after whitespace when humanizers merge lines.
_DISTINCTIVE_LABEL = re.compile(
    r"^(body\s+paragraph\s+\d+|literature\s+review|executive\s+summary|"
    r"concluding\s+paragraph|works\s+cited|counterargument|counter-argument|"
    r"references|bibliography)$",
    re.I,
)

_SECTION_ALIASES: dict[str, tuple[str, ...]] = {
    "concluding paragraph": ("conclusion",),
    "conclusion": ("concluding paragraph",),
    "references": ("bibliography", "works cited"),
}


def _title_section(raw: str) -> str:
    t = re.sub(r"\s+", " ", raw.strip().lower())
    if t.startswith("body paragraph "):
        n = t.split()[-1]
        return f"Body paragraph {n}"
    if t == "concluding paragraph":
        return "Concluding Paragraph"
    return " ".join(w.capitalize() if w.isalpha() else w for w in t.split())


def _canonical_label_key(label: str) -> str:
    return re.sub(r"\s+", " ", label.strip().lower())


def extract_format_section_labels(text: str) -> list[str]:
    """Section labels from any assignment brief (bullets, lists, standalone lines)."""
    if not (text or "").strip():
        return []
    found: list[str] = []
    for pat in (_BULLET_SECTION, _STANDALONE_SECTION):
        for m in pat.finditer(text):
            found.append(_title_section(m.group(1)))
    if _INCLUDE_BLOCK.search(text):
        for m in _STANDALONE_SECTION.finditer(text):
            found.append(_title_section(m.group(1)))
    for m in _SECTION_LIST.finditer(text):
        chunk = m.group(1)
        for part in re.split(r",|;|/|\band\b", chunk):
            label = part.strip(" -:·")
            if 1 <= len(label.split()) <= 6 and not label.endswith("."):
                titled = _title_section(label) if label.islower() or len(label.split()) > 1 else label.strip()
                if len(titled) >= 3:
                    found.append(titled)
    found.extend(_extract_required_sections(text))
    if _REF_LIST_SIGNAL.search(text):
        found.append("References")
    deduped: list[str] = []
    seen: set[str] = set()
    for item in found:
        key = _canonical_label_key(item)
        if key and key not in seen:
            seen.add(key)
            deduped.append(item.strip())
    return deduped[:24]


def _match_variants(label: str) -> list[str]:
    norm = _canonical_label_key(label)
    variants = [label.strip()]
    for alias in _SECTION_ALIASES.get(norm, ()):
        variants.append(_title_section(alias))
    return variants


def _label_pattern(label: str) -> re.Pattern[str]:
    words = label.strip().split()
    inner = r"\s+".join(re.escape(w) for w in words)
    return re.compile(rf"(?i)\b{inner}\b")


def _following_text(text: str, end: int) -> str:
    return text[end:].lstrip()


def _normalize_spaces(text: str) -> str:
    return collapse_internal_spaces(text)


def _looks_like_citation_start(following: str) -> bool:
    return bool(_CITATION_START.match(following.strip()))


def _looks_like_real_section_start(label_key: str, following: str, *, explicit: bool = False) -> bool:
    """Body after a split should read like a new section, not mid-sentence prose."""
    following = following.strip()
    if not following:
        return True
    first = following.split()[0]
    if _DISTINCTIVE_LABEL.match(label_key):
        return True
    if label_key in ("references", "bibliography", "works cited"):
        return _looks_like_citation_start(following) or bool(re.match(r"[A-Z\[]", first))
    meta_starts = (
        "suddenly",
        "appears",
        "often",
        "sometimes",
        "typically",
        "usually",
        "may",
        "might",
        "can",
        "will",
        "should",
        "would",
        "is",
        "are",
        "was",
        "were",
        "has",
        "have",
    )
    if first.lower() in meta_starts:
        return False
    if explicit:
        return True
    if label_key in _PROSE_SENSITIVE:
        return first[0].isupper()
    return first[0].isupper() or _DISTINCTIVE_LABEL.match(label_key)


def _valid_split_at(
    text: str,
    start: int,
    end: int,
    label_key: str,
    *,
    explicit_labels: set[str] | None = None,
) -> bool:
    explicit = bool(explicit_labels and label_key in explicit_labels)
    following = _following_text(text, end)
    if not _looks_like_real_section_start(label_key, following, explicit=explicit):
        return False

    if start == 0:
        return True

    before = text[:start].rstrip()
    if not before:
        return True

    prev_char = before[-1]
    if prev_char in ".!?\n":
        return True

    if _DISTINCTIVE_LABEL.match(label_key):
        return True

    if label_key in ("references", "bibliography", "works cited"):
        return _looks_like_citation_start(following)

    if explicit:
        if start > 0 and text[start - 1].isspace():
            return True
        return prev_char in ".!?\n"

    if label_key in _PROSE_SENSITIVE:
        return False

    return False


def _collect_heading_splits(
    text: str,
    labels: Iterable[str],
    used_labels: set[str],
) -> list[tuple[int, int, str]]:
    """First valid occurrence of each label only; respects document-level used_labels."""
    explicit_keys = {_canonical_label_key(l) for l in labels}
    candidates: list[tuple[int, int, str, str, int]] = []
    for label in labels:
        label_key = _canonical_label_key(label)
        if label_key in used_labels:
            continue
        for variant in _match_variants(label):
            rx = _label_pattern(variant)
            for m in rx.finditer(text):
                if not _valid_split_at(
                    text, m.start(), m.end(), label_key, explicit_labels=explicit_keys
                ):
                    continue
                candidates.append(
                    (m.start(), m.end(), text[m.start() : m.end()].strip(), label_key, len(variant))
                )
                break

    candidates.sort(key=lambda x: (x[0], -x[4]))
    chosen: list[tuple[int, int, str, str]] = []
    last_end = -1
    for start, end, matched, label_key, _ in candidates:
        if start < last_end or label_key in used_labels:
            continue
        chosen.append((start, end, matched, label_key))
        last_end = end
    return chosen


def split_paragraph_by_requirement_headings(
    text: str,
    labels: list[str],
    used_labels: set[str] | None = None,
) -> list[tuple[str | None, str]]:
    """
    Split one paragraph into (heading, body) segments using requirement labels.
    Returns [(None, full_text)] when no split applies.
    """
    stripped = (text or "").strip()
    if not stripped or not labels:
        return [(None, stripped)]

    used = used_labels if used_labels is not None else set()
    explicit_keys = {_canonical_label_key(l) for l in labels}

    heading_line, body_after_newline = split_embedded_heading_paragraph(stripped)
    if body_after_newline is not None:
        key = _canonical_label_key(heading_line)
        if key in used:
            return [(None, stripped)]
        if _valid_split_at(stripped, 0, len(heading_line), key, explicit_labels=explicit_keys):
            used.add(key)
            return [(heading_line, _normalize_spaces(body_after_newline))]
        return [(None, stripped)]

    splits = _collect_heading_splits(stripped, labels, used)
    if not splits:
        return [(None, stripped)]

    for _, _, _, label_key in splits:
        used.add(label_key)

    segments: list[tuple[str | None, str]] = []
    pos = 0
    for i, (start, end, matched, _label_key) in enumerate(splits):
        if start > pos:
            pre = stripped[pos:start].strip()
            if pre:
                segments.append((None, pre))
        next_start = splits[i + 1][0] if i + 1 < len(splits) else len(stripped)
        body = _normalize_spaces(stripped[end:next_start].strip())
        segments.append((matched, body))
        pos = next_start
    if pos < len(stripped):
        tail = stripped[pos:].strip()
        if tail:
            if segments and segments[-1][0] and not segments[-1][1]:
                segments[-1] = (segments[-1][0], tail)
            else:
                segments.append((None, tail))
    return segments


def _insert_paragraph_after(paragraph: Paragraph, text: str) -> Paragraph:
    new_el = OxmlElement("w:p")
    paragraph._p.addnext(new_el)
    new_paragraph = Paragraph(new_el, paragraph._parent)
    if text:
        new_paragraph.add_run(text)
    return new_paragraph


def normalize_document_internal_spaces(document) -> None:
    """Collapse humanizer double-spaces in every paragraph."""
    from formatter.headings import set_plain_paragraph_text

    for paragraph in document.paragraphs:
        new = _normalize_spaces(paragraph.text)
        if new != paragraph.text:
            set_plain_paragraph_text(paragraph, new)


def expand_requirement_heading_paragraphs(document, labels: list[str]) -> int:
    """Split merged requirement headings; each required section at most once in the document."""
    if not labels:
        return 0
    used_labels: set[str] = set()
    inserted = 0
    idx = 0
    while idx < len(document.paragraphs):
        paragraph = document.paragraphs[idx]
        segments = split_paragraph_by_requirement_headings(
            paragraph.text, labels, used_labels=used_labels
        )
        if len(segments) == 1 and segments[0][0] is None:
            idx += 1
            continue

        anchor = paragraph
        first = True
        for heading, body in segments:
            if first:
                if heading:
                    paragraph.text = heading
                    if body:
                        anchor = _insert_paragraph_after(paragraph, body)
                        inserted += 1
                else:
                    paragraph.text = body
                first = False
                continue
            if heading:
                anchor = _insert_paragraph_after(anchor, heading)
                inserted += 1
            if body:
                anchor = _insert_paragraph_after(anchor, body)
                inserted += 1
        idx += 1
    normalize_document_internal_spaces(document)
    return inserted


def normalize_label_for_match(text: str, labels: list[str]) -> bool:
    """True if paragraph text equals a requirement section label."""
    norm = _canonical_label_key(text)
    for label in labels:
        for variant in _match_variants(label):
            if norm == _canonical_label_key(variant):
                return True
    return False
